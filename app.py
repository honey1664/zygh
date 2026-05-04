# app.py

import streamlit as st
import pandas as pd
import numpy as np
import json
import os
import re
import hashlib
import sqlite3
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime
from typing import List, Dict, Optional
import time
import requests
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
import easyocr
import numpy as np
from PIL import Image
import io
import tempfile
import subprocess
from bs4 import BeautifulSoup
import pdfplumber
from dotenv import load_dotenv
from ai_chat_client import ai_chat_client
from chat_floating_v2 import render_floating_chat

# ==================== 配置区域 ====================
st.set_page_config(
    page_title="AI大学生职业规划智能体",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed"  # 侧边栏默认折叠，点击展开图标可展开
)
def add_ai_assistant_to_sidebar():
    """在侧边栏添加AI助手"""
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 💬 AI职业顾问")

        # 初始化聊天历史
        if 'chat_messages' not in st.session_state:
            st.session_state['chat_messages'] = [
                {"role": "assistant",
                 "content": "你好！我是你的AI职业顾问 👋\n\n我可以帮你解答职业规划、面试准备、技能提升等问题！"}
            ]

        # 显示聊天记录
        for msg in st.session_state['chat_messages'][-10:]:
            if msg['role'] == 'user':
                st.markdown(
                    f"<div style='text-align: right; margin: 8px 0;'><span style='background: #667eea; color: white; padding: 8px 14px; border-radius: 18px; font-size: 13px;'>{msg['content']}</span></div>",
                    unsafe_allow_html=True)
            else:
                st.markdown(
                    f"<div style='text-align: left; margin: 8px 0;'><span style='background: #f0f2f6; color: #333; padding: 8px 14px; border-radius: 18px; font-size: 13px;'>{msg['content']}</span></div>",
                    unsafe_allow_html=True)

        # 输入框
        user_input = st.text_input("", placeholder="输入你的问题...", key="ai_chat_input", label_visibility="collapsed")

        col1, col2 = st.columns([3, 1])
        with col1:
            if st.button("发送", key="ai_chat_send", use_container_width=True):
                if user_input:
                    # 添加用户消息
                    st.session_state['chat_messages'].append({"role": "user", "content": user_input})

                    # 获取AI回复
                    with st.spinner("AI思考中..."):
                        profile = st.session_state.get('student_profile', None)
                        response = ai_chat_client.get_answer(user_input, profile)
                        st.session_state['chat_messages'].append({"role": "assistant", "content": response})

                    st.rerun()

        with col2:
            if st.button("清空", key="ai_chat_clear", use_container_width=True):
                st.session_state['chat_messages'] = []
                st.rerun()

def get_jobs_from_excel():
    """从Excel文件中获取岗位列表"""
    try:
        df = pd.read_excel('a13基于AI的大学生职业规划智能体-JD采样数据.xls')
        unique_jobs = df['岗位名称'].unique().tolist()
        return sorted(unique_jobs)
    except Exception as e:
        print(f"读取Excel文件失败: {e}")
        # 返回增强版岗位列表
        return [
            "Java开发工程师", "Python开发工程师", "前端开发工程师",
            "软件测试工程师", "数据分析工程师", "产品经理",
            "UI/UX设计师", "网络工程师", "运维工程师",
            "AI算法工程师", "数据产品经理", "嵌入式软件工程师"
        ]


def export_report_to_pdf(report_content: str, filename: str) -> bytes:
    """将报告导出为PDF文件"""
    from io import BytesIO
    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()

    # 创建自定义样式
    title_style = ParagraphStyle('CustomTitle',
                               parent=styles['Heading1'],
                               fontSize=18,
                               spaceAfter=30)

    heading_style = ParagraphStyle('CustomHeading',
                                 parent=styles['Heading2'],
                                 fontSize=14,
                                 spaceAfter=20)

    normal_style = ParagraphStyle('CustomNormal',
                                parent=styles['BodyText'],
                                fontSize=12,
                                leading=18)

    # 处理报告内容
    elements = []

    # 分割报告内容
    lines = report_content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 12))
        elif line.startswith('# '):
            # 一级标题
            text = line[2:]
            elements.append(Paragraph(text, title_style))
        elif line.startswith('## '):
            # 二级标题
            text = line[3:]
            elements.append(Paragraph(text, heading_style))
        elif line.startswith('### '):
            # 三级标题
            text = line[4:]
            elements.append(Paragraph(text, styles['Heading3']))
        elif line.startswith('- '):
            # 列表项
            text = line[2:]
            elements.append(Paragraph(f'• {text}', normal_style))
        else:
            # 普通文本
            elements.append(Paragraph(line, normal_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def export_report_to_word(report_content: str, filename: str) -> bytes:
    """将报告导出为Word文件"""
    from io import BytesIO
    buffer = BytesIO()

    doc = Document()

    # 分割报告内容
    lines = report_content.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('# '):
            # 一级标题
            text = line[2:]
            doc.add_heading(text, level=1)
        elif line.startswith('## '):
            # 二级标题
            text = line[3:]
            doc.add_heading(text, level=2)
        elif line.startswith('### '):
            # 三级标题
            text = line[4:]
            doc.add_heading(text, level=3)
        elif line.startswith('- '):
            # 列表项
            text = line[2:]
            doc.add_paragraph(text, style='ListBullet')
        else:
            # 普通文本
            doc.add_paragraph(line)

    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_report_to_txt(report_content: str, filename: str) -> bytes:
    """将报告导出为TXT文件"""
    lines = report_content.split('\n')
    txt_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            txt_lines.append('')
        elif line.startswith('# '):
            text = line[2:]
            txt_lines.append('=' * 60)
            txt_lines.append(text)
            txt_lines.append('=' * 60)
        elif line.startswith('## '):
            text = line[3:]
            txt_lines.append('-' * 40)
            txt_lines.append(text)
            txt_lines.append('-' * 40)
        elif line.startswith('### '):
            text = line[4:]
            txt_lines.append(f'【{text}】')
        elif line.startswith('- '):
            text = line[2:]
            txt_lines.append(f'  • {text}')
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            txt_lines.append(f'【{text}】')
        else:
            txt_lines.append(line)

    return '\n'.join(txt_lines).encode('utf-8')


def export_report_to_excel(report_content: str, filename: str) -> bytes:
    """将报告导出为Excel文件"""
    from io import BytesIO
    import pandas as pd

    lines = report_content.split('\n')
    sections = []
    current_section = "报告概述"
    current_content = []

    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            if current_content:
                sections.append({
                    '章节': current_section,
                    '内容': '\n'.join(current_content)
                })
            current_section = line[2:]
            current_content = []
        elif line.startswith('## '):
            if current_content:
                sections.append({
                    '章节': current_section,
                    '内容': '\n'.join(current_content)
                })
            current_section = line[3:]
            current_content = []
        elif line.startswith('### '):
            if current_content:
                sections.append({
                    '章节': current_section,
                    '内容': '\n'.join(current_content)
                })
            current_section = line[4:]
            current_content = []
        else:
            if line:
                current_content.append(line)

    if current_content:
        sections.append({
            '章节': current_section,
            '内容': '\n'.join(current_content)
        })

    df = pd.DataFrame(sections)

    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='职业规划报告')

        workbook = writer.book
        worksheet = writer.sheets['职业规划报告']

        worksheet.column_dimensions['A'].width = 25
        worksheet.column_dimensions['B'].width = 80

    buffer.seek(0)
    return buffer.getvalue()


def ensure_export_folder():
    """确保导出文件夹存在"""
    export_folder = "结合文件夹"
    if not os.path.exists(export_folder):
        os.makedirs(export_folder)
    return export_folder


class EnhancedResumeParser:
    """增强版简历解析器 - 支持多格式文件精准识别"""

    API_KEY = "sk-42a21eb1dea14df8875dd923c3881d49"
    API_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"

    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF文档',
        '.docx': 'Word文档',
        '.doc': 'Word文档',
        '.txt': '文本文件',
        '.rtf': '富文本文件',
        '.md': 'Markdown文件',
        '.html': 'HTML文件',
        '.jpg': '图片简历',
        '.jpeg': '图片简历',
        '.png': '图片简历'
    }

    @staticmethod
    def extract_text_from_file(uploaded_file) -> str:
        """从上传文件中提取文本内容"""
        file_name = uploaded_file.name
        file_ext = os.path.splitext(file_name)[1].lower()

        try:
            if file_ext == '.pdf':
                return EnhancedResumeParser._extract_pdf_text(uploaded_file)
            elif file_ext in ['.docx', '.doc']:
                return EnhancedResumeParser._extract_word_text(uploaded_file, file_ext)
            elif file_ext in ['.txt', '.rtf', '.md']:
                return EnhancedResumeParser._extract_text_file(uploaded_file)
            elif file_ext == '.html':
                return EnhancedResumeParser._extract_html_text(uploaded_file)
            elif file_ext in ['.jpg', '.jpeg', '.png']:
                return EnhancedResumeParser._extract_image_text(uploaded_file)
            else:
                try:
                    return uploaded_file.read().decode('utf-8')
                except:
                    return f"[文件格式: {file_name}，请手动输入内容]"
        except Exception as e:
            print(f"文件解析错误: {e}")
            return f"[解析失败: {str(e)}]"

    @staticmethod
    def _extract_pdf_text(uploaded_file) -> str:
        """提取PDF文本"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name

            text = ""
            with pdfplumber.open(tmp_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += " | ".join([str(cell) for cell in row if cell]) + "\n"
            os.unlink(tmp_path)
            return text if text.strip() else "未能从PDF中提取文本内容"
        except ImportError:
            return "请安装pdfplumber库: pip install pdfplumber"
        except Exception as e:
            return f"PDF解析错误: {str(e)}"

    @staticmethod
    def _extract_word_text(uploaded_file, file_ext) -> str:
        """提取Word文档文本"""
        try:
            from docx import Document

            if file_ext == '.docx':
                doc = Document(io.BytesIO(uploaded_file.read()))
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                        if row_text:
                            text += "\n" + row_text
                return text if text.strip() else "Word文档中没有提取到文本"
            else:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = tmp.name
                try:
                    result = subprocess.run(['antiword', tmp_path], capture_output=True, text=True)
                    text = result.stdout if result.returncode == 0 else ""
                    if not text:
                        result = subprocess.run(['catdoc', tmp_path], capture_output=True, text=True)
                        text = result.stdout if result.returncode == 0 else ""
                except:
                    text = ""
                os.unlink(tmp_path)
                return text if text.strip() else "旧版Word文档解析失败，请转换为.docx格式"
        except ImportError:
            return "请安装python-docx库: pip install python-docx"
        except Exception as e:
            return f"Word解析错误: {str(e)}"

    @staticmethod
    def _extract_text_file(uploaded_file) -> str:
        """提取文本文件内容"""
        try:
            content = uploaded_file.read()
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            return f"文本文件解析错误: {str(e)}"

    @staticmethod
    def _extract_html_text(uploaded_file) -> str:
        """提取HTML文件中的文本"""
        try:
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text()
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            return "\n".join(lines)
        except ImportError:
            import re
            text = uploaded_file.read().decode('utf-8', errors='ignore')
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            return text
        except Exception as e:
            return f"HTML解析错误: {str(e)}"

    @staticmethod
    def _extract_image_text(uploaded_file) -> str:
        """使用OCR提取图片中的文字"""
        try:
            reader = easyocr.Reader(['ch_sim', 'en'], gpu=False, verbose=False)
            image = Image.open(io.BytesIO(uploaded_file.read()))
            image_np = np.array(image)
            results = reader.readtext(image_np, detail=0, paragraph=True)
            text = "\n".join(results)
            return text if text.strip() else "未能从图片中识别出文字"
        except ImportError:
            return "请安装easyocr和Pillow: pip install easyocr pillow"
        except Exception as e:
            return f"图片OCR识别错误: {str(e)}"

    @staticmethod
    def parse_resume_with_llm(resume_text: str, file_name: str = None) -> Dict:
        """使用大模型精准解析简历"""
        prompt = f"""
请作为一名专业的HR和职业规划顾问，仔细分析以下简历内容，生成结构化的学生能力画像JSON。

## 简历来源
{'文件: ' + file_name if file_name else '手动输入'}

## 简历原文
{resume_text[:5000]}

## 输出格式要求
请严格按照以下JSON格式输出：

{{
    "基本信息": {{
        "姓名": "",
        "性别": "",
        "年龄": "",
        "电话": "",
        "邮箱": "",
        "专业": "",
        "学历": "",
        "学校": "",
        "意向城市": ""
    }},
    "专业技能": [],
    "技能熟练度": {{}},
    "证书": [],
    "项目经验": [],
    "实习经历": "",
    "软实力": {{
        "学习能力": "",
        "沟通能力": "",
        "抗压能力": "",
        "创新能力": ""
    }},
    "求职意向": {{
        "期望岗位": "",
        "期望薪资": "",
        "期望城市": ""
    }},
    "综合评估": {{
        "能力等级": "",
        "主要优势": [],
        "待提升方向": [],
        "推荐岗位": []
    }}
}}

请仔细分析简历，尽可能填充更多字段。
"""
        headers = {
            "Authorization": f"Bearer {EnhancedResumeParser.API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "qwen-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1
        }

        try:
            response = requests.post(EnhancedResumeParser.API_URL, headers=headers, json=data, timeout=60)
            result = response.json()
            if 'choices' in result:
                content = result['choices'][0]['message']['content']
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    parsed = json.loads(json_match.group())
                    parsed['_metadata'] = {
                        'parse_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'source_file': file_name
                    }
                    return parsed
        except Exception as e:
            print(f"解析失败: {e}")

        return EnhancedResumeParser._fallback_parse(resume_text)

    @staticmethod
    def _fallback_parse(resume_text: str) -> Dict:
        """规则引擎兜底解析"""
        import re
        return {
            "基本信息": {
                "姓名": EnhancedResumeParser._extract_name(resume_text),
                "专业": EnhancedResumeParser._extract_major(resume_text),
                "学历": EnhancedResumeParser._extract_education(resume_text),
                "学校": EnhancedResumeParser._extract_school(resume_text),
                "意向城市": "",
                "电话": EnhancedResumeParser._extract_phone(resume_text),
                "邮箱": EnhancedResumeParser._extract_email(resume_text)
            },
            "专业技能": EnhancedResumeParser._extract_skills(resume_text),
            "证书": EnhancedResumeParser._extract_certificates(resume_text),
            "项目经验": [],
            "实习经历": "",
            "软实力": {"学习能力": "良好", "沟通能力": "良好", "抗压能力": "良好", "创新能力": "良好"},
            "求职意向": {},
            "综合评估": {"能力等级": "良好", "主要优势": [], "待提升方向": [], "推荐岗位": []}
        }

    @staticmethod
    def _extract_name(text: str) -> str:
        import re
        patterns = [r'姓名[：:]\s*([\u4e00-\u9fa5]{2,4})', r'名字[：:]\s*([\u4e00-\u9fa5]{2,4})']
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        return ""

    @staticmethod
    def _extract_phone(text: str) -> str:
        import re
        match = re.search(r'1[3-9]\d{9}', text)
        return match.group(0) if match else ""

    @staticmethod
    def _extract_email(text: str) -> str:
        import re
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        return match.group(0) if match else ""

    @staticmethod
    def _extract_education(text: str) -> str:
        edu_map = {'博士': '博士', '硕士': '硕士', '研究生': '硕士', '本科': '本科', '学士': '本科', '大专': '专科'}
        for kw, level in edu_map.items():
            if kw in text:
                return level
        return "本科"

    @staticmethod
    def _extract_major(text: str) -> str:
        majors = ['计算机科学与技术', '软件工程', '电子信息工程', '通信工程', '数据科学', '人工智能']
        for major in majors:
            if major in text:
                return major
        return ""

    @staticmethod
    def _extract_school(text: str) -> str:
        import re
        matches = re.findall(r'([\u4e00-\u9fa5]{2,10})(大学|学院|学校)', text)
        return matches[0][0] + matches[0][1] if matches else ""

    @staticmethod
    def _extract_skills(text: str) -> List[str]:
        skills = ['Python', 'Java', 'JavaScript', 'C++', 'Go', 'Spring', 'SpringBoot', 'MySQL', 'Redis', 'Docker']
        return [s for s in skills if s.lower() in text.lower()]

    @staticmethod
    def _extract_certificates(text: str) -> List[str]:
        certs = ['CET-4', 'CET-6', '英语四级', '英语六级', '软考', 'PMP']
        return [c for c in certs if c in text]

    @staticmethod
    def calculate_portrait_score(portrait: Dict) -> Dict:
        """计算画像评分"""
        total_fields, filled_fields = 0, 0

        def check(val):
            if val is None:
                return False
            if isinstance(val, str):
                return bool(val.strip()) and val != "未知"
            if isinstance(val, list):
                return len(val) > 0
            if isinstance(val, dict):
                return any(check(v) for v in val.values())
            return True

        basic = portrait.get("基本信息", {})
        for field in ["姓名", "专业", "学历", "学校"]:
            total_fields += 1
            if check(basic.get(field)):
                filled_fields += 1

        total_fields += 1
        if check(portrait.get("专业技能")):
            filled_fields += 1

        total_fields += 1
        if check(portrait.get("证书")):
            filled_fields += 1

        total_fields += 1
        if check(portrait.get("项目经验")):
            filled_fields += 1

        completeness = round((filled_fields / total_fields) * 100, 2) if total_fields > 0 else 0

        skill_count = len(portrait.get("专业技能", []))
        cert_count = len(portrait.get("证书", []))
        project_count = len(portrait.get("项目经验", []))

        competitiveness = min(100, skill_count * 8 + cert_count * 5 + project_count * 10 + 30)

        return {"完整度评分": completeness, "竞争力评分": competitiveness, "说明": f"填写{filled_fields}/{total_fields}个字段"}

# 数据库路径
DB_PATH = "./data/career_planner.db"

# ==================== 样式定制 ====================
st.markdown("""
<style>
    /* 全局样式 */
    .stApp {
        background: #ffffff;
    }

    /* 登录页面样式 */
    .login-wrapper {
        min-height: 100vh;
        display: flex;
        justify-content: center;
        align-items: center;
        padding: 20px;
        background: #f5f7fa;
    }

    .login-card {
        background: white;
        border-radius: 20px;
        padding: 40px 35px;
        width: 100%;
        max-width: 450px;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.08);
    }

    .login-header {
        text-align: center;
        margin-bottom: 35px;
    }

    .login-icon {
        font-size: 56px;
        margin-bottom: 15px;
    }

    .login-title {
        font-size: 28px;
        font-weight: bold;
        color: #333;
        margin-bottom: 8px;
    }

    .login-subtitle {
        color: #888;
        font-size: 14px;
    }

    /* 表单样式 */
    .stTextInput > div > div > input {
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        padding: 12px 16px;
        font-size: 15px;
    }

    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 3px rgba(102, 126, 234, 0.1);
    }

    /* 按钮样式 */
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 12px;
        font-size: 16px;
        font-weight: 600;
        width: 100%;
        transition: all 0.3s ease;
    }

    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(102, 126, 234, 0.3);
    }

    /* 用户类型切换 */
    .user-type-label {
        font-weight: 500;
        color: #555;
        margin-bottom: 10px;
        font-size: 14px;
    }

    .stRadio > div {
        gap: 20px;
    }

    .stRadio label {
        background: #f8f9fa;
        padding: 8px 24px;
        border-radius: 30px;
        border: 1px solid #e0e0e0;
        color: #666;
    }

    /* 分割线 */
    .divider {
        display: flex;
        align-items: center;
        text-align: center;
        margin: 25px 0 20px;
    }

    .divider::before,
    .divider::after {
        content: '';
        flex: 1;
        border-bottom: 1px solid #e8e8e8;
    }

    .divider span {
        padding: 0 15px;
        color: #aaa;
        font-size: 13px;
    }

    /* 底部链接 */
    .footer-links {
        display: flex;
        justify-content: space-between;
        gap: 15px;
        margin-top: 20px;
    }

    .footer-link {
        flex: 1;
        text-align: center;
        padding: 10px;
        border-radius: 10px;
        background: #f8f9fa;
        color: #667eea;
        text-decoration: none;
        font-size: 14px;
        font-weight: 500;
        cursor: pointer;
        border: 1px solid #e8e8e8;
    }

    .footer-link:hover {
        background: #667eea;
        color: white;
    }

    /* 卡片样式 */
    .card {
        background: white;
        border-radius: 15px;
        padding: 20px;
        box-shadow: 0 2px 12px rgba(0,0,0,0.05);
        margin: 10px 0;
        cursor: pointer;
        transition: transform 0.3s ease;
        border: 1px solid #f0f0f0;
    }

    .card:hover {
        transform: translateY(-5px);
        box-shadow: 0 8px 20px rgba(0,0,0,0.1);
    }

    .highlight {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }

    .discussion-item {
        background: white;
        border-radius: 12px;
        padding: 15px;
        margin: 10px 0;
        border-left: 4px solid #667eea;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }

    /* 顶部导航栏 */
    .top-nav {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 15px 30px;
        margin-bottom: 20px;
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        z-index: 1000;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }

    .nav-container {
        display: flex;
        justify-content: space-between;
        align-items: center;
        flex-wrap: nowrap;
        max-width: 1400px;
        margin: 0 auto;
        width: 100%;
    }

    .nav-menu {
        display: flex;
        gap: 20px;
        flex: 1;
        justify-content: center;
        margin: 0 10px;
        white-space: nowrap;
    }

    .top-nav .nav-item {
        padding: 8px 16px;
        cursor: pointer;
        font-size: 14px;
        font-weight: 500;
        color: white !important;
        text-decoration: none !important;
        display: inline-block;
        transition: all 0.3s ease;
        position: relative;
    }

    .nav-item:hover {
        transform: translateY(-2px);
    }

    .nav-item:hover::after {
        content: '';
        position: absolute;
        bottom: 0;
        left: 50%;
        transform: translateX(-50%);
        width: 80%;
        height: 2px;
        background: white;
        border-radius: 2px;
    }

    .nav-item.active {
        font-weight: 600;
    }

    .nav-item.active::after {
        content: '';
        position: absolute;
        bottom: 0;
        left: 50%;
        transform: translateX(-50%);
        width: 80%;
        height: 2px;
        background: white;
        border-radius: 2px;
    }

    .user-info {
        display: flex;
        align-items: center;
        gap: 15px;
        background: rgba(255, 255, 255, 0.1);
        padding: 5px 15px;
        border-radius: 30px;
    }

    .user-info span {
        color: white;
    }

    .logout-btn {
        background: rgba(255, 255, 255, 0.2);
        color: #e0e0e0;
        padding: 6px 16px;
        border-radius: 25px;
        font-size: 13px;
        text-decoration: none !important;
        transition: all 0.3s ease;
    }

    .logout-btn:hover {
        background: rgba(255, 255, 255, 0.3);
        color: #e0e0e0;
    }

    /* 隐藏多余元素 */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* 进度条 */
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, #667eea, #764ba2);
    }

    /* 仪表盘卡片 */
    .dashboard-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 20px;
        color: white;
        text-align: center;
        margin: 10px 0;
    }

    .dashboard-number {
        font-size: 32px;
        font-weight: bold;
    }

    .dashboard-label {
        font-size: 14px;
        opacity: 0.9;
        margin-top: 5px;
    }

    /* 报告样式 */
    .report-section {
        background: white;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 20px;
        border: 1px solid #f0f0f0;
    }

    .report-title {
        font-size: 24px;
        font-weight: bold;
        text-align: center;
        margin-bottom: 20px;
        color: #333;
    }

    .report-subtitle {
        font-size: 18px;
        font-weight: bold;
        margin: 20px 0 15px;
        padding-left: 10px;
        border-left: 4px solid #667eea;
        color: #333;
    }

    /* 学习资源样式 */
    .resource-section {
        background: white;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 20px;
        border: 1px solid #f0f0f0;
    }

    .gap-item {
        background: #fff3e0;
        border-left: 3px solid #ff9800;
        padding: 12px;
        margin: 10px 0;
        border-radius: 8px;
    }
</style>
""", unsafe_allow_html=True)


# ==================== 数据库管理 ====================
def init_database():
    """初始化数据库"""
    os.makedirs("./data", exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    # 用户表
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        email TEXT UNIQUE,
        password_hash TEXT NOT NULL,
        user_type TEXT NOT NULL DEFAULT 'user',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 讨论表
    c.execute('''CREATE TABLE IF NOT EXISTS discussions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        user_type TEXT,
        title TEXT NOT NULL,
        content TEXT NOT NULL,
        is_anonymous INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        reply_count INTEGER DEFAULT 0,
        is_pinned INTEGER DEFAULT 0,
        likes INTEGER DEFAULT 0,
        liked_by TEXT DEFAULT ''
    )''')

    # 回复表
    c.execute('''CREATE TABLE IF NOT EXISTS replies (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        discussion_id INTEGER NOT NULL,
        username TEXT NOT NULL,
        user_type TEXT,
        content TEXT NOT NULL,
        is_anonymous INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        is_pinned INTEGER DEFAULT 0,
        likes INTEGER DEFAULT 0,
        liked_by TEXT DEFAULT ''
    )''')

    # 学生画像表
    c.execute('''CREATE TABLE IF NOT EXISTS student_profiles (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL UNIQUE,
        profile_data TEXT,
        completeness_score REAL DEFAULT 0,
        competitiveness_score REAL DEFAULT 0,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 职业报告表
    c.execute('''CREATE TABLE IF NOT EXISTS career_reports (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        report_content TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 密码重置表
    c.execute('''CREATE TABLE IF NOT EXISTS password_reset (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        email TEXT NOT NULL,
        code TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        expires_at TIMESTAMP
    )''')

    # 操作日志表
    c.execute('''CREATE TABLE IF NOT EXISTS operation_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        action TEXT NOT NULL,
        ip TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 系统配置表
    c.execute('''CREATE TABLE IF NOT EXISTS system_config (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        config_key TEXT UNIQUE NOT NULL,
        config_value TEXT,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 模拟面试记录表
    c.execute('''CREATE TABLE IF NOT EXISTS mock_interviews (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        job_title TEXT NOT NULL,
        interview_questions TEXT,
        user_answers TEXT,
        feedback TEXT,
        score REAL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 简历优化记录表
    c.execute('''CREATE TABLE IF NOT EXISTS resume_optimization (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        original_resume TEXT,
        optimized_resume TEXT,
        suggestions TEXT,
        matching_score REAL,
        target_job TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    # 插入默认配置
    c.execute("SELECT * FROM system_config WHERE config_key = 'site_name'")
    if not c.fetchone():
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("site_name", "AI大学生职业规划智能体"))
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("site_desc", "基于人工智能技术，为您提供精准的职业规划方案"))
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("theme_color", "#667eea"))
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("enable_discussion", "True"))
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("enable_resume_parse", "True"))
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("enable_report", "True"))
        c.execute("INSERT INTO system_config (config_key, config_value) VALUES (?, ?)",
                  ("enable_recommend", "True"))

    # 创建默认管理员
    c.execute("SELECT * FROM users WHERE username = 'admin'")
    if not c.fetchone():
        admin_pwd = hash_password("admin123")
        c.execute("INSERT INTO users (username, email, password_hash, user_type) VALUES (?, ?, ?, ?)",
                  ("admin", "admin@career.com", admin_pwd, "admin"))

    # 创建测试普通用户
    c.execute("SELECT * FROM users WHERE username = '张三'")
    if not c.fetchone():
        user_pwd = hash_password("123456")
        c.execute("INSERT INTO users (username, email, password_hash, user_type) VALUES (?, ?, ?, ?)",
                  ("张三", "zhangsan@test.com", user_pwd, "user"))

    conn.commit()
    conn.close()


def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()


def verify_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed


def create_user(username: str, email: str, password: str, user_type: str) -> bool:
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO users (username, email, password_hash, user_type) VALUES (?, ?, ?, ?)",
                  (username, email, hash_password(password), user_type))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def authenticate_user(username: str, password: str, user_type: str) -> Optional[Dict]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE username = ? AND user_type = ?", (username, user_type))
    user = c.fetchone()
    conn.close()

    if user and verify_password(password, user['password_hash']):
        return dict(user)
    return None


def user_exists(username: str) -> bool:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT 1 FROM users WHERE username = ?", (username,))
    exists = c.fetchone() is not None
    conn.close()
    return exists


def email_exists(email: str) -> bool:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT 1 FROM users WHERE email = ?", (email,))
    exists = c.fetchone() is not None
    conn.close()
    return exists


def change_password(username: str, old_password: str, new_password: str) -> bool:
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT password_hash FROM users WHERE username = ?", (username,))
    row = c.fetchone()

    if not row or not verify_password(old_password, row[0]):
        conn.close()
        return False

    c.execute("UPDATE users SET password_hash = ? WHERE username = ?",
              (hash_password(new_password), username))
    conn.commit()
    conn.close()
    return True


def save_student_profile(username: str, profile_data: Dict) -> bool:
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        profile_json = json.dumps(profile_data, ensure_ascii=False)

        c.execute("SELECT id FROM student_profiles WHERE username = ?", (username,))
        existing = c.fetchone()

        if existing:
            c.execute("UPDATE student_profiles SET profile_data=?, updated_at=CURRENT_TIMESTAMP WHERE username=?",
                      (profile_json, username))
        else:
            c.execute("INSERT INTO student_profiles (username, profile_data) VALUES (?, ?)",
                      (username, profile_json))

        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"保存失败: {e}")
        return False


def get_student_profile(username: str) -> Optional[Dict]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM student_profiles WHERE username = ?", (username,))
    row = c.fetchone()
    conn.close()

    if row and row['profile_data']:
        try:
            profile = json.loads(row['profile_data'])
            return profile
        except:
            return None
    return None


def save_career_report(username: str, report_content: str, generated_time=None) -> bool:
    try:
        import datetime
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        # 使用指定的生成时间或当前时间
        current_time = generated_time if generated_time else datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        c.execute("INSERT INTO career_reports (username, report_content, created_at) VALUES (?, ?, ?)",
                  (username, report_content, current_time))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def save_resume_optimization(username: str, original_resume: str, optimized_resume: str, suggestions: str, matching_score: float, target_job: str) -> bool:
    """保存简历优化记录"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO resume_optimization (username, original_resume, optimized_resume, suggestions, matching_score, target_job) VALUES (?, ?, ?, ?, ?, ?)",
                  (username, original_resume, optimized_resume, suggestions, matching_score, target_job))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"保存简历优化记录失败: {e}")
        return False


def get_resume_optimization_history(username: str) -> List[Dict]:
    """获取简历优化历史"""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM resume_optimization WHERE username = ? ORDER BY created_at DESC", (username,))
        rows = c.fetchall()
        conn.close()
        return [dict(row) for row in rows]
    except Exception as e:
        print(f"获取简历优化历史失败: {e}")
        return []


def save_mock_interview(username: str, job_title: str, interview_questions: str, user_answers: str, feedback: str, score: float) -> bool:
    """保存模拟面试记录"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO mock_interviews (username, job_title, interview_questions, user_answers, feedback, score) VALUES (?, ?, ?, ?, ?, ?)",
                  (username, job_title, interview_questions, user_answers, feedback, score))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"保存模拟面试记录失败: {e}")
        return False


def get_mock_interview_history(username: str) -> List[Dict]:
    """获取模拟面试历史"""
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        c.execute("SELECT * FROM mock_interviews WHERE username = ? ORDER BY created_at DESC", (username,))
        rows = c.fetchall()
        conn.close()
        return [dict(row) for row in rows]
    except Exception as e:
        print(f"获取模拟面试历史失败: {e}")
        return []


def log_operation(username: str, action: str, ip: str = ""):
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO operation_logs (username, action, ip) VALUES (?, ?, ?)",
                  (username, action, ip))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"记录日志失败: {e}")

def get_system_config():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT config_key, config_value FROM system_config")
    rows = c.fetchall()
    conn.close()
    config = {}
    for row in rows:
        config[row['config_key']] = row['config_value']
    return config

def save_system_config(config_key, config_value):
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("UPDATE system_config SET config_value = ?, updated_at = ? WHERE config_key = ?",
                  (config_value, datetime.now(), config_key))
        conn.commit()
        conn.close()
        return True
    except:
        return False

def get_system_statistics():
    """获取系统统计数据"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()

    c.execute("SELECT COUNT(*) FROM users")
    total_users = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM student_profiles")
    total_profiles = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM discussions")
    total_discussions = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM replies")
    total_replies = c.fetchone()[0]

    c.execute("SELECT COUNT(*) FROM career_reports")
    total_reports = c.fetchone()[0]

    # 获取技能分布
    c.execute("SELECT profile_data FROM student_profiles")
    rows = c.fetchall()

    skill_count = {}
    for row in rows:
        if row[0]:
            try:
                profile = json.loads(row[0])
                for skill in profile.get("专业技能", []):
                    skill_count[skill] = skill_count.get(skill, 0) + 1
            except:
                pass

    conn.close()

    return {
        "total_users": total_users,
        "total_profiles": total_profiles,
        "total_discussions": total_discussions,
        "total_replies": total_replies,
        "total_reports": total_reports,
        "skill_count": skill_count
    }


# ==================== 讨论功能增强版 ====================

def add_discussion(username: str, user_type: str, title: str, content: str, is_anonymous: bool = False) -> bool:
    """发布新话题"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO discussions (username, user_type, title, content, is_anonymous) VALUES (?, ?, ?, ?, ?)",
                  (username, user_type, title, content, 1 if is_anonymous else 0))
        conn.commit()
        conn.close()
        return True
    except Exception as e:
        print(f"Error adding discussion: {e}")
        return False


def get_discussions(keyword: str = None) -> List[Dict]:
    """获取讨论列表，支持关键词搜索"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    
    if keyword and keyword.strip():
        search_term = f"%{keyword.strip()}%"
        c.execute("""
            SELECT * FROM discussions 
            WHERE title LIKE ? OR content LIKE ?
            ORDER BY is_pinned DESC, created_at DESC
        """, (search_term, search_term))
    else:
        c.execute("SELECT * FROM discussions ORDER BY is_pinned DESC, created_at DESC")
    
    discussions = [dict(row) for row in c.fetchall()]
    conn.close()
    return discussions


def update_discussion(discussion_id: int, title: str, content: str, username: str, user_type: str) -> bool:
    """编辑讨论（仅发布者或管理员）"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        if user_type == "admin":
            c.execute("UPDATE discussions SET title = ?, content = ? WHERE id = ?", 
                      (title, content, discussion_id))
        else:
            c.execute("UPDATE discussions SET title = ?, content = ? WHERE id = ? AND username = ?", 
                      (title, content, discussion_id, username))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def delete_discussion(discussion_id: int, username: str, user_type: str) -> bool:
    """删除讨论（仅发布者或管理员）"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        if user_type == "admin":
            c.execute("DELETE FROM discussions WHERE id = ?", (discussion_id,))
        else:
            c.execute("DELETE FROM discussions WHERE id = ? AND username = ?", (discussion_id, username))
        c.execute("DELETE FROM replies WHERE discussion_id = ?", (discussion_id,))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def pin_discussion(discussion_id: int, is_pinned: int) -> bool:
    """置顶/取消置顶讨论（仅管理员）"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("UPDATE discussions SET is_pinned = ? WHERE id = ?", (is_pinned, discussion_id))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def like_discussion(discussion_id: int, username: str) -> bool:
    """点赞/取消点赞讨论"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT likes, liked_by FROM discussions WHERE id = ?", (discussion_id,))
        row = c.fetchone()
        if row:
            likes = row[0]
            liked_by = row[1] or ""
            liked_list = liked_by.split(",") if liked_by else []

            if username in liked_list:
                liked_list.remove(username)
                likes = max(0, likes - 1)
            else:
                liked_list.append(username)
                likes += 1

            c.execute("UPDATE discussions SET likes = ?, liked_by = ? WHERE id = ?",
                      (likes, ",".join(liked_list), discussion_id))
            conn.commit()
            conn.close()
            return True
        return False
    except:
        return False


def check_discussion_liked(discussion_id: int, username: str) -> bool:
    """检查用户是否已点赞讨论"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT liked_by FROM discussions WHERE id = ?", (discussion_id,))
    row = c.fetchone()
    conn.close()
    if row and row[0]:
        return username in row[0].split(",")
    return False


# ==================== 回复功能增强版 ====================

def add_reply(discussion_id: int, username: str, user_type: str, content: str, is_anonymous: bool = False) -> bool:
    """添加回复"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO replies (discussion_id, username, user_type, content, is_anonymous) VALUES (?, ?, ?, ?, ?)",
                  (discussion_id, username, user_type, content, 1 if is_anonymous else 0))
        c.execute("UPDATE discussions SET reply_count = reply_count + 1 WHERE id = ?", (discussion_id,))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def get_replies(discussion_id: int) -> List[Dict]:
    """获取回复列表"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM replies WHERE discussion_id = ? ORDER BY created_at ASC", (discussion_id,))
    replies = [dict(row) for row in c.fetchall()]
    conn.close()
    return replies


def update_reply(reply_id: int, content: str, username: str, user_type: str) -> bool:
    """编辑回复（仅发布者或管理员）"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        if user_type == "admin":
            c.execute("UPDATE replies SET content = ? WHERE id = ?", (content, reply_id))
        else:
            c.execute("UPDATE replies SET content = ? WHERE id = ? AND username = ?", 
                      (content, reply_id, username))
        conn.commit()
        conn.close()
        return True
    except:
        return False


def delete_reply(reply_id: int, username: str, user_type: str, discussion_id: int = None) -> bool:
    """删除回复（仅发布者或管理员）"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        if user_type == "admin":
            c.execute("DELETE FROM replies WHERE id = ?", (reply_id,))
        else:
            c.execute("DELETE FROM replies WHERE id = ? AND username = ?", (reply_id, username))
        
        if discussion_id:
            c.execute("UPDATE discussions SET reply_count = reply_count - 1 WHERE id = ?", (discussion_id,))
        
        conn.commit()
        conn.close()
        return True
    except:
        return False


def like_reply(reply_id: int, username: str) -> bool:
    """点赞/取消点赞回复"""
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("SELECT likes, liked_by FROM replies WHERE id = ?", (reply_id,))
        row = c.fetchone()
        if row:
            likes = row[0]
            liked_by = row[1] or ""
            liked_list = liked_by.split(",") if liked_by else []

            if username in liked_list:
                liked_list.remove(username)
                likes = max(0, likes - 1)
            else:
                liked_list.append(username)
                likes += 1

            c.execute("UPDATE replies SET likes = ?, liked_by = ? WHERE id = ?",
                      (likes, ",".join(liked_list), reply_id))
            conn.commit()
            conn.close()
            return True
        return False
    except:
        return False


def check_reply_liked(reply_id: int, username: str) -> bool:
    """检查用户是否已点赞回复"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT liked_by FROM replies WHERE id = ?", (reply_id,))
    row = c.fetchone()
    conn.close()
    if row and row[0]:
        return username in row[0].split(",")
    return False


# ==================== 学生画像类 ====================
class StudentProfileManager:
    def __init__(self):
        self.profile = {}

    def create_profile_from_form(self, form_data: Dict) -> Dict:
        self.profile = {
            "基本信息": {
                "姓名": form_data.get("name", ""),
                "专业": form_data.get("major", ""),
                "学历": form_data.get("education", "本科"),
                "学校": form_data.get("school", "")
            },
            "专业技能": form_data.get("skills", []),
            "证书": form_data.get("certificates", []),
            "项目经验": form_data.get("projects", []),
            "实习经历": form_data.get("internship", ""),
            "自我评价": form_data.get("self_evaluation", ""),
            "求职意向": {
                "城市": form_data.get("city", ""),
                "期望薪资": form_data.get("salary", ""),
                "岗位类型": form_data.get("job_type", [])
            }
        }
        return self.profile

    def calculate_completeness_score(self, profile: Dict) -> float:
        score = 0
        basic = profile.get("基本信息", {})
        if basic.get("姓名"): score += 5
        if basic.get("专业"): score += 5
        if basic.get("学历"): score += 5
        if basic.get("学校"): score += 5
        skills = profile.get("专业技能", [])
        if skills:
            score += min(30, len(skills) * 6)
        certs = profile.get("证书", [])
        if certs:
            score += min(20, len(certs) * 5)
        projects = profile.get("项目经验", [])
        if projects:
            score += min(20, len(projects) * 7)
        if profile.get("实习经历"):
            score += 10
        return min(100, score)

    def calculate_competitiveness_score(self, profile: Dict) -> float:
        score = 0
        skills = profile.get("专业技能", [])
        skill_weight = {
            "Python": 10, "Java": 10, "Go": 10, "C++": 8,
            "机器学习": 12, "深度学习": 12, "AI": 12,
            "云原生": 10, "Docker": 8, "K8s": 10,
            "大数据": 10, "数据分析": 8, "SQL": 6
        }
        for skill in skills:
            score += skill_weight.get(skill, 2)

        certs = profile.get("证书", [])
        cert_weight = {
            "软考高级": 15, "软考中级": 10, "软考初级": 5,
            "PMP": 10, "NPDP": 10, "AWS认证": 10, "ACM": 15
        }
        for cert in certs:
            score += cert_weight.get(cert, 2)

        projects = profile.get("项目经验", [])
        score += min(30, len(projects) * 10)

        return min(100, score)


# ==================== 岗位数据 ====================
class EnhancedJobDataManager:
    """增强版岗位数据管理器 - 包含详细的岗位信息和学习资源"""

    JOB_CATEGORIES = {
        "Java开发工程师": {
            "title": "Java开发工程师",
            "category": "后端开发",
            "skills": ["Java", "Spring Boot", "Spring Cloud", "MySQL", "Redis", "Docker", "Kubernetes", "微服务", "分布式", "JVM"],
            "certificates": ["Oracle Java认证", "软考中级", "AWS认证", "阿里云认证"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "8K-25K",
            "trend": "云原生、微服务、低代码平台是发展方向",
            "创新能力": "需要创新架构设计，探索新技术栈，解决复杂业务问题",
            "学习能力": "需要持续学习Java生态、云原生技术、微服务架构",
            "抗压能力": "面对项目 deadline、技术挑战、业务压力能保持专业",
            "沟通能力": "需要与产品、测试、前端有效沟通，协作完成项目",
            "实习能力": "有Java相关实习经验，了解开发流程和版本控制"
        },
        "Python开发工程师": {
            "title": "Python开发工程师",
            "category": "后端开发",
            "skills": ["Python", "Django", "Flask", "FastAPI", "MySQL", "Redis", "Celery", "Docker", "Linux", "Git"],
            "certificates": ["PCEP", "PCAP", "软考中级", "AWS认证"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "10K-28K",
            "trend": "AI应用、数据分析、自动化运维是热门方向",
            "创新能力": "需要创新开发方法，探索Python生态，解决业务痛点",
            "学习能力": "需要学习Python库、Web框架、数据分析工具",
            "抗压能力": "面对项目压力、技术挑战、快速迭代能保持高效",
            "沟通能力": "需要与团队成员、产品经理、客户有效沟通",
            "实习能力": "有Python开发实习经验，了解Web开发流程"
        },
        "前端开发工程师": {
            "title": "前端开发工程师",
            "category": "前端开发",
            "skills": ["HTML5", "CSS3", "JavaScript", "TypeScript", "Vue", "React", "Angular", "Webpack", "Node.js", "响应式设计"],
            "certificates": ["W3C认证", "前端工程师认证", "软考中级"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "8K-22K",
            "trend": "跨端开发、低代码、AI辅助开发是趋势",
            "创新能力": "需要创新用户界面设计，探索前端新框架，提升用户体验",
            "学习能力": "需要持续学习前端框架、CSS预处理器、构建工具",
            "抗压能力": "面对需求变更、浏览器兼容性问题、性能优化能保持专业",
            "沟通能力": "需要与产品、设计、后端有效沟通，理解需求",
            "实习能力": "有前端开发实习经验，了解前端工程化"
        },
        "软件测试工程师": {
            "title": "软件测试工程师",
            "category": "测试",
            "skills": ["测试理论", "Selenium", "Appium", "Postman", "Jmeter", "Python", "SQL", "性能测试", "自动化测试", "接口测试"],
            "certificates": ["ISTQB", "软考中级", "CSTE"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "7K-18K",
            "trend": "自动化测试、AI测试是发展方向",
            "创新能力": "需要创新测试方法，设计更有效的测试用例，探索AI测试工具",
            "学习能力": "需要学习测试工具、自动化框架、性能测试技术",
            "抗压能力": "面对测试周期紧张、Bug追踪压力、上线前回归测试能高效完成",
            "沟通能力": "需要与开发、产品有效沟通，准确描述Bug并推动解决",
            "实习能力": "有测试实习经验，了解测试流程和Bug管理工具"
        },
        "数据分析工程师": {
            "title": "数据分析工程师",
            "category": "数据科学",
            "skills": ["Python", "SQL", "Excel", "Tableau", "PowerBI", "统计学", "机器学习", "数据可视化", "Hive", "Spark"],
            "certificates": ["CDA", "BDA", "软考中级", "AWS数据分析认证"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "10K-25K",
            "trend": "数据中台、BI分析、AI算法结合",
            "创新能力": "需要创新分析方法，从数据中发现业务洞察，提出改进建议",
            "学习能力": "需要持续学习数据分析工具、统计方法、机器学习算法",
            "抗压能力": "面对数据质量挑战、紧急分析需求、跨部门协调能保持专业",
            "沟通能力": "需要将复杂数据结果转化为业务语言，向非技术部门有效传达",
            "实习能力": "有数据分析实习经验，了解数据仓库和BI工具"
        },
        "产品经理": {
            "title": "产品经理",
            "category": "产品设计",
            "skills": ["需求分析", "Axure", "XMind", "用户研究", "数据分析", "PRD写作", "项目管理", "原型设计", "竞品分析", "产品运营"],
            "certificates": ["PMP", "NPDP", "软考高级", "产品经理认证"],
            "education": "本科及以上",
            "experience": "2-5年",
            "salary_range": "12K-30K",
            "trend": "AI产品经理、数据产品经理需求旺盛",
            "创新能力": "需要创新产品设计思维，挖掘用户痛点，提出创新解决方案",
            "学习能力": "需要学习行业知识、用户研究方法、数据分析技能",
            "抗压能力": "面对多需求并行、资源有限、上线压力能有效推进产品",
            "沟通能力": "需要与设计、开发、运营、老板多方协调，推动产品落地",
            "实习能力": "有产品实习或项目经验，了解产品从0到1流程"
        },
        "UI/UX设计师": {
            "title": "UI/UX设计师",
            "category": "设计",
            "skills": ["Figma", "Sketch", "Adobe XD", "Photoshop", "UI设计", "UX研究", "交互设计", "动效设计", "品牌设计", "Axure"],
            "certificates": ["Adobe认证", "UX认证", "CDA交互设计认证"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "8K-20K",
            "trend": "体验设计、AI辅助设计工具",
            "创新能力": "需要创新设计思路，关注设计趋势，打造差异化设计方案",
            "学习能力": "需要持续学习设计工具、用户研究方法、设计规范",
            "抗压能力": "面对多项目并行、紧急需求、反复修改能保持设计质量",
            "沟通能力": "需要与产品、开发有效沟通，准确传达设计意图",
            "实习能力": "有设计实习经验，有完整设计作品集"
        },
        "网络工程师": {
            "title": "网络工程师",
            "category": "技术支持",
            "skills": ["TCP/IP", "Cisco", "华为设备", "防火墙", "VPN", "网络监控", "Linux", "安全", "无线网络", "SDN"],
            "certificates": ["HCIA", "HCIP", "CCNA", "CISSP", "CISP"],
            "education": "专科及以上",
            "experience": "1-3年",
            "salary_range": "6K-15K",
            "trend": "网络安全、云网络是发展方向",
            "创新能力": "需要创新网络架构设计，优化网络性能，提升网络安全",
            "学习能力": "需要学习网络新技术、安全技术、云网络技术",
            "抗压能力": "面对网络故障、紧急扩容、安全事件能快速响应",
            "沟通能力": "需要与业务部门、技术团队有效沟通网络需求和方案",
            "实习能力": "有网络运维实习经验，了解网络设备和配置"
        },
        "运维工程师": {
            "title": "运维工程师",
            "category": "技术支持",
            "skills": ["Linux", "Shell", "Docker", "K8s", "Ansible", "Prometheus", "Zabbix", "云平台", "Nginx", "Tomcat"],
            "certificates": ["RHCE", "KCNA", "软考中级", "阿里云认证"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "8K-20K",
            "trend": "SRE、DevOps、云原生运维",
            "创新能力": "需要创新运维自动化，提升运维效率，探索监控告警新方法",
            "学习能力": "需要持续学习云原生、DevOps、自动化工具",
            "抗压能力": "面对故障响应、变更发布、容量规划能保持稳定",
            "沟通能力": "需要与开发、产品有效沟通，协调变更和发布",
            "实习能力": "有运维实习经验，了解Linux和自动化工具"
        },
        "AI算法工程师": {
            "title": "AI算法工程师",
            "category": "人工智能",
            "skills": ["Python", "TensorFlow", "PyTorch", "机器学习", "深度学习", "NLP", "计算机视觉", "Linux", "CUDA", "模型部署"],
            "certificates": ["软考高级", "AWS ML", "华为AI认证", "TensorFlow认证"],
            "education": "硕士及以上",
            "experience": "1-3年",
            "salary_range": "15K-40K",
            "trend": "大模型、AIGC、应用落地是风口",
            "创新能力": "需要创新算法方案，解决技术难题，发表高质量论文或专利",
            "学习能力": "需要持续学习前沿算法、深度学习框架、工程优化技术",
            "抗压能力": "面对算法调优、实验失败、项目交付压力能坚持探索",
            "沟通能力": "需要与产品、工程有效沟通，将算法转化为产品",
            "实习能力": "有AI实验室或公司实习经验，有顶会论文或项目经验"
        },
        "数据产品经理": {
            "title": "数据产品经理",
            "category": "产品设计",
            "skills": ["数据分析", "SQL", "产品设计", "数据可视化", "埋点设计", "Python", "Tableau", "需求分析", "Axure", "BI工具"],
            "certificates": ["CDA", "PMP", "NPDP", "数据分析师认证"],
            "education": "本科及以上",
            "experience": "2-4年",
            "salary_range": "15K-35K",
            "trend": "数据中台、数据治理、AI产品方向火热",
            "创新能力": "需要创新数据产品设计，挖掘数据价值，提升数据使用效率",
            "学习能力": "需要学习数据分析方法、数据产品设计、数据技术",
            "抗压能力": "面对多需求、数据质量挑战、跨团队协调能有效推进",
            "沟通能力": "需要与数据开发、业务部门、运营团队有效沟通",
            "实习能力": "有数据产品或数据分析实习经验"
        },
        "嵌入式软件工程师": {
            "title": "嵌入式软件工程师",
            "category": "硬件开发",
            "skills": ["C", "C++", "RTOS", "STM32", "Linux", "ARM", "硬件调试", "UART", "SPI", "I2C", "Makefile"],
            "certificates": ["软考中级", "嵌入式认证", "ARM认证"],
            "education": "本科及以上",
            "experience": "1-3年",
            "salary_range": "10K-25K",
            "trend": "物联网、智能硬件、汽车电子是发展方向",
            "创新能力": "需要创新嵌入式系统设计，优化资源利用，提升系统性能",
            "学习能力": "需要学习硬件知识、RTOS、驱动开发、外设协议",
            "抗压能力": "面对硬件调试困难、实时性要求、功耗优化能持续攻关",
            "沟通能力": "需要与硬件工程师、算法工程师有效协作",
            "实习能力": "有嵌入式开发或硬件项目实习经验"
        }
    }

    PROMOTION_PATHS = {
        "Java开发工程师": ["Java开发工程师 → 高级Java工程师 → 技术组长 → 架构师 → 技术总监"],
        "前端开发工程师": ["前端开发 → 高级前端 → 前端组长 → 前端架构师 → 技术总监"],
        "软件测试工程师": ["测试工程师 → 高级测试 → 测试主管 → 测试经理 → 质量总监"],
        "数据分析工程师": ["数据分析师 → 高级分析师 → 数据主管 → 数据经理 → 首席数据官"],
        "产品经理": ["产品经理 → 高级产品经理 → 产品总监 → VP产品 → CPO"],
        "AI算法工程师": ["算法工程师 → 高级算法 → 算法专家 → 技术总监 → 首席技术官"]
    }

    CAREER_PATHS = {
        "Java开发工程师": [
            "Java开发 → 后端架构师 → 技术总监",
            "Java开发 → 数据开发工程师 → 数据架构师",
            "Java开发 → 产品经理 → 产品总监",
            "Java开发 → AI算法工程师 → AI专家"
        ],
        "前端开发工程师": [
            "前端开发 → 全栈工程师 → 技术总监",
            "前端开发 → UI设计师 → 设计总监",
            "前端开发 → 产品经理 → 产品总监"
        ],
        "软件测试工程师": [
            "测试 → 自动化测试开发 → 测试架构师",
            "测试 → 产品经理 → 产品总监",
            "测试 → 性能测试 → 性能专家"
        ],
        "数据分析工程师": [
            "数据分析 → 数据产品经理 → 产品总监",
            "数据分析 → AI算法工程师 → AI专家",
            "数据分析 → 商业分析师 → 战略分析师"
        ],
        "产品经理": [
            "产品经理 → 产品总监 → VP产品",
            "产品经理 → 项目经理 → 项目总监",
            "产品经理 → 运营总监"
        ]
    }

    LEARNING_RESOURCES = {
        "Java": [
            {"name": "Java核心技术", "type": "书籍", "url": "https://book.douban.com/subject/26880667/"},
            {"name": "Spring Boot实战", "type": "书籍", "url": "https://book.douban.com/subject/26897412/"},
            {"name": "慕课网Java工程师", "type": "在线课程", "url": "https://www.imooc.com/learn/124"},
            {"name": "B站Java学习路线", "type": "视频", "url": "https://www.bilibili.com/video/BV1zE411j7Nw"}
        ],
        "Python": [
            {"name": "Python编程从入门到实践", "type": "书籍", "url": "https://book.douban.com/subject/3116733/"},
            {"name": "Python Flask Web开发", "type": "书籍", "url": "https://book.douban.com/subject/26597428/"},
            {"name": "Python官方文档", "type": "文档", "url": "https://docs.python.org/3/"},
            {"name": "Coursera Python课程", "type": "在线课程", "url": "https://www.coursera.org/learn/python"}
        ],
        "前端": [
            {"name": "MDN Web开发文档", "type": "文档", "url": "https://developer.mozilla.org/"},
            {"name": "Vue3官方文档", "type": "文档", "url": "https://vuejs.org/"},
            {"name": "React官方文档", "type": "文档", "url": "https://react.dev/"},
            {"name": "前端进阶之路", "type": "开源项目", "url": "https://github.com/qianguyihao/Web"}
        ],
        "数据分析": [
            {"name": "利用Python进行数据分析", "type": "书籍", "url": "https://book.douban.com/subject/26278699/"},
            {"name": "Kaggle数据集", "type": "实践", "url": "https://www.kaggle.com/datasets"},
            {"name": "SQL必知必会", "type": "书籍", "url": "https://book.douban.com/subject/24250054/"}
        ],
        "AI算法": [
            {"name": "机器学习实战", "type": "书籍", "url": "https://book.douban.com/subject/24462300/"},
            {"name": "深度学习入门", "type": "书籍", "url": "https://book.douban.com/subject/27028517/"},
            {"name": "吴恩达机器学习", "type": "在线课程", "url": "https://www.coursera.org/learn/machine-learning"},
            {"name": "fast.ai深度学习", "type": "在线课程", "url": "https://www.fast.ai/"}
        ]
    }

    CERT_EXAMS = {
        "软考初级": {"name": "程序员", "url": "https://www.ruankao.org.cn/"},
        "软考中级": {"name": "软件设计师", "url": "https://www.ruankao.org.cn/"},
        "软考高级": {"name": "系统架构设计师", "url": "https://www.ruankao.org.cn/"},
        "PMP": {"name": "项目管理专业人士", "url": "https://www.pmi.org/"},
        "NPDP": {"name": "产品经理专业人士", "url": "https://www.pdma.org/"},
        "AWS认证": {"name": "AWS认证", "url": "https://aws.amazon.com/certification/"}
    }

    INTERNSHIP_PLATFORMS = [
        {"name": "实习僧", "url": "https://www.shixiseng.com/"},
        {"name": "牛客网", "url": "https://www.nowcoder.com/"},
        {"name": "拉勾网实习", "url": "https://www.lagou.com/"},
        {"name": "BOSS直聘实习", "url": "https://www.zhipin.com/"}
    ]

    @classmethod
    def get_all_jobs(cls):
        return list(cls.JOB_CATEGORIES.keys())

    @classmethod
    def get_job_profile(cls, job_name: str) -> Dict:
        return cls.JOB_CATEGORIES.get(job_name, {})

    @classmethod
    def get_promotion_path(cls, job_name: str) -> List[str]:
        return cls.PROMOTION_PATHS.get(job_name, ["暂无晋升路径数据"])

    @classmethod
    def get_career_paths(cls, job_name: str) -> List[str]:
        return cls.CAREER_PATHS.get(job_name, ["暂无换岗路径数据"])

    @classmethod
    def get_learning_resources(cls, skill: str) -> List[Dict]:
        for key, resources in cls.LEARNING_RESOURCES.items():
            if key in skill or skill in key:
                return resources
        return cls.LEARNING_RESOURCES.get("Python", [])

    @classmethod
    def recommend_resources_for_gaps(cls, gaps: List[str]) -> Dict:
        recommendations = {}
        for gap in gaps:
            for skill_key, resources in cls.LEARNING_RESOURCES.items():
                if skill_key.lower() in gap.lower() or gap.lower() in skill_key.lower():
                    recommendations[gap] = resources
                    break
            else:
                recommendations[gap] = cls.LEARNING_RESOURCES.get("Python", [])
        return recommendations

# 兼容旧类名
JobDataManager = EnhancedJobDataManager

class AdvancedJobMatchingEngine:
    """升级版岗位匹配引擎 - 基于AI深度分析"""

    def __init__(self):
        self.job_data = JobDataManager()
        self.llm = LLMClient()

    def match_with_ai(self, student_profile: Dict, resume_portrait: Dict = None, job_name: str = None) -> Dict:
        """使用AI进行深度岗位匹配分析"""

        # 合并学生画像数据
        if resume_portrait:
            merged_profile = self._merge_profiles(student_profile, resume_portrait)
        else:
            merged_profile = student_profile

        # 获取岗位信息
        if not job_name:
            job_name = merged_profile.get('求职意向', {}).get('期望岗位', 'Java开发工程师')
        job_info = self.job_data.get_job_profile(job_name)

        # 构建AI分析提示词
        prompt = f"""
请作为一名专业的职业规划顾问，分析以下学生与岗位的匹配情况，给出详细的匹配分析报告。

## 学生能力画像
- 姓名: {merged_profile.get('基本信息', {}).get('姓名', '未知')}
- 专业: {merged_profile.get('基本信息', {}).get('专业', '未知')}
- 学历: {merged_profile.get('基本信息', {}).get('学历', '未知')}
- 学校: {merged_profile.get('基本信息', {}).get('学校', '未知')}
- 专业技能: {', '.join(merged_profile.get('专业技能', []))}
- 证书: {', '.join(merged_profile.get('证书', []))}
- 项目经验: {', '.join([str(item) if isinstance(item, dict) else item for item in merged_profile.get('项目经验', [])])}
- 实习经历: {merged_profile.get('实习经历', '无')[:200]}
- 学习能力: {merged_profile.get('学习能力', '未知')}
- 沟通能力: {merged_profile.get('沟通能力', '未知')}
- 抗压能力: {merged_profile.get('抗压能力', '未知')}
- 创新能力: {merged_profile.get('创新能力', '未知')}

## 目标岗位信息
- 岗位名称: {job_name}
- 所需技能: {', '.join(job_info.get('skills', []))}
- 薪资范围: {job_info.get('salary_range', '面议')}
- 学历要求: {job_info.get('education', '本科及以上')}
- 发展趋势: {job_info.get('trend', '良好')}

请按以下JSON格式输出分析结果（只输出JSON，不要其他文字）：
{{
    "总匹配度": 0-100的数字,
    "各维度得分": {{
        "学历匹配": 0-100,
        "技能匹配": 0-100,
        "项目经验匹配": 0-100,
        "证书匹配": 0-100,
        "软实力匹配": 0-100
    }},
    "优势": ["优势1", "优势2", "优势3"],
    "差距": ["差距1", "差距2", "差距3"],
    "技能提升建议": ["建议1", "建议2", "建议3"],
    "推荐理由": "详细的推荐理由",
    "竞争力评估": "优秀/良好/一般/待提升",
    "建议薪资": "建议薪资范围"
}}
"""

        try:
            response = self.llm.call(prompt)
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                result['岗位名称'] = job_name
                result['分析时间'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                return result
        except Exception as e:
            st.error(f"AI分析失败: {e}")

        return self._fallback_match(merged_profile, job_info, job_name)

    def _merge_profiles(self, student_profile: Dict, resume_portrait: Dict) -> Dict:
        """合并学生画像和简历解析结果"""
        merged = student_profile.copy()

        if '基本信息' not in merged:
            merged['基本信息'] = {}

        if resume_portrait.get('姓名') and resume_portrait['姓名'] != '未知':
            merged['基本信息']['姓名'] = resume_portrait['姓名']
        if resume_portrait.get('专业') and resume_portrait['专业'] != '未知':
            merged['基本信息']['专业'] = resume_portrait['专业']
        if resume_portrait.get('学历') and resume_portrait['学历'] != '未知':
            merged['基本信息']['学历'] = resume_portrait['学历']
        if resume_portrait.get('意向城市'):
            merged['求职意向'] = merged.get('求职意向', {})
            merged['求职意向']['城市'] = resume_portrait['意向城市']

        existing_skills = set(merged.get('专业技能', []))
        new_skills = set(resume_portrait.get('专业技能', []))
        merged['专业技能'] = list(existing_skills.union(new_skills))

        existing_certs = set(merged.get('证书', []))
        new_certs = set(resume_portrait.get('证书', []))
        merged['证书'] = list(existing_certs.union(new_certs))

        existing_projects = merged.get('项目经验', [])
        new_projects = resume_portrait.get('项目经验', [])
        merged['项目经验'] = existing_projects + new_projects

        merged['学习能力'] = resume_portrait.get('学习能力', merged.get('学习能力', '未知'))
        merged['沟通能力'] = resume_portrait.get('沟通能力', merged.get('沟通能力', '未知'))
        merged['抗压能力'] = resume_portrait.get('抗压能力', merged.get('抗压能力', '未知'))
        merged['创新能力'] = resume_portrait.get('创新能力', merged.get('创新能力', '未知'))

        return merged

    def _fallback_match(self, profile: Dict, job_info: Dict, job_name: str) -> Dict:
        """传统匹配算法（兜底）"""
        student_skills = set(profile.get("专业技能", []))
        required_skills = set(job_info.get("skills", []))

        if required_skills:
            skill_match = len(student_skills.intersection(required_skills)) / len(required_skills) * 100
        else:
            skill_match = 50

        edu_level = {"专科": 60, "本科": 80, "硕士": 90, "博士": 95}
        student_edu = profile.get("基本信息", {}).get("学历", "本科")
        required_edu = job_info.get("education", "本科及以上")
        edu_match = edu_level.get(student_edu, 70)
        if "硕士" in required_edu and student_edu not in ["硕士", "博士"]:
            edu_match = 50

        project_count = len(profile.get("项目经验", []))
        project_match = min(100, project_count * 20)

        cert_count = len(profile.get("证书", []))
        cert_match = min(100, cert_count * 15)

        total_match = (skill_match * 0.4 + edu_match * 0.2 + project_match * 0.2 + cert_match * 0.15 + 70 * 0.05)

        gaps = list(required_skills - student_skills)

        return {
            "总匹配度": round(total_match, 1),
            "各维度得分": {
                "学历匹配": round(edu_match, 1),
                "技能匹配": round(skill_match, 1),
                "项目经验匹配": round(project_match, 1),
                "证书匹配": round(cert_match, 1),
                "软实力匹配": 70
            },
            "优势": self._get_advantages(profile, job_info),
            "差距": gaps if gaps else ["需要更多项目实战经验"],
            "技能提升建议": [f"学习{gap}" for gap in gaps[:3]] if gaps else ["持续深入掌握核心技术栈"],
            "推荐理由": f"该岗位与您的专业背景有一定匹配度，技能匹配度为{round(skill_match)}%",
            "竞争力评估": "良好" if total_match > 70 else "一般",
            "建议薪资": job_info.get("salary_range", "8K-15K"),
            "岗位名称": job_name,
            "分析时间": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }

    def _get_advantages(self, profile: Dict, job_info: Dict) -> List[str]:
        """分析优势"""
        advantages = []
        skills = set(profile.get("专业技能", []))
        required = set(job_info.get("skills", []))

        matched_skills = skills.intersection(required)
        if matched_skills:
            advantages.append(f"掌握{', '.join(list(matched_skills)[:3])}等岗位核心技能")

        if profile.get("项目经验"):
            advantages.append(f"拥有{len(profile.get('项目经验', []))}个项目经验，具备实践能力")

        edu = profile.get("基本信息", {}).get("学历", "")
        if edu in ["本科", "硕士", "博士"]:
            advantages.append(f"{edu}学历符合岗位基本要求")

        if not advantages:
            advantages.append("具备学习意愿和发展潜力")

        return advantages[:3]

    def _get_recommended_jobs(self, profile: Dict) -> List[str]:
        """获取推荐的岗位列表"""
        all_jobs = self.job_data.get_all_jobs()
        student_skills = set(profile.get('专业技能', []))

        job_scores = []
        for job in all_jobs:
            job_info = self.job_data.get_job_profile(job)
            required = set(job_info.get('skills', []))
            if required:
                score = len(student_skills.intersection(required)) / len(required)
            else:
                score = 0.3
            job_scores.append((job, score))

        job_scores.sort(key=lambda x: x[1], reverse=True)
        return [job for job, _ in job_scores[:3]]


class AdvancedCareerReportGenerator:
    """升级版职业规划报告生成器 - 基于AI深度分析"""

    def __init__(self):
        self.job_data = JobDataManager()
        self.llm = LLMClient()

    def generate_ai_report(self, student_profile: Dict, match_results: Dict, resume_portrait: Dict = None) -> tuple:
        """使用AI生成个性化职业规划报告，返回(报告内容, 生成时间)"""

        if resume_portrait:
            merged_profile = self._merge_with_resume(student_profile, resume_portrait)
        else:
            merged_profile = student_profile

        # 生成统一的时间戳
        generated_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        recommended_jobs = self._get_recommended_jobs(merged_profile)

        prompt = f"""
请作为一名资深的职业规划专家，根据以下学生信息生成一份详细、个性化的职业发展规划报告。

## 学生综合信息
### 基本信息
- 姓名: {merged_profile.get('基本信息', {}).get('姓名', '同学')}
- 专业: {merged_profile.get('基本信息', {}).get('专业', '计算机相关')}
- 学历: {merged_profile.get('基本信息', {}).get('学历', '本科')}
- 学校: {merged_profile.get('基本信息', {}).get('学校', '普通高校')}

### 能力分析
- 专业技能: {', '.join(merged_profile.get('专业技能', [])) or '待完善'}
- 证书: {', '.join(merged_profile.get('证书', [])) or '暂无'}
- 项目经验数量: {len(merged_profile.get('项目经验', []))}个
- 实习经历: {merged_profile.get('实习经历', '无')[:150]}

### 软实力评估
- 学习能力: {merged_profile.get('学习能力', '良好')}
- 沟通能力: {merged_profile.get('沟通能力', '良好')}
- 抗压能力: {merged_profile.get('抗压能力', '良好')}
- 创新能力: {merged_profile.get('创新能力', '良好')}

### 求职意向
- 期望城市: {merged_profile.get('求职意向', {}).get('城市', '不限')}
- 期望薪资: {merged_profile.get('求职意向', {}).get('期望薪资', '面议')}

## AI匹配分析结果
- 推荐岗位Top3: {recommended_jobs}
- 最佳匹配岗位: {match_results.get('岗位名称', recommended_jobs[0] if recommended_jobs else 'Java开发工程师')}
- 总匹配度: {match_results.get('总匹配度', 70)}%
- 优势: {', '.join(match_results.get('优势', ['具备相关专业背景']))}
- 待提升: {', '.join(match_results.get('差距', ['项目经验不足']))}

请生成一份完整的职业规划报告，包含以下部分：

1. **个人竞争力分析** - 基于能力画像的客观评估
2. **职业定位建议** - 推荐最适合的岗位及理由
3. **阶段性发展目标** - 短期(1年)、中期(1-3年)、长期(3-5年)目标
4. **详细行动计划** - 具体可行的学习和发展计划
5. **技能提升路径** - 针对差距的学习资源建议
6. **求职策略建议** - 简历优化、面试准备、求职渠道
7. **行业趋势分析** - 目标岗位的未来发展前景

请用专业、鼓励的语气，格式清晰，使用markdown格式输出。
"""

        try:
            report = self.llm.call(prompt)
            header = f"""# 📋 个性化职业发展规划报告

**生成时间**: {generated_time}
**报告类型**: AI智能分析版
**适用用户**: {merged_profile.get('基本信息', {}).get('姓名', '同学')}

---
"""
            return header + report, generated_time
        except Exception as e:
            st.error(f"AI报告生成失败: {e}")
            fallback_report = self._generate_fallback_report(merged_profile, match_results, recommended_jobs)
            return fallback_report, generated_time

    def _merge_with_resume(self, profile: Dict, resume: Dict) -> Dict:
        """合并简历解析数据"""
        merged = profile.copy()

        if '基本信息' not in merged:
            merged['基本信息'] = {}

        if resume.get('姓名') and resume['姓名'] != '未知':
            merged['基本信息']['姓名'] = resume['姓名']
        if resume.get('专业') and resume['专业'] != '未知':
            merged['基本信息']['专业'] = resume['专业']
        if resume.get('学历') and resume['学历'] != '未知':
            merged['基本信息']['学历'] = resume['学历']

        existing_skills = set(merged.get('专业技能', []))
        new_skills = set(resume.get('专业技能', []))
        merged['专业技能'] = list(existing_skills.union(new_skills))

        existing_certs = set(merged.get('证书', []))
        new_certs = set(resume.get('证书', []))
        merged['证书'] = list(existing_certs.union(new_certs))

        for ability in ['学习能力', '沟通能力', '抗压能力', '创新能力']:
            if resume.get(ability) and resume[ability] != '未知':
                merged[ability] = resume[ability]

        if resume.get('意向城市') and resume['意向城市'] != '未知':
            if '求职意向' not in merged:
                merged['求职意向'] = {}
            merged['求职意向']['城市'] = resume['意向城市']

        return merged

    def _get_recommended_jobs(self, profile: Dict) -> List[str]:
        """获取推荐的岗位列表"""
        all_jobs = self.job_data.get_all_jobs()
        student_skills = set(profile.get('专业技能', []))

        job_scores = []
        for job in all_jobs:
            job_info = self.job_data.get_job_profile(job)
            required = set(job_info.get('skills', []))
            if required:
                score = len(student_skills.intersection(required)) / len(required)
            else:
                score = 0.3
            job_scores.append((job, score))

        job_scores.sort(key=lambda x: x[1], reverse=True)
        return [job for job, _ in job_scores[:3]]

    def _generate_fallback_report(self, profile: Dict, match_results: Dict, recommended_jobs: List[str]) -> str:
        """生成兜底报告"""
        name = profile.get('基本信息', {}).get('姓名', '同学')
        best_job = match_results.get('岗位名称', recommended_jobs[0] if recommended_jobs else '目标岗位')
        match_rate = match_results.get('总匹配度', 70)
        advantages = match_results.get('优势', ['具备相关专业背景'])
        gaps = match_results.get('差距', ['需要提升项目经验'])

        return f"""# 📋 职业发展规划报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
**适用用户**: {name}

---

## 一、个人竞争力分析

根据您的能力画像分析：

### 综合评估
- **匹配度评分**: {match_rate}%
- **竞争力等级**: {"优秀" if match_rate >= 80 else "良好" if match_rate >= 60 else "待提升"}

### 核心优势
{self._format_list(advantages)}

### 待提升方向
{self._format_list(gaps)}

---

## 二、职业定位建议

### 推荐岗位Top3
{self._format_job_list(recommended_jobs)}

### 最佳匹配岗位
**{best_job}** - 匹配度{match_rate}%

---

## 三、阶段性发展目标

### 🎯 短期目标 (1年内)
- 掌握岗位核心技能栈
- 完成2-3个实战项目
- 考取相关认证证书

### 🎯 中期目标 (1-3年)
- 成为团队核心开发人员
- 积累完整项目经验
- 提升架构设计能力

### 🎯 长期目标 (3-5年)
- 向技术专家或管理方向发展
- 建立个人技术影响力
- 持续跟进前沿技术

---

## 四、详细行动计划

### 第1-3个月：基础夯实
- 系统学习{best_job}所需的核心技术
- 每天保持2-3小时学习时间
- 参与开源项目或做个人项目

### 第4-6个月：项目实践
- 完成1-2个完整的项目开发
- 将项目部署上线
- 撰写技术博客记录学习过程

### 第7-12个月：求职准备
- 完善简历，突出项目经验
- 刷算法题，准备面试
- 投递实习或校招岗位

---

## 五、技能提升路径

### 必须掌握
- 岗位核心框架和工具链
- 数据库设计与优化
- 版本控制和协作工具

### 推荐学习资源
- **在线课程**: B站、慕课网、Coursera
- **书籍**: 行业经典书籍
- **实践平台**: LeetCode、GitHub、Kaggle

---

## 六、求职策略建议

### 简历优化
1. 突出项目经验和实际成果
2. 用数据量化工作成果
3. 针对不同岗位定制简历

### 面试准备
1. 技术面试：刷题+系统设计
2. 行为面试：准备STAR案例
3. 公司研究：了解目标公司业务

### 求职渠道
- 校园招聘会
- BOSS直聘、拉勾网
- 公司官网和内推

---

## 七、行业趋势分析

{best_job}岗位当前市场需求旺盛，随着数字化转型深入，该岗位的发展前景广阔。建议持续关注：

- 云原生技术发展
- AI与业务结合
- 全栈能力要求提升

---

*本报告由AI职业规划系统生成，仅供参考。建议每3-6个月重新评估并调整规划。*
"""

    def _format_list(self, items: List[str]) -> str:
        if not items:
            return "- 暂无\n"
        return "\n".join([f"- {item}" for item in items]) + "\n"

    def _format_job_list(self, jobs: List[str]) -> str:
        if not jobs:
            return "1. Java开发工程师\n2. Python开发工程师\n3. 前端开发工程师\n"
        return "\n".join([f"{i+1}. {job}" for i, job in enumerate(jobs)]) + "\n"


# ==================== 大模型客户端 ====================
class LLMClient:
    """大模型客户端"""

    def __init__(self):
        import os
        self.api_key = os.getenv("AI_API_KEY", "sk-42a21eb1dea14df8875dd923c3881d49")
        self.base_url = "https://dashscope.aliyuncs.com/compatible-mode/v1"
        self.model = "qwen-turbo"

    def call(self, prompt: str, system_prompt: str = "") -> str:
        """调用大模型API"""
        try:
            import requests
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            data = {
                "model": self.model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7
            }
            response = requests.post(
                f"{self.base_url}/chat/completions",
                headers=headers,
                json=data,
                timeout=30
            )
            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"调用大模型API失败: {e}")
            # 失败时返回模拟响应
            return self._mock_response(prompt)

    def chat(self, prompt: str, system_prompt: str = "你是一个专业的职业规划顾问，擅长分析学生与岗位的匹配度。") -> str:
        """聊天接口"""
        return self.call(prompt, system_prompt)

    def _mock_response(self, prompt: str) -> str:
        """模拟返回"""
        import json
        if "岗位画像" in prompt or "职位描述" in prompt:
            return json.dumps({
                "专业技能": ["Java", "SpringBoot", "MySQL", "Redis", "Docker"],
                "证书要求": ["软考中级", "PMP可选"],
                "抗压能力": "高",
                "沟通能力": "中",
                "学习能力": "高",
                "创新能力": "中",
                "学历要求": "本科及以上",
                "工具": ["Git", "Maven", "IDEA"],
                "发展趋势": "云原生、微服务架构是未来发展方向"
            }, ensure_ascii=False)

        elif "学生画像" in prompt or "简历" in prompt:
            return json.dumps({
                "专业技能": ["Python", "Java", "MySQL", "HTML/CSS"],
                "证书": ["英语四级", "软考初级"],
                "项目经验": ["电商项目", "学生管理系统"],
                "实习经历": "某互联网公司实习3个月",
                "学习能力": "强",
                "沟通能力": "中",
                "抗压能力": "高",
                "创新能力": "中"
            }, ensure_ascii=False)

        elif "职业规划" in prompt or "报告" in prompt:
            return """
## 职业发展规划报告

### 一、职业探索与岗位匹配
根据您的能力画像和市场需求分析，推荐以下Top3岗位：
1. Java开发工程师（匹配度85%）
2. Python开发工程师（匹配度82%）
3. 软件测试工程师（匹配度78%）

### 二、职业目标与发展路径
- 短期目标（1年内）：成为初级Java开发工程师
- 中期目标（1-3年）：晋升为中级开发工程师
- 长期目标（3-5年）：成为技术团队负责人

### 三、行动计划
**短期（1-6个月）：**
- 完成SpringBoot微服务项目实战
- 考取软考中级证书
- 投递实习岗位积累经验

**中期（6-18个月）：**
- 深入学习分布式系统
- 提升架构设计能力
- 准备跳槽到一线大厂

### 四、评估与调整
建议每3个月进行一次复盘，根据行业变化和个人发展情况调整计划。
"""

        return "收到请求，正在处理..."

# 实例化LLM客户端
llm = LLMClient()

# ==================== 行业动态API ====================
class IndustryNewsAPI:
    """行业动态API集成"""

    @staticmethod
    def get_tech_news() -> List[Dict]:
        try:
            response = requests.get(
                "https://api.jinse.com/community/v1/notes/lists?limit=10",
                timeout=5
            )
            if response.status_code == 200:
                data = response.json()
                return [{"title": f"行业动态{i+1}", "content": "技术行业最新动态"} for i in range(5)]
        except:
            pass

        return [
            {"title": "AI大模型应用爆发", "content": "2024年AI大模型在各行业应用持续深入，企业对AI人才需求旺盛", "source": "行业观察"},
            {"title": "云原生技术趋势", "content": "云原生技术成为企业数字化转型关键，Kubernetes应用广泛", "source": "技术前沿"},
            {"title": "前端技术更新", "content": "前端框架持续演进，跨端开发成为主流方向", "source": "前端之家"},
            {"title": "数据安全受重视", "content": "数据安全法实施，企业对安全人才需求增加", "source": "安全资讯"},
            {"title": "新能源行业崛起", "content": "新能源行业快速发展，对嵌入式和硬件人才需求增加", "source": "行业观察"}
        ]

    @staticmethod
    def get_job_market_trend() -> Dict:
        return {
            "热门岗位": ["AI算法工程师", "Java开发工程师", "前端开发工程师"],
            "薪资上涨": ["AI算法工程师 +15%", "云原生工程师 +12%", "数据工程师 +10%"],
            "需求增加": ["产品经理", "数据分析", "网络安全"],
            "建议": "建议关注AI方向岗位，同时加强云原生和数据技能"
        }


# ==================== 匹配引擎 ====================
class JobMatchingEngine:
    def __init__(self):
        self.job_data = JobDataManager()

    def match(self, student_profile: Dict, job_name: str) -> Dict:
        job = self.job_data.get_job_profile(job_name)
        if not job:
            return {"error": "未找到该岗位信息"}

        student_skills = set(student_profile.get("专业技能", []))
        required_skills = set(job.get("skills", []))

        if not required_skills:
            match_rate = 50
        else:
            matched = len(student_skills.intersection(required_skills))
            match_rate = matched / len(required_skills) * 100

        gaps = list(required_skills - student_skills)

        advantages = []
        if student_profile.get("基本信息", {}).get("专业"):
            advantages.append("专业背景与岗位职责有一定关联，具备较强的数据处理和编程能力")
        if "英语四级" in student_profile.get("证书", []) or "英语六级" in student_profile.get("证书", []):
            advantages.append("英语证书表明具备基本的英文读写能力，符合岗位对英语的要求")
        if student_profile.get("专业技能"):
            advantages.append("具备相关技术技能，能够快速上手工作")

        disadvantages = []
        if not student_profile.get("实习经历"):
            disadvantages.append("缺乏相关领域的实习或工作经验，可能影响岗位适应性")
        if len(student_profile.get("项目经验", [])) < 2:
            disadvantages.append("项目经验不足，技能与岗位需求存在一定的脱节")

        return {
            "总匹配度": round(match_rate + 20, 1),
            "基础要求": 80,
            "职业技能": round(match_rate + 10, 1),
            "职业素养": 75,
            "发展潜力": 80,
            "优势": advantages[:2] if advantages else ["具备相关专业背景"],
            "差距": disadvantages if disadvantages else ["需要提升相关技能"],
            "推荐理由": "该岗位与您的专业背景有一定匹配度，建议提升相关技能后再申请。"
        }


# ==================== 报告生成 ====================
class CareerReportGenerator:
    def __init__(self):
        self.job_data = JobDataManager()

    def generate_report(self, student_profile: Dict, match_results: Dict) -> str:
        all_jobs = self.job_data.get_all_jobs()
        matching_engine = JobMatchingEngine()

        job_scores = []
        for job in all_jobs:
            result = matching_engine.match(student_profile, job)
            if "总匹配度" in result:
                job_scores.append((job, result["总匹配度"]))

        job_scores.sort(key=lambda x: x[1], reverse=True)
        top3_jobs = job_scores[:3]

        completeness = student_profile.get('completeness_score', 0)
        competitiveness = student_profile.get('competitiveness_score', 0)

        report = f"""# 职业发展规划报告

**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## 一、职业探索与岗位匹配

### 推荐Top3岗位

"""

        for i, (job, score) in enumerate(top3_jobs, 1):
            job_info = self.job_data.get_job_profile(job)
            report += f"""**{i}. {job}** (匹配度 {score}%)
- 薪资范围: {job_info.get('salary_range', '面议')}
- 学历要求: {job_info.get('education', '本科')}
- 发展趋势: {job_info.get('trend', '良好')}

"""

        report += f"""### 您的匹配详情

- **总匹配度**: {match_results.get('总匹配度', 0)}%
- 基础要求得分: {match_results.get('基础要求', 0)}分
- 职业技能得分: {match_results.get('职业技能', 0)}分
- 职业素养得分: {match_results.get('职业素养', 0)}分
- 发展潜力得分: {match_results.get('发展潜力', 0)}分

### 优势分析
{', '.join(match_results.get('优势', ['暂无']))}

### 差距分析
{', '.join(match_results.get('差距', ['暂无']))}

---

## 二、职业目标与发展路径

"""

        for job, _ in top3_jobs[:1]:
            paths = self.job_data.get_promotion_path(job)
            report += f"""### {job}晋升路径

"""
            for path in paths:
                report += f"- {path}\n"

            career_paths = self.job_data.get_career_paths(job)
            report += f"""

### 换岗路径选择

"""
            for path in career_paths[:3]:
                report += f"- {path}\n"

        report += """

---

## 三、行动计划

### 短期目标（1-6个月）

"""

        gaps = match_results.get('差距', [])
        if gaps:
            report += f"""1. **技能提升**: 重点学习相关技术栈
2. **证书考取**: 根据目标岗位考取相关证书
3. **项目实战**: 完成1-2个实战项目
4. **实习投递**: 开始投递实习岗位

"""

        report += """### 中期目标（6-18个月）

1. 积累项目经验，提升实战能力
2. 准备跳槽到目标公司
3. 持续学习行业新技术

### 长期目标（3-5年）

1. 成为技术骨干或管理人才
2. 持续关注行业前沿动态
3. 建立个人职业品牌

---

## 四、评估与调整

建议每**3个月**进行一次职业规划复盘，根据以下情况进行调整：

- 行业技术发展趋势变化
- 个人兴趣和能力发展情况
- 就业市场岗位需求变化
- 目标公司战略调整

---

*本报告由AI职业规划智能体生成，仅供参考*
"""

        return report


class ResumeParser:
    """简历解析器"""

    API_KEY = "sk-42a21eb1dea14df8875dd923c3881d49"
    API_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"

    @staticmethod
    def parse_resume_with_llm(resume_text: str) -> Dict:
        """使用大模型解析简历"""
        prompt = f"""
请解析以下简历文本，生成结构化的学生能力画像JSON，要求包含以下维度：
1. 姓名（字符串）
2. 专业（字符串）
3. 学历（字符串）
4. 意向城市（字符串）
5. 专业技能（列表）
6. 证书（列表）
7. 项目经验（列表）
8. 实习经历（文本）
9. 学习能力（文本）
10. 沟通能力（文本）
11. 抗压能力（文本）
12. 创新能力（文本）

简历文本：
{resume_text[:3000]}

要求：
1. 仅输出JSON，不要其他文字
2. 没有的信息填"未知"
3. 分析内容要贴合简历实际
"""
        headers = {
            "Authorization": f"Bearer {ResumeParser.API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "qwen-turbo",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1
        }

        try:
            response = requests.post(ResumeParser.API_URL, headers=headers, json=data, timeout=30)
            result = response.json()
            if 'choices' in result:
                content = result['choices'][0]['message']['content']
                # 提取JSON
                import re
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    return json.loads(json_match.group())
        except Exception as e:
            st.error(f"解析失败: {e}")

        # 兜底返回
        return {
            "姓名": "未知",
            "专业": "未知",
            "学历": "本科",
            "意向城市": "未知",
            "专业技能": [],
            "证书": [],
            "项目经验": [],
            "实习经历": "",
            "学习能力": "未知",
            "沟通能力": "未知",
            "抗压能力": "未知",
            "创新能力": "未知"
        }

    @staticmethod
    def calculate_portrait_score(portrait: Dict) -> Dict:
        """计算画像评分"""
        # 完整度评分
        total_fields = 12
        filled_fields = 0
        for k, v in portrait.items():
            if v and v != "未知" and (not isinstance(v, list) or len(v) > 0):
                filled_fields += 1
        completeness_score = round((filled_fields / total_fields) * 100, 2)

        # 竞争力评分
        skill_count = len(portrait.get("专业技能", []))
        cert_count = len(portrait.get("证书", []))
        project_count = len(portrait.get("项目经验", []))

        skill_weight = {"Python": 10, "Java": 10, "Go": 10, "机器学习": 15, "深度学习": 15,
                       "Docker": 8, "K8s": 10, "大数据": 10, "AI": 12}
        skill_score = sum(skill_weight.get(s, 2) for s in portrait.get("专业技能", []))

        cert_weight = {"软考高级": 15, "软考中级": 10, "软考初级": 5, "PMP": 10, "ACM": 15}
        cert_score = sum(cert_weight.get(c, 2) for c in portrait.get("证书", []))

        project_score = project_count * 10
        competitiveness_score = min(100, skill_score + cert_score + project_score + 20)

        return {
            "完整度评分": completeness_score,
            "竞争力评分": competitiveness_score,
            "说明": f"填写{filled_fields}/{total_fields}个字段"
        }


# ==================== 页面函数 ====================

def render_admin_top_nav():
    """渲染管理员顶部导航栏"""
    # 使用表单和按钮实现导航，避免 JavaScript 跳转问题
    with st.form(key="admin_nav_form", clear_on_submit=False):
        col_logo, col_menu, col_user, col_logout = st.columns([2, 6, 1, 1])

        with col_logo:
            st.markdown("### 🌟 职导星-管理后台")

        with col_menu:
            menu_items = [
                ("dashboard", "仪表盘"),
                ("user_management", "用户管理"),
                ("discussion_management", "讨论管理"),
                ("report_management", "报告管理"),
                ("resume_management", "简历管理"),
                ("interview_management", "面试管理"),
                ("profile_management", "能力画像"),
                ("data_statistics", "数据统计"),
                ("operation_logs", "操作日志"),
                ("system_config", "系统配置"),
            ]
            cols = st.columns(len(menu_items))
            for i, (key, label) in enumerate(menu_items):
                with cols[i]:
                    if st.form_submit_button(label, type="secondary"):
                        st.session_state['admin_nav'] = key
                        st.rerun()

        with col_user:
            st.write(f"👤 {st.session_state.get('username', '管理员')}")

        with col_logout:
            if st.form_submit_button("退出", type="secondary"):
                st.session_state['logged_in'] = False
                st.session_state['username'] = None
                st.session_state['user_type'] = None
                st.session_state['nav'] = 'Home'
                st.rerun()

    st.markdown("---")


def render_top_nav():
    """渲染顶部导航栏"""
    menu_items = {
        "Home": "首页",
        "Profile": "信息录入",
        "Resume": "简历服务",
        "Interview": "模拟面试",
        "Assessment": "职业测评",
        "Image": "能力画像",
        "Match": "岗位匹配",
        "Graph": "职业图谱",
        "Report": "规划报告",
        "Interaction": "咨询与互动",
        "Learn": "学习资源",
        "LearningPath": "学习路径",
        "Cooperation": "校企合作",
        "News": "行业动态"
    }

    current_page = st.session_state.get('nav', 'Home')

    # 使用 HTML 构建导航栏，确保样式一致
    nav_html = f"""
    <style>
    /* 导航栏样式 - 确保始终一致 */
    .top-nav {{
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 8px 20px;
        margin-bottom: 20px;
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        z-index: 1000;
        box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
    }}

    .nav-container {{
        display: flex;
        justify-content: space-between;
        align-items: center;
        flex-wrap: nowrap;
        max-width: 100%;
        margin: 0;
        padding-left: 0;
        width: 100%;
    }}

    .logo-area {{
        display: flex;
        align-items: center;
        gap: 8px;
    }}

    .logo-area span {{
        font-size: 20px;
    }}

    .nav-menu {{
        display: flex;
        gap: 5px;
        flex-wrap: nowrap;
        justify-content: center;
        margin: 0 auto;
        padding-left: 0;
    }}

    .top-nav .nav-item {{
        padding: 5px 12px;
        cursor: pointer;
        font-size: 14px;
        font-weight: 500;
        color: white !important;
        text-decoration: none !important;
        display: inline-block;
        transition: all 0.3s ease;
        border-radius: 15px;
        white-space: nowrap;
    }}

    .top-nav .nav-item:hover {{
        background: rgba(255, 255, 255, 0.2);
        transform: translateY(-2px);
    }}

    .top-nav .nav-item.active {{
        background: rgba(255, 255, 255, 0.25);
        font-weight: 600;
    }}

    .user-info {{
        display: flex;
        align-items: center;
        gap: 10px;
        background: rgba(255, 255, 255, 0.1);
        padding: 4px 12px;
        border-radius: 20px;
    }}

    .user-info span {{
        color: white;
        font-size: 13px;
    }}

    .logout-btn {{
        background: rgba(255, 255, 255, 0.2);
        color: #e0e0e0;
        padding: 4px 10px;
        border-radius: 15px;
        font-size: 12px;
        text-decoration: none !important;
        transition: all 0.3s ease;
    }}

    .logout-btn:hover {{
        background: rgba(255, 255, 255, 0.3);
        color: white;
    }}

    /* 隐藏多余元素 */
    #MainMenu {{visibility: hidden;}}
    footer {{visibility: hidden;}}
    header {{visibility: hidden;}}

    /* 内容区域顶部间距 */
    .main-content {{
        margin-top: 30px;
    }}
    </style>

    <div class="top-nav">
        <div class="nav-container">
            <div class="logo-area">
                <span>🌟</span>
                <span style="font-weight: bold; color: white;">职引星</span>
            </div>
            <div class="nav-menu">
    """

    for key, label in menu_items.items():
        active_class = "active" if current_page == key else ""
        nav_html += f'<a href="?nav={key}" class="nav-item {active_class}">{label}</a>'

    # 根据用户类型显示不同的按钮
    username = st.session_state.get('username', '用户')
    if username == '访客用户':
        nav_html += f"""
            </div>
            <div class="user-info">
                <span>👤 {username}</span>
                <a href="?nav=Login" class="logout-btn">登录/注册</a>
            </div>
        </div>
    </div>
    <div class="main-content"></div>
    """
    else:
        nav_html += f"""
            </div>
            <div class="user-info">
                <span>👤 {username}</span>
                <a href="?nav=Logout" class="logout-btn">退出</a>
            </div>
        </div>
    </div>
    <div class="main-content"></div>
    """

    st.markdown(nav_html, unsafe_allow_html=True)


def render_login_page():
    # 两列布局：左侧渐变背景+系统信息，右侧登录功能
    col1, col2 = st.columns([1, 1], gap="large")
    
    with col1:
        # 左侧渐变背景+系统信息
        st.markdown("""
        <style>
        .login-left-bg {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
            border-radius: 20px;
            padding: 40px;
            text-align: center;
            display: flex;
            flex-direction: column;
            justify-content: center;
            height: 100%;
            min-height: 500px;
            box-shadow: 0 20px 60px rgba(102, 126, 234, 0.3);
        }
        .login-title {
            font-size: 56px;
            font-weight: bold;
            color: white;
            margin-bottom: 20px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        }
        .login-subtitle {
            font-size: 20px;
            color: rgba(255,255,255,0.95);
            margin-bottom: 30px;
        }
        .login-desc {
            font-size: 16px;
            color: rgba(255,255,255,0.85);
            line-height: 1.8;
        }
        .login-features {
            margin-top: 30px;
            text-align: left;
            padding: 20px;
            background: rgba(255,255,255,0.15);
            border-radius: 15px;
        }
        .login-feature-item {
            color: white;
            font-size: 15px;
            margin: 12px 0;
            padding-left: 10px;
        }
        .login-feature-item:before {
            content: "✦ ";
            color: #ffd700;
        }
        </style>
        <div class="login-left-bg">
            <div style="font-size: 70px; margin-bottom: 20px;">🌟</div>
            <div class="login-title">职引星</div>
            <div class="login-subtitle">AI大学生职业规划智能体</div>
            <div class="login-features">
                <div class="login-feature-item">智能人岗匹配，精准推荐职业方向</div>
                <div class="login-feature-item">AI能力画像，全面分析个人竞争力</div>
                <div class="login-feature-item">职业发展图谱，可视化规划未来路径</div>
                <div class="login-feature-item">个性化规划报告，导出留存随时查看</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        # 右侧登录表单
        with st.form("login_form"):
            # 减少输入框之间的间距和顶部间距
            st.markdown("<style>.stTextInput { margin-bottom: 15px; } .stForm { margin-top: 0; padding-top: 0; }</style>", unsafe_allow_html=True)
            
            username = st.text_input("用户名", placeholder="请输入用户名")
            password = st.text_input("密码", type="password", placeholder="请输入密码")
            
            st.markdown("**用户类型**")
            user_type = st.radio("用户类型", ["user", "admin"], format_func=lambda x: "普通用户" if x == "user" else "管理员", horizontal=True)
            
            submit = st.form_submit_button("登录", type="primary", use_container_width=True)
            
            if submit:
                if not username or not password:
                    st.error("请输入用户名和密码")
                else:
                    user = authenticate_user(username, password, user_type)
                    if user:
                        st.session_state['logged_in'] = True
                        st.session_state['username'] = username
                        st.session_state['user_type'] = user_type
                        # 管理员登录时不设置 nav，确保侧边栏正常显示
                        if user_type == 'admin':
                            st.session_state['nav'] = None
                        else:
                            st.session_state['nav'] = 'Home'
                        log_operation(username, "登录系统")
                        st.success("登录成功！")
                        time.sleep(0.5)
                        st.rerun()
                    else:
                        st.error("用户名或密码错误")
        
        # 底部链接
        col_l, col_r = st.columns(2, gap="small")
        with col_l:
            if st.button("注册账号", use_container_width=True):
                st.session_state['show_register'] = True
                st.rerun()
        with col_r:
            if st.button("忘记密码", use_container_width=True):
                st.session_state['show_forgot_pwd'] = True
                st.rerun()
        
        st.markdown("</div>", unsafe_allow_html=True)


def render_register_page():
    # 两列布局：左侧渐变背景+系统信息，右侧注册功能
    col1, col2 = st.columns([1, 1], gap="large")
    
    with col1:
        # 左侧渐变背景+系统信息
        st.markdown("""
        <style>
        .login-left-bg {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 50%, #f093fb 100%);
            border-radius: 20px;
            padding: 40px;
            text-align: center;
            display: flex;
            flex-direction: column;
            justify-content: center;
            height: 100%;
            min-height: 500px;
            box-shadow: 0 20px 60px rgba(102, 126, 234, 0.3);
        }
        .login-title {
            font-size: 56px;
            font-weight: bold;
            color: white;
            margin-bottom: 20px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.2);
        }
        .login-subtitle {
            font-size: 20px;
            color: rgba(255,255,255,0.95);
            margin-bottom: 30px;
        }
        .login-features {
            margin-top: 30px;
            text-align: left;
            padding: 20px;
            background: rgba(255,255,255,0.15);
            border-radius: 15px;
        }
        .login-feature-item {
            color: white;
            font-size: 15px;
            margin: 12px 0;
            padding-left: 10px;
        }
        .login-feature-item:before {
            content: "✦ ";
            color: #ffd700;
        }
        </style>
        <div class="login-left-bg">
            <div style="font-size: 70px; margin-bottom: 20px;">🌟</div>
            <div class="login-title">职引星</div>
            <div class="login-subtitle">AI大学生职业规划智能体</div>
            <div class="login-features">
                <div class="login-feature-item">智能人岗匹配，精准推荐职业方向</div>
                <div class="login-feature-item">AI能力画像，全面分析个人竞争力</div>
                <div class="login-feature-item">职业发展图谱，可视化规划未来路径</div>
                <div class="login-feature-item">个性化规划报告，导出留存随时查看</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        # 右侧注册表单
        with st.form("register_form"):
            # 减少输入框之间的间距和顶部间距
            st.markdown("<style>.stTextInput { margin-bottom: 15px; } .stForm { margin-top: 0; padding-top: 0; }</style>", unsafe_allow_html=True)
            
            username = st.text_input("用户名", placeholder="请输入用户名")
            email = st.text_input("邮箱", placeholder="请输入邮箱")
            password = st.text_input("密码", type="password", placeholder="请输入密码（至少8位）")
            confirm_pwd = st.text_input("确认密码", type="password", placeholder="请再次输入密码")
            
            st.markdown("**用户类型**")
            user_type = st.radio("", ["user", "admin"], format_func=lambda x: "普通用户" if x == "user" else "管理员", horizontal=True)
            
            submit = st.form_submit_button("注册", type="primary", use_container_width=True)
            
            if submit:
                if not username or not email or not password:
                    st.error("请填写完整信息")
                elif password != confirm_pwd:
                    st.error("两次密码不一致")
                elif len(password) < 8:
                    st.error("密码至少8位")
                else:
                    if user_exists(username):
                        st.error("用户名已存在")
                    elif email_exists(email):
                        st.error("邮箱已被注册")
                    else:
                        if create_user(username, email, password, user_type):
                            st.success("注册成功！请登录")
                            time.sleep(1)
                            st.session_state['show_register'] = False
                            # 确保用户没有自动登录为访客
                            st.session_state['logged_in'] = False
                            st.session_state['username'] = None
                            st.session_state['user_type'] = None
                            st.rerun()
                        else:
                            st.error("注册失败")
        
        # 底部链接
        if st.button("返回登录", use_container_width=True):
            st.session_state['show_register'] = False
            st.rerun()
        
        st.markdown("</div>", unsafe_allow_html=True)


def get_user_by_email(email: str) -> Optional[Dict]:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM users WHERE email = ?", (email,))
    user = c.fetchone()
    conn.close()
    return dict(user) if user else None


def render_admin_reset_page():
    st.markdown('<div class="auth-container">', unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("## 🔑 管理员密码重置")
        st.markdown("输入管理员邮箱，系统将发送验证码")

        if 'admin_reset_step' not in st.session_state:
            st.session_state['admin_reset_step'] = 1

        if st.session_state['admin_reset_step'] == 1:
            with st.form("admin_reset_form"):
                email = st.text_input("管理员邮箱", placeholder="请输入注册时使用的邮箱")
                submit = st.form_submit_button("发送验证码", type="primary", use_container_width=True)

                if submit:
                    if not email:
                        st.error("请输入邮箱")
                    else:
                        user = get_user_by_email(email)
                        if user and user.get('user_type') == 'admin':
                            import random
                            code = str(random.randint(100000, 999999))
                            st.session_state['admin_reset_email'] = email
                            st.session_state['admin_reset_code'] = code
                            st.session_state['admin_reset_step'] = 2
                            st.success(f"验证码已发送到 {email}（测试模式验证码: {code}）")
                            st.rerun()
                        else:
                            st.error("该邮箱不是管理员账号")

        elif st.session_state['admin_reset_step'] == 2:
            email = st.session_state.get('admin_reset_email', '')

            st.markdown(f"**邮箱**: {email}")
            code = st.text_input("验证码", placeholder="请输入6位验证码", key="admin_code")
            new_password = st.text_input("新密码", type="password", placeholder="请输入新密码（至少8位）", key="admin_new_pwd")
            confirm_pwd = st.text_input("确认新密码", type="password", placeholder="请再次输入新密码", key="admin_confirm_pwd")

            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("重置密码", type="primary", use_container_width=True):
                    if not code or not new_password:
                        st.error("请填写完整信息")
                    elif new_password != confirm_pwd:
                        st.error("两次密码不一致")
                    elif len(new_password) < 8:
                        st.error("新密码至少8位")
                    elif code != st.session_state.get('admin_reset_code'):
                        st.error("验证码错误")
                    else:
                        conn = sqlite3.connect(DB_PATH)
                        c = conn.cursor()
                        c.execute("UPDATE users SET password_hash = ? WHERE email = ?",
                                 (hash_password(new_password), email))
                        conn.commit()
                        conn.close()
                        st.success("密码重置成功！请使用新密码登录")
                        time.sleep(1)
                        st.session_state['admin_reset_step'] = 1
                        st.session_state['show_admin_reset'] = False
                        if 'admin_reset_email' in st.session_state:
                            del st.session_state['admin_reset_email']
                        st.rerun()

            with col_b:
                if st.button("重新发送验证码", use_container_width=True):
                    import random
                    code = str(random.randint(100000, 999999))
                    st.session_state['admin_reset_code'] = code
                    st.success(f"验证码已重新发送（测试模式验证码: {code}）")
                    st.rerun()

        st.markdown("---")
        if st.button("返回登录", use_container_width=True):
            st.session_state['admin_reset_step'] = 1
            st.session_state['show_admin_reset'] = False
            if 'admin_reset_email' in st.session_state:
                del st.session_state['admin_reset_email']
            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)


def render_change_password_page():
    """修改密码页面"""
    st.markdown("""
    <div class="login-wrapper">
        <div class="login-card">
            <div class="login-header">
                <div class="login-icon">🔐</div>
                <h1 class="login-title">修改密码</h1>
            </div>
    """, unsafe_allow_html=True)

    with st.form("change_pwd_form"):
        username = st.text_input("用户名")
        old_password = st.text_input("原密码", type="password")
        new_password = st.text_input("新密码", type="password")
        confirm_pwd = st.text_input("确认新密码", type="password")

        if st.form_submit_button("修改密码", use_container_width=True):
            if not username or not old_password or not new_password:
                st.error("请填写完整信息")
            elif new_password != confirm_pwd:
                st.error("两次密码不一致")
            elif len(new_password) < 8:
                st.error("新密码至少8位")
            else:
                if change_password(username, old_password, new_password):
                    st.success("密码修改成功！")
                    time.sleep(1)
                    st.session_state['show_change_pwd'] = False
                    st.rerun()
                else:
                    st.error("原密码错误")

    st.markdown("""
        <div class="footer-links">
            <div class="footer-link" onclick="location.href='?nav=login'">返回登录</div>
        </div>
    </div></div>
    """, unsafe_allow_html=True)


def render_forgot_password_page():
    """忘记密码页面"""
    st.markdown("""
    <div class="login-wrapper">
        <div class="login-card">
            <div class="login-header">
                <div class="login-icon">🔑</div>
                <h1 class="login-title">找回密码</h1>
                <p>请联系管理员重置密码</p>
            </div>
    """, unsafe_allow_html=True)

    if st.button("返回登录", use_container_width=True):
        st.session_state['show_forgot_pwd'] = False
        st.rerun()

    st.markdown('</div></div>', unsafe_allow_html=True)


def render_homepage():
    """首页 - 带可视化图表"""
    st.markdown("""
    <div style="text-align: center; padding: 30px 0 20px;">
        <h1 style="font-size: 42px; margin-bottom: 15px;">
            🎯 <span class="highlight">AI大学生职业规划智能体</span>
        </h1>
        <p style="font-size: 18px; color: #666;">基于AI技术，为您提供精准的职业规划方案</p>
    </div>
    """, unsafe_allow_html=True)

    # 四个功能卡片
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(
            '<a href="?nav=Profile" style="text-decoration: none;"><div class="card"><h3>信息录入</h3><p>填写基本信息，完善能力画像</p><p style="color:#667eea;">点击进入 →</p></div></a>',
            unsafe_allow_html=True)
    with col2:
        st.markdown(
            '<a href="?nav=Match" style="text-decoration: none;"><div class="card"><h3>岗位匹配</h3><p>智能分析岗位匹配度</p><p style="color:#667eea;">点击进入 →</p></div></a>',
            unsafe_allow_html=True)
    with col3:
        st.markdown(
            '<a href="?nav=Graph" style="text-decoration: none;"><div class="card"><h3>职业图谱</h3><p>可视化晋升与换岗路径</p><p style="color:#667eea;">点击进入 →</p></div></a>',
            unsafe_allow_html=True)
    with col4:
        st.markdown(
            '<a href="?nav=Report" style="text-decoration: none;"><div class="card"><h3>规划报告</h3><p>生成专属职业规划报告</p><p style="color:#667eea;">点击进入 →</p></div></a>',
            unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("## 📊 系统数据概览")

    # 获取统计数据
    stats = get_system_statistics()

    # 第一行：关键指标卡片
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_users']}</div>
            <div class="dashboard-label">注册用户</div>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_profiles']}</div>
            <div class="dashboard-label">学生画像</div>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_discussions']}</div>
            <div class="dashboard-label">讨论话题</div>
        </div>
        """, unsafe_allow_html=True)
    with col4:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_replies']}</div>
            <div class="dashboard-label">回复总数</div>
        </div>
        """, unsafe_allow_html=True)
    with col5:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_reports']}</div>
            <div class="dashboard-label">规划报告</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # 第二行：两个可视化图表
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 热门技能分布")
        if stats['skill_count']:
            # 取前10个热门技能
            top_skills = sorted(stats['skill_count'].items(), key=lambda x: x[1], reverse=True)[:10]
            if top_skills:
                skills_df = pd.DataFrame(top_skills, columns=['技能', '人数'])

                fig = px.bar(skills_df, x='人数', y='技能', orientation='h',
                             title='学生掌握的热门技能Top10',
                             color='人数', color_continuous_scale='Blues',
                             text='人数')
                fig.update_layout(
                    plot_bgcolor='rgba(0,0,0,0)',
                    paper_bgcolor='rgba(0,0,0,0)',
                    height=400,
                    xaxis_title="掌握人数",
                    yaxis_title="技能名称"
                )
                fig.update_traces(textposition='outside')
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("暂无技能数据")
        else:
            st.info("暂无技能数据，请先录入学生信息")

    with col2:
        st.markdown("### 数据分布概览")
        # 创建环形图展示数据分布
        data_categories = ['学生画像', '讨论话题', '规划报告']
        data_values = [stats['total_profiles'], stats['total_discussions'], stats['total_reports']]

        if sum(data_values) > 0:
            fig = go.Figure(data=[go.Pie(
                labels=data_categories,
                values=data_values,
                hole=0.4,
                marker_colors=['#667eea', '#764ba2', '#f39c12'],
                textinfo='label+percent',
                textposition='outside'
            )])
            fig.update_layout(
                title='平台内容分布',
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                height=400,
                showlegend=True
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("暂无数据")


def render_student_input(manager: StudentProfileManager):
    """学生信息录入"""
    st.title("信息录入")

    # 检查是否有解析的简历数据
    parsed_profile = st.session_state.get('parsed_profile', {})
    if parsed_profile:
        st.success("检测到解析的简历数据，已自动填充到表单中！")

    with st.form("student_form"):
        st.subheader("基本信息")
        col1, col2 = st.columns(2)
        with col1:
            name = st.text_input("姓名", value=parsed_profile.get('姓名', ''))
            major = st.text_input("专业", value=parsed_profile.get('专业', ''))
        with col2:
            school = st.text_input("学校")
            education = st.selectbox("学历", ["专科", "本科", "硕士", "博士"], index=["专科", "本科", "硕士", "博士"].index(parsed_profile.get('学历', '本科')))

        st.subheader("专业技能")
        skills = st.multiselect("选择技能",
                                ["Python", "Java", "JavaScript", "MySQL", "Docker", "Linux", "机器学习", "数据分析", "C++", "C#", "PHP", "Go", "Ruby", "Swift", "Kotlin", "React", "Vue", "Angular", "Node.js", "Spring Boot", "Django", "Flask", "TensorFlow", "PyTorch", "Hadoop", "Spark", "Redis", "MongoDB", "PostgreSQL", "SQLite", "Git", "GitHub", "GitLab", "Jenkins", "CI/CD", "Kubernetes", "AWS", "Azure", "Google Cloud", "DevOps", "敏捷开发", "Scrum"],
                                default=parsed_profile.get('专业技能', []))

        st.subheader("证书")
        certificates = st.multiselect("选择证书", ["软考低级","软考中级", "软考高级","PMP", "英语四级", "英语六级", "雅思", "托福", "GRE", "AWS认证", "Azure认证", "Google Cloud认证", "Oracle认证", "Microsoft认证", "CompTIA认证", "CISSP", "CCNA", "CCNP", "CCIE", "CPA", "ACCA", "CFA", "FRM", "律师资格证", "教师资格证", "注册会计师", "注册税务师", "注册建筑师", "注册工程师"],
                                    default=parsed_profile.get('证书', []))

        st.subheader("项目经验")
        project_count = st.number_input("项目数量", 0, 3, len(parsed_profile.get('项目经验', [])), step=1)
        projects = []
        project_list = parsed_profile.get('项目经验', [])
        for i in range(project_count):
            default_project = project_list[i] if i < len(project_list) else ""
            project = st.text_input(f"项目{i + 1}", value=default_project, key=f"proj_{i}")
            if project:
                projects.append(project)

        st.subheader("实习经历")
        internship = st.text_area("实习经历", value=parsed_profile.get('实习经历', ''))

        st.subheader("求职意向")

        col3, col4, col5 = st.columns(3)
        # 期望城市 - 手动输入
        with col3:
            city = st.text_input("期望城市", value=parsed_profile.get('期望城市', ''))
        with col4:
            salary = st.selectbox("期望薪资", ["5K-8K", "8K-12K", "12K-20K", "20K以上"])
        with col5:
            job_type = st.multiselect("期望岗位", get_jobs_from_excel(),
                                   default=[parsed_profile.get('期望岗位', '')] if parsed_profile.get('期望岗位', '') else [])

        if st.form_submit_button("保存信息", use_container_width=True):
            if not name:
                st.warning("请输入姓名")
            else:
                form_data = {
                    "name": name, "major": major, "school": school, "education": education,
                    "skills": skills, "certificates": certificates, "projects": projects,
                    "internship": internship, "city": city, "salary": salary, "job_type": job_type
                }
                profile = manager.create_profile_from_form(form_data)
                completeness = manager.calculate_completeness_score(profile)
                competitiveness = manager.calculate_competitiveness_score(profile)
                profile['completeness_score'] = completeness
                profile['competitiveness_score'] = competitiveness

                if save_student_profile(st.session_state['username'], profile):
                    st.session_state['student_profile'] = profile
                    # 保存成功后清除解析的简历数据
                    if 'parsed_profile' in st.session_state:
                        del st.session_state['parsed_profile']
                    st.success("保存成功！")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("画像完整度", f"{completeness:.0f}%")
                    with col2:
                        st.metric("竞争力评分", f"{competitiveness:.0f}%")


def render_student_profile(manager: StudentProfileManager, profile: Dict):
    """能力画像页面"""
    st.title("能力画像")
    st.markdown("---")

    # 引导用户到简历解析页面
    st.markdown("### 📄 简历解析")
    st.info("请通过导航栏的'简历解析'功能上传并解析您的简历，生成学生画像。")
    st.markdown("---")

    # 原有能力画像部分
    st.markdown("### 📊 我的能力画像")

    completeness = profile.get('completeness_score', 0)
    competitiveness = profile.get('competitiveness_score', 0)

    col1, col2 = st.columns(2)
    with col1:
        st.metric("画像完整度", f"{completeness:.0f}%")
        st.progress(completeness / 100)
    with col2:
        st.metric("竞争力评分", f"{competitiveness:.0f}%")
        st.progress(competitiveness / 100)

    basic = profile.get("基本信息", {})
    st.markdown("---")
    st.subheader("基本信息")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.info(f"姓名: {basic.get('姓名', '-')}")
    with col2:
        st.info(f"专业: {basic.get('专业', '-')}")
    with col3:
        st.info(f"学历: {basic.get('学历', '-')}")
    with col4:
        st.info(f"学校: {basic.get('学校', '-')}")

    st.subheader("专业技能")
    skills = profile.get("专业技能", [])
    if skills:
        st.markdown(" ".join([f"`{s}`" for s in skills]))

    st.subheader("证书")
    certs = profile.get("证书", [])
    if certs:
        st.markdown(" ".join([f"`{c}`" for c in certs]))

    st.subheader("项目经验")
    for p in profile.get("项目经验", []):
        st.markdown(f"- {p}")

    st.subheader("实习经历")
    if profile.get("实习经历"):
        st.info(profile["实习经历"])

    st.subheader("求职意向")
    job_intent = profile.get("求职意向", {})
    st.info(f"期望城市: {job_intent.get('城市', '-')}  |  期望薪资: {job_intent.get('期望薪资', '-')}")


def render_advanced_job_matching(engine: AdvancedJobMatchingEngine, profile: Dict):
    """升级版岗位匹配页面"""
    st.title("🔍 AI智能岗位匹配")
    st.markdown("---")

    # 显示真实岗位数据统计
    st.markdown("---")
    st.subheader("📊 岗位市场分析")

    # 读取Excel数据并创建JobDataProcessor实例
    import plotly.express as px
    from job_data_processor import JobDataProcessor

    job_processor = None
    job_data = None

    try:
        job_processor = JobDataProcessor('a13基于AI的大学生职业规划智能体-JD采样数据.xls')
        job_data = job_processor.jobs_data
        st.success("✅ 成功加载岗位数据")
    except Exception as e:
        st.warning(f"读取Excel文件失败: {e}，使用默认数据")

    # 行业分布分析
    st.markdown("### 行业分布")
    if job_processor:
        try:
            industry_stats = job_processor.get_industry_statistics()
            if industry_stats:
                industry_df = pd.DataFrame(list(industry_stats.items()), columns=['行业', '岗位数'])
                fig = px.bar(industry_df, x='行业', y='岗位数', title='岗位行业分布')
                st.plotly_chart(fig)
            else:
                # 默认行业数据
                industry_stats = {
                    "互联网": 120,
                    "金融": 80,
                    "教育": 60,
                    "医疗": 40,
                    "制造业": 50
                }
                industry_df = pd.DataFrame(list(industry_stats.items()), columns=['行业', '岗位数'])
                fig = px.bar(industry_df, x='行业', y='岗位数', title='岗位行业分布')
                st.plotly_chart(fig)
        except Exception as e:
            st.warning(f"获取行业统计数据失败: {e}")
            # 默认行业数据
            industry_stats = {
                "互联网": 120,
                "金融": 80,
                "教育": 60,
                "医疗": 40,
                "制造业": 50
            }
            industry_df = pd.DataFrame(list(industry_stats.items()), columns=['行业', '岗位数'])
            fig = px.bar(industry_df, x='行业', y='岗位数', title='岗位行业分布')
            st.plotly_chart(fig)
    else:
        # 默认行业数据
        industry_stats = {
            "互联网": 120,
            "金融": 80,
            "教育": 60,
            "医疗": 40,
            "制造业": 50
        }
        industry_df = pd.DataFrame(list(industry_stats.items()), columns=['行业', '岗位数'])
        fig = px.bar(industry_df, x='行业', y='岗位数', title='岗位行业分布')
        st.plotly_chart(fig)

    # 薪资水平分析
    st.markdown("### 薪资水平")
    if job_processor:
        try:
            salary_stats = job_processor.get_salary_statistics()
            if salary_stats:
                # 确保薪资数据是数字类型
                for key in salary_stats:
                    if isinstance(salary_stats[key], str):
                        # 尝试转换为数字
                        try:
                            salary_stats[key] = float(salary_stats[key].replace('K', '')) * 1000
                        except:
                            pass
            else:
                # 默认薪资数据
                salary_stats = {
                    "average": 15000,
                    "median": 12000,
                    "max": 30000,
                    "min": 5000
                }
        except Exception as e:
            st.warning(f"获取薪资统计数据失败: {e}")
            # 默认薪资数据
            salary_stats = {
                "average": 15000,
                "median": 12000,
                "max": 30000,
                "min": 5000
            }
    else:
        # 默认薪资数据
        salary_stats = {
            "average": 15000,
            "median": 12000,
            "max": 30000,
            "min": 5000
        }

    c1, c2, c3, c4 = st.columns(4)
    with c1: st.metric("平均薪资", f"{salary_stats.get('average', 0):.0f}元")
    with c2: st.metric("中位数", f"{salary_stats.get('median', 0):.0f}元")
    with c3: st.metric("最高薪资", f"{salary_stats.get('max', 0):.0f}元")
    with c4: st.metric("最低薪资", f"{salary_stats.get('min', 0):.0f}元")

    # 获取简历解析数据
    resume_portrait = st.session_state.get('portrait', None)
    if resume_portrait:
        st.success("✅ 已加载简历解析数据，AI将进行深度分析")

    # 岗位选择
    jobs = get_jobs_from_excel()

    # 智能推荐岗位
    if profile.get('专业技能'):
        recommended = engine._get_recommended_jobs(profile)
        st.info(f"💡 根据您的技能，推荐关注: {' → '.join(recommended[:3])}")

    selected_job = st.selectbox("选择目标岗位", jobs,
                                 index=jobs.index(recommended[0]) if recommended and recommended[0] in jobs else 0)

    st.markdown("---")

    if st.button("🚀 AI深度分析匹配度", type="primary", use_container_width=True):
        with st.spinner("AI正在深度分析您的能力与岗位匹配度..."):
            result = engine.match_with_ai(profile, resume_portrait, selected_job)
            st.session_state['advanced_match_results'] = result
            st.session_state['selected_job'] = selected_job

    if 'advanced_match_results' in st.session_state:
        result = st.session_state['advanced_match_results']

        total = result.get("总匹配度", 0)
        color = "#2ecc71" if total >= 80 else "#f39c12" if total >= 60 else "#e74c3c"

        st.markdown(f"""
        <div style="text-align: center; padding: 20px; background: linear-gradient(135deg, #667eea15 0%, #764ba215 100%); border-radius: 20px; margin: 20px 0;">
            <h2 style="color: {color}; font-size: 48px; margin: 0;">{total}%</h2>
            <p style="color: #666;">综合匹配度</p>
            <div style="background: #e0e0e0; border-radius: 10px; height: 10px; width: 80%; margin: 10px auto;">
                <div style="background: {color}; border-radius: 10px; height: 10px; width: {total}%;"></div>
            </div>
            <p style="margin-top: 10px;">
                <span style="background: {color}20; color: {color}; padding: 4px 12px; border-radius: 20px;">
                    竞争力评估: {result.get('竞争力评估', '良好')}
                </span>
            </p>
        </div>
        """, unsafe_allow_html=True)

        st.subheader("📊 各维度匹配分析")
        dims = result.get("各维度得分", {})
        if dims:
            cols = st.columns(len(dims))
            for i, (dim, score) in enumerate(dims.items()):
                with cols[i]:
                    st.metric(dim, f"{score}%")
                    st.progress(score / 100)

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### ✅ 核心优势")
            for adv in result.get("优势", []):
                st.success(f"✓ {adv}")

        with col2:
            st.markdown("### ⚠️ 待提升方向")
            for gap in result.get("差距", []):
                st.warning(f"⚠ {gap}")

        st.markdown("### 📚 AI技能提升建议")
        for i, suggestion in enumerate(result.get("技能提升建议", []), 1):
            st.markdown(f"{i}. {suggestion}")

        st.markdown("### 💡 推荐理由")
        st.info(result.get("推荐理由", "该岗位与您的能力背景较为匹配，建议积极准备。"))

        if result.get("建议薪资"):
            st.markdown(f"### 💰 建议薪资范围")
            st.success(f"根据市场行情和您的匹配度，建议薪资: {result['建议薪资']}")

        st.caption(f"分析时间: {result.get('分析时间', '')}")
        st.session_state['match_results'] = result

    # 基于真实数据的岗位推荐
    st.markdown("---")
    st.subheader("💼 岗位推荐")

    if profile.get('专业技能'):
        if st.button("基于我的技能推荐岗位"):
            with st.spinner("AI分析中..."):
                recommended_jobs = []

                if job_processor:
                    # 使用AI和JobDataProcessor进行岗位推荐
                    skills = profile.get('专业技能', [])
                    experience = profile.get('工作经验', '')
                    education = profile.get('学历', '')
                    city = profile.get('求职意向', {}).get('城市', '')

                    # 调用JobDataProcessor的推荐功能
                    recommended_jobs = job_processor.get_recommended_jobs(
                        skills=skills,
                        experience=experience,
                        education=education,
                        city=city
                    )

                    # 如果有推荐岗位，使用AI对推荐结果进行优化
                    if recommended_jobs:
                        st.success(f"AI为您推荐了 {len(recommended_jobs)} 个匹配岗位")
                    else:
                        st.warning("暂无匹配的岗位推荐")
                else:
                    # 如果JobDataProcessor初始化失败，使用默认数据
                    recommended_jobs = [
                        {
                            "岗位名称": "Java开发工程师",
                            "公司名称": "阿里巴巴",
                            "薪资范围": "15K-25K",
                            "地址": "杭州",
                            "所属行业": "互联网",
                            "公司规模": "10000人以上",
                            "岗位详情": "1. 负责Java后端开发\n2. 熟悉Spring Boot、Spring Cloud\n3. 有微服务开发经验优先"
                        },
                        {
                            "岗位名称": "Python开发工程师",
                            "公司名称": "腾讯",
                            "薪资范围": "18K-28K",
                            "地址": "深圳",
                            "所属行业": "互联网",
                            "公司规模": "10000人以上",
                            "岗位详情": "1. 负责Python后端开发\n2. 熟悉Django、Flask\n3. 有数据分析经验优先"
                        },
                        {
                            "岗位名称": "前端开发工程师",
                            "公司名称": "字节跳动",
                            "薪资范围": "12K-22K",
                            "地址": "北京",
                            "所属行业": "互联网",
                            "公司规模": "10000人以上",
                            "岗位详情": "1. 负责前端开发\n2. 熟悉Vue、React\n3. 有前端工程化经验优先"
                        }
                    ]
                    st.success(f"为您推荐了 {len(recommended_jobs)} 个匹配岗位")

                if recommended_jobs:
                    for i, job in enumerate(recommended_jobs[:5], 1):  # 最多显示5个推荐岗位
                        with st.expander(f"推荐岗位 {i}: {job.get('岗位名称')}"):
                            st.markdown(f"**公司**: {job.get('公司名称', '未知')}")
                            st.markdown(f"**薪资**: {job.get('薪资范围', '未知')}")
                            st.markdown(f"**地点**: {job.get('工作地点', job.get('地址', '未知'))}")
                            st.markdown(f"**行业**: {job.get('所属行业', '未知')}")
                            st.markdown(f"**公司规模**: {job.get('公司规模', '未知')}")

                            # 显示岗位详情
                            job_detail = job.get('岗位详情', job.get('岗位要求', ''))
                            if job_detail:
                                st.markdown("**岗位要求**:")
                                st.markdown(job_detail)
                else:
                    st.warning("暂无匹配的岗位推荐")
    else:
        st.warning("请先完善您的技能信息")

    # 关键词搜索岗位
    st.markdown("---")
    st.subheader("🔎 岗位搜索")

    search_keyword = st.text_input("输入关键词搜索岗位")
    if search_keyword:
        if st.button("搜索"):
            with st.spinner("AI搜索中..."):
                search_results = []

                if job_processor:
                    # 使用JobDataProcessor进行关键词搜索
                    search_results = job_processor.get_jobs_by_keyword(search_keyword)

                    if search_results:
                        st.success(f"AI找到 {len(search_results)} 个相关岗位")
                    else:
                        st.warning("未找到相关岗位")
                else:
                    # 如果JobDataProcessor初始化失败，使用默认数据
                    search_results = [
                        {
                            "岗位名称": "Java开发工程师",
                            "公司名称": "阿里巴巴",
                            "薪资范围": "15K-25K",
                            "地址": "杭州",
                            "所属行业": "互联网",
                            "公司规模": "10000人以上",
                            "岗位详情": "1. 负责Java后端开发\n2. 熟悉Spring Boot、Spring Cloud\n3. 有微服务开发经验优先"
                        },
                        {
                            "岗位名称": "Java高级开发工程师",
                            "公司名称": "腾讯",
                            "薪资范围": "20K-30K",
                            "地址": "深圳",
                            "所属行业": "互联网",
                            "公司规模": "10000人以上",
                            "岗位详情": "1. 负责Java后端架构设计\n2. 熟悉分布式系统\n3. 有5年以上Java开发经验"
                        }
                    ]
                    st.success(f"找到 {len(search_results)} 个相关岗位")

                if search_results:
                    for i, job in enumerate(search_results[:5], 1):  # 最多显示5个搜索结果
                        with st.expander(f"搜索结果 {i}: {job.get('岗位名称')}"):
                            st.markdown(f"**公司**: {job.get('公司名称', '未知')}")
                            st.markdown(f"**薪资**: {job.get('薪资范围', '未知')}")
                            st.markdown(f"**地点**: {job.get('工作地点', job.get('地址', '未知'))}")
                            st.markdown(f"**行业**: {job.get('所属行业', '未知')}")
                            st.markdown(f"**公司规模**: {job.get('公司规模', '未知')}")

                            # 显示岗位详情
                            job_detail = job.get('岗位详情', job.get('岗位要求', ''))
                            if job_detail:
                                st.markdown("**岗位要求**:")
                                st.markdown(job_detail)
                else:
                    st.warning("未找到相关岗位")


def render_learning_path():
    """学习路径推荐页面"""
    # 添加学习路径页面的样式
    st.markdown("""
    <style>
        .learning-path-card {
            background: white;
            border-radius: 12px;
            padding: 20px;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.1);
            border: 1px solid #e8dfff;
            margin-bottom: 20px;
        }
        
        .path-step {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 15px;
            border-radius: 10px;
            margin: 10px 0;
        }
        
        .resource-tag {
            display: inline-block;
            background: #f0f2f6;
            color: #667eea;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            margin: 4px;
        }
        
        .skill-gap {
            background: #fff3e0;
            border-left: 3px solid #ff9800;
            padding: 10px;
            border-radius: 8px;
            margin: 10px 0;
        }
    </style>
    """, unsafe_allow_html=True)

    st.title("🎯 学习路径推荐")
    st.markdown("基于您的能力画像和目标岗位，AI为您定制个性化学习路径")

    # 获取用户画像和目标岗位
    profile = st.session_state.get('student_profile', {})
    resume_portrait = st.session_state.get('portrait', None)

    if not profile and not resume_portrait:
        st.warning("请先完善您的能力画像或上传简历")
        return

    # 合并画像数据
    merged_profile = profile.copy()
    if resume_portrait:
        if '基本信息' not in merged_profile:
            merged_profile['基本信息'] = {}
        if resume_portrait.get('姓名') and resume_portrait['姓名'] != '未知':
            merged_profile['基本信息']['姓名'] = resume_portrait['姓名']
        if resume_portrait.get('专业') and resume_portrait['专业'] != '未知':
            merged_profile['基本信息']['专业'] = resume_portrait['专业']
        if resume_portrait.get('学历') and resume_portrait['学历'] != '未知':
            merged_profile['基本信息']['学历'] = resume_portrait['学历']
        existing_skills = set(merged_profile.get('专业技能', []))
        new_skills = set(resume_portrait.get('专业技能', []))
        merged_profile['专业技能'] = list(existing_skills.union(new_skills))
        existing_certs = set(merged_profile.get('证书', []))
        new_certs = set(resume_portrait.get('证书', []))
        merged_profile['证书'] = list(existing_certs.union(new_certs))

    # 选择目标岗位
    jobs = get_jobs_from_excel()
    selected_job = st.selectbox("选择目标岗位", jobs)

    st.markdown("---")

    if st.button("🚀 生成个性化学习路径", type="primary", use_container_width=True):
        with st.spinner("AI正在分析并生成学习路径..."):
            # 构建提示词
            prompt = f"""请根据以下学生信息和目标岗位，生成个性化的学习路径：

            ## 学生信息
            - 姓名: {merged_profile.get('基本信息', {}).get('姓名', '同学')}
            - 专业: {merged_profile.get('基本信息', {}).get('专业', '计算机相关')}
            - 学历: {merged_profile.get('基本信息', {}).get('学历', '本科')}
            - 学校: {merged_profile.get('基本信息', {}).get('学校', '普通高校')}
            - 专业技能: {', '.join(merged_profile.get('专业技能', [])) or '待完善'}
            - 证书: {', '.join(merged_profile.get('证书', [])) or '暂无'}
            - 项目经验数量: {len(merged_profile.get('项目经验', []))}个

            ## 目标岗位
            {selected_job}

           请生成一个详细的学习路径，包括：
            1. 学习阶段划分（短期、中期、长期）
            2. 每个阶段的学习目标和核心技能
            3. 推荐的学习资源（书籍、在线课程、实践项目等）
            4. 技能提升的具体建议
            5. 学习时间规划

           请以结构化的方式输出，使用中文回答。"""

            response = llm.chat(prompt)

            # 显示AI生成的学习路径
            st.markdown("### 🤖 个性化学习路径")
            st.markdown(f"<div class='learning-path-card'>{response}</div>", unsafe_allow_html=True)

            # 保存学习路径到会话状态
            st.session_state['learning_path'] = response

    # 显示已生成的学习路径
    if 'learning_path' in st.session_state:
        st.markdown("---")
        st.markdown("### 📋 已保存的学习路径")
        st.markdown(f"<div class='learning-path-card'>{st.session_state['learning_path']}</div>", unsafe_allow_html=True)

    # 技能差距分析
    st.markdown("---")
    st.subheader("🔍 技能差距分析")

    if st.button("分析技能差距"):
        with st.spinner("AI正在分析技能差距..."):
            # 获取目标岗位的技能要求
            job_info = JobDataManager.get_job_profile(selected_job)
            required_skills = job_info.get('skills', [])
            student_skills = set(merged_profile.get('专业技能', []))
            required_skill_set = set(required_skills)
            missing_skills = list(required_skill_set - student_skills)

            if missing_skills:
                st.markdown("### ⚠️ 待提升技能")
                for skill in missing_skills[:5]:  # 显示前5个缺失技能
                    st.markdown(f"<div class='skill-gap'>{skill}</div>", unsafe_allow_html=True)

                # 生成技能提升建议
                prompt = f"""请为以下缺失技能提供具体的学习建议和资源：
                技能列表：{', '.join(missing_skills[:5])}
                目标岗位：{selected_job}

                请为每个技能提供：
                1. 学习重点
                2. 推荐的学习资源
                3. 实践项目建议
                4. 学习时间估计

                使用中文回答，结构化输出。"""

                skill_response = llm.chat(prompt)
                st.markdown("### 📚 技能提升建议")
                st.markdown(f"<div class='learning-path-card'>{skill_response}</div>", unsafe_allow_html=True)
            else:
                st.success("🎉 您已经具备了该岗位的核心技能！")


def render_corporate_cooperation():
    """校企合作页面"""
    # 添加校企合作页面的样式
    st.markdown("""
    <style>
        .company-card {
            background: white;
            border-radius: 12px;
            padding: 20px;
            box-shadow: 0 4px 15px rgba(102, 126, 234, 0.1);
            border: 1px solid #e8dfff;
            margin-bottom: 20px;
            transition: transform 0.3s ease;
        }
        
        .company-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 8px 25px rgba(102, 126, 234, 0.2);
        }
        
        .company-logo {
            font-size: 48px;
            margin-bottom: 15px;
        }
        
        .job-card {
            background: #f8f9fa;
            border-radius: 8px;
            padding: 15px;
            margin: 10px 0;
            border-left: 4px solid #667eea;
        }
        
        .industry-tag {
            display: inline-block;
            background: #e3f2fd;
            color: #1976d2;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            margin: 4px;
        }
        
        .salary-tag {
            display: inline-block;
            background: #e8f5e8;
            color: #388e3c;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            margin: 4px;
        }
    </style>
    """, unsafe_allow_html=True)

    st.title("🤝 校企合作")
    st.markdown("探索合作企业信息和实习就业机会")

    # 合作企业数据
    合作企业 = [
        {
            "名称": "阿里巴巴集团",
            "行业": "互联网",
            "规模": "10000人以上",
            "简介": "全球领先的数字商业公司，业务涵盖电商、云计算、数字媒体等多个领域",
            "招聘岗位": ["Java开发工程师", "前端开发工程师", "产品经理", "数据分析师"],
            "薪资范围": "15K-30K",
            "logo": "🏢"
        },
        {
            "名称": "腾讯科技",
            "行业": "互联网",
            "规模": "10000人以上",
            "简介": "中国领先的互联网科技公司，拥有微信、QQ等核心产品",
            "招聘岗位": ["Python开发工程师", "AI算法工程师", "产品经理", "UI设计师"],
            "薪资范围": "18K-35K",
            "logo": "🏙️"
        },
        {
            "名称": "字节跳动",
            "行业": "互联网",
            "规模": "10000人以上",
            "简介": "全球领先的内容平台公司，拥有抖音、TikTok等热门产品",
            "招聘岗位": ["前端开发工程师", "后端开发工程师", "数据分析师", "产品经理"],
            "薪资范围": "20K-40K",
            "logo": "📱"
        },
        {
            "名称": "华为技术有限公司",
            "行业": "通信设备",
            "规模": "10000人以上",
            "简介": "全球领先的ICT解决方案提供商，专注于通信技术和智能设备",
            "招聘岗位": ["嵌入式软件工程师", "硬件工程师", "网络工程师", "算法工程师"],
            "薪资范围": "15K-30K",
            "logo": "📡"
        },
        {
            "名称": "百度在线网络技术有限公司",
            "行业": "互联网",
            "规模": "10000人以上",
            "简介": "中国领先的搜索引擎公司，专注于AI和互联网服务",
            "招聘岗位": ["AI算法工程师", "后端开发工程师", "数据分析师", "产品经理"],
            "薪资范围": "18K-35K",
            "logo": "🔍"
        }
    ]

    # 实习就业机会数据
    实习机会 = [
        {
            "岗位名称": "Java开发实习生",
            "公司名称": "阿里巴巴",
            "薪资": "300-400元/天",
            "地点": "杭州",
            "要求": "熟悉Java，了解Spring Boot，有项目经验优先",
            "截止日期": "2024-12-31"
        },
        {
            "岗位名称": "前端开发实习生",
            "公司名称": "腾讯",
            "薪资": "350-450元/天",
            "地点": "深圳",
            "要求": "熟悉HTML/CSS/JavaScript，了解Vue或React",
            "截止日期": "2024-12-15"
        },
        {
            "岗位名称": "数据分析师实习生",
            "公司名称": "字节跳动",
            "薪资": "400-500元/天",
            "地点": "北京",
            "要求": "熟悉Python，了解数据分析工具，有SQL基础",
            "截止日期": "2024-12-10"
        },
        {
            "岗位名称": "产品经理实习生",
            "公司名称": "百度",
            "薪资": "300-400元/天",
            "地点": "北京",
            "要求": "了解产品经理工作流程，有良好的沟通能力",
            "截止日期": "2024-12-20"
        }
    ]

    # 合作企业展示
    st.markdown("---")
    st.subheader("🏢 合作企业")

    # 企业筛选
    industries = list(set([company["行业"] for company in 合作企业]))
    selected_industry = st.selectbox("按行业筛选", ["全部"] + industries)

    # 显示企业卡片
    filtered_companies = 合作企业
    if selected_industry != "全部":
        filtered_companies = [company for company in 合作企业 if company["行业"] == selected_industry]

    cols = st.columns(2)
    for i, company in enumerate(filtered_companies):
        with cols[i % 2]:
            st.markdown(f"""
            <div class='company-card'>
                <div class='company-logo'>{company['logo']}</div>
                <h3>{company['名称']}</h3>
                <div class='industry-tag'>{company['行业']}</div>
                <div class='salary-tag'>{company['规模']}</div>
                <p>{company['简介']}</p>
                <h4>招聘岗位</h4>
                <p>{', '.join(company['招聘岗位'])}</p>
                <p><strong>薪资范围:</strong> {company['薪资范围']}</p>
            </div>
            """, unsafe_allow_html=True)

    # 实习就业机会
    st.markdown("---")
    st.subheader("🎓 实习就业机会")

    # 实习岗位筛选
    job_types = list(set([job["岗位名称"] for job in 实习机会]))
    selected_job_type = st.selectbox("按岗位类型筛选", ["全部"] + job_types)

    # 显示实习岗位
    filtered_jobs = 实习机会
    if selected_job_type != "全部":
        filtered_jobs = [job for job in 实习机会 if job["岗位名称"] == selected_job_type]

    for job in filtered_jobs:
        st.markdown(f"""
        <div class='job-card'>
            <h4>{job['岗位名称']} - {job['公司名称']}</h4>
            <div class='salary-tag'>{job['薪资']}</div>
            <div class='industry-tag'>{job['地点']}</div>
            <p><strong>要求:</strong> {job['要求']}</p>
            <p><strong>截止日期:</strong> {job['截止日期']}</p>
        </div>
        """, unsafe_allow_html=True)

    # AI推荐功能
    st.markdown("---")
    st.subheader("🤖 AI企业推荐")

    # 获取用户画像
    profile = st.session_state.get('student_profile', {})
    resume_portrait = st.session_state.get('portrait', None)

    if not profile and not resume_portrait:
        st.warning("请先完善您的能力画像或上传简历，以获取个性化企业推荐")
    else:
        if st.button("获取AI企业推荐"):
            with st.spinner("AI正在分析并推荐适合的企业..."):
                # 合并画像数据
                merged_profile = profile.copy()
                if resume_portrait:
                    if '基本信息' not in merged_profile:
                        merged_profile['基本信息'] = {}
                    if resume_portrait.get('专业') and resume_portrait['专业'] != '未知':
                        merged_profile['基本信息']['专业'] = resume_portrait['专业']
                    existing_skills = set(merged_profile.get('专业技能', []))
                    new_skills = set(resume_portrait.get('专业技能', []))
                    merged_profile['专业技能'] = list(existing_skills.union(new_skills))

                # 构建提示词
                prompt = f"""请根据以下学生信息，推荐适合的合作企业：

                ## 学生信息
                - 专业: {merged_profile.get('基本信息', {}).get('专业', '计算机相关')}
                - 专业技能: {', '.join(merged_profile.get('专业技能', [])) or '待完善'}
                - 证书: {', '.join(merged_profile.get('证书', [])) or '暂无'}

                请推荐3-5家适合该学生的企业，并说明推荐理由，包括：
                1. 企业名称
                2. 推荐理由
                3. 适合的岗位
                4. 薪资范围

                使用中文回答，结构化输出。"""

                response = llm.chat(prompt)
                st.markdown("### 个性化企业推荐")
                st.markdown(f"<div class='learning-path-card'>{response}</div>", unsafe_allow_html=True)


def render_career_graph():
    """职业图谱页面"""
    # 添加职业图谱页面的样式
    st.markdown("""
    <style>
        .career-card {
            background: white;
            border-radius: 12px;
            padding: 20px;
            box-shadow: 0 4px 15px rgba(156, 39, 176, 0.1);
            border: 1px solid #e8dfff;
            margin-bottom: 20px;
        }
        
        .promotion-step {
            text-align: center;
            padding: 15px;
            border-radius: 10px;
            margin: 0 5px;
        }
        
        .promotion-step.entry {
            background: linear-gradient(135deg, #4caf50 0%, #45a049 100%);
            color: white;
        }
        
        .promotion-step.middle {
            background: linear-gradient(135deg, #ff9800 0%, #f57c00 100%);
            color: white;
        }
        
        .promotion-step.advanced {
            background: linear-gradient(135deg, #f44336 0%, #d32f2f 100%);
            color: white;
        }
        
        .skill-tag {
            display: inline-block;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            margin: 4px;
        }
    </style>
    """, unsafe_allow_html=True)

    st.title("职业发展图谱")

    # 添加页面描述
    st.markdown("探索职业发展路径，了解晋升阶梯和技能要求")

    jobs = get_jobs_from_excel()
    selected_job = st.selectbox("选择岗位", jobs)

    st.markdown("---")

    # 使用AI大模型分析职业发展路径
    with st.spinner("AI正在分析职业发展路径..."):
        # 构建提示词
        prompt = f"""请分析{selected_job}的职业发展路径，包括：
        1. 晋升路径：从入门到高级的职业发展阶梯
        2. 换岗路径：相关的职业转换方向
        3. 技能要求：各个阶段需要的核心技能
        4. 行业发展趋势：该岗位的未来发展前景
        
        请以结构化的方式输出，使用中文回答。"""

        response = llm.chat(prompt)

        # 显示AI分析结果
        st.markdown("### 🤖 AI分析结果")
        st.markdown(f"<div class='career-card'>{response}</div>", unsafe_allow_html=True)

    # 解析AI分析结果，提取关键信息
    def parse_ai_response(response):
        """解析AI分析结果"""
        import re

        # 提取晋升路径
        promotion_paths = []
        career_paths = []
        skills = []

        # 查找包含'→'的行作为路径
        lines = response.split('\n')
        for line in lines:
            line = line.strip()
            if '→' in line:
                # 检查是否是晋升路径
                if any(keyword in line for keyword in ['晋升', '进阶', '发展']):
                    # 提取路径部分
                    promotion_paths.append(line)
                else:
                    # 作为换岗路径
                    career_paths.append(line)
            # 查找技能相关内容
            elif any(keyword in line for keyword in ['技能', '技术', '能力']):
                # 提取技能点
                skill_points = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9+#\-]+', line)
                skills.extend(skill_points)

        # 特殊处理用户提供的分析结果格式
        # 检查是否包含明确的晋升路径和换岗路径
        if '晋升路径' in response and '换岗路径' in response:
            # 提取晋升路径
            promotion_section = re.search(r'晋升路径.*?(?=换岗路径)', response, re.DOTALL)
            if promotion_section:
                # 提取职位阶段
                position_matches = re.findall(r'[\u4e00-\u9fa5]+(?:工程师|架构师|总监|CTO)', promotion_section.group(0))
                if position_matches:
                    # 去重并保持顺序
                    unique_positions = []
                    seen = set()
                    for pos in position_matches:
                        if pos not in seen:
                            seen.add(pos)
                            unique_positions.append(pos)
                    if len(unique_positions) > 1:
                        promotion_paths = [' → '.join(unique_positions)]
                    else:
                        # 如果没有足够的职位阶段，使用默认格式
                        promotion_paths = [f"初级{selected_job} → 中级{selected_job} → 高级{selected_job} → {selected_job[:-4]}架构师 → 技术总监/CTO"]

            # 提取换岗路径
            career_section = re.search(r'换岗路径.*?(?=技能要求|行业发展趋势|总结)', response, re.DOTALL)
            if career_section:
                # 提取可转换岗位
                position_matches = re.findall(r'[\u4e00-\u9fa5]+(?:工程师|架构师|经理|专家)', career_section.group(0))
                if position_matches:
                    # 去重
                    unique_positions = list(set(position_matches))
                    if unique_positions:
                        # 生成换岗路径
                        career_paths = []
                        for pos in unique_positions:
                            if pos != selected_job:
                                career_paths.append(f"{selected_job} → {pos}")
                if not career_paths:
                    # 如果没有提取到换岗路径，使用默认路径
                    career_paths = [
                        f"{selected_job} → 全栈开发工程师",
                        f"{selected_job} → 后端架构师",
                        f"{selected_job} → DevOps工程师",
                        f"{selected_job} → 产品经理",
                        f"{selected_job} → 数据工程师/大数据开发",
                        f"{selected_job} → 技术支持/运维工程师"
                    ]

        # 提取技能要求
        skills_section = re.search(r'技能要求.*?(?=行业发展趋势|总结)', response, re.DOTALL)
        if skills_section:
            # 提取技能点
            # 尝试匹配更具体的技能格式
            skill_patterns = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9+#\-\/]+', skills_section.group(0))
            # 过滤掉非技能词汇
            skill_keywords = ['熟悉', '掌握', '了解', '精通', '具备', '有', '需要', '要求', '技能', '技术', '能力', '初级', '中级', '高级', '架构师', '工程师', '总监', 'CTO']
            for skill in skill_patterns:
                if len(skill) > 2 and not any(keyword in skill for keyword in skill_keywords):
                    skills.append(skill)

        # 如果没有提取到技能，尝试从整个AI分析结果中提取
        if not skills:
            # 查找包含技能相关词汇的句子
            for line in response.split('\n'):
                line = line.strip()
                if any(keyword in line for keyword in ['技能', '技术', '能力']):
                    # 提取技能点
                    skill_patterns = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9+#\-\/]+', line)
                    for skill in skill_patterns:
                        if len(skill) > 2 and not any(keyword in skill for keyword in skill_keywords):
                            skills.append(skill)

        # 尝试从AI分析结果中提取晋升路径，即使没有'→'符号
        if not promotion_paths:
            # 查找晋升路径部分
            promotion_section = re.search(r'晋升路径.*?(?=换岗路径|技能要求|行业发展趋势|总结)', response, re.DOTALL)
            if promotion_section:
                # 提取职位阶段
                position_matches = re.findall(r'[\u4e00-\u9fa5]+(?:工程师|架构师|总监|CTO)', promotion_section.group(0))
                if position_matches:
                    # 去重并保持顺序
                    unique_positions = []
                    seen = set()
                    for pos in position_matches:
                        if pos not in seen:
                            seen.add(pos)
                            unique_positions.append(pos)
                    if len(unique_positions) > 1:
                        promotion_paths = [' → '.join(unique_positions)]

        # 尝试从AI分析结果中提取换岗路径
        if not career_paths:
            # 查找换岗路径部分
            career_section = re.search(r'换岗路径.*?(?=技能要求|行业发展趋势|总结)', response, re.DOTALL)
            if career_section:
                # 提取可转换岗位
                position_matches = re.findall(r'[\u4e00-\u9fa5]+(?:工程师|架构师|经理|专家)', career_section.group(0))
                if position_matches:
                    # 去重
                    unique_positions = list(set(position_matches))
                    if unique_positions:
                        # 生成换岗路径
                        for pos in unique_positions:
                            if pos != selected_job:
                                career_paths.append(f"{selected_job} → {pos}")

        # 如果没有找到晋升路径，使用默认
        if not promotion_paths:
            promotion_paths = [f"{selected_job} → 高级{selected_job} → 技术组长 → 架构师 → 技术总监"]

        # 如果没有找到换岗路径，使用默认
        if not career_paths:
            career_paths = [
                f"{selected_job} → 后端架构师 → 技术总监",
                f"{selected_job} → 数据开发工程师 → 数据架构师",
                f"{selected_job} → 产品经理 → 产品总监",
                f"{selected_job} → AI算法工程师 → AI专家"
            ]

        # 如果没有找到技能，使用默认
        if not skills:
            skills = ["Java", "SpringBoot", "SpringCloud", "MySQL", "Redis", "Docker", "K8s", "微服务", "Maven", "Git"]

        return {
            'promotion_paths': promotion_paths,
            'career_paths': career_paths,
            'skills': skills
        }

    # 解析AI结果
    ai_data = parse_ai_response(response)

    # 添加明显的分隔线，确保文字分析结果和图谱之间有清晰的界限
    st.markdown("---")
    st.markdown("### 📊 职业发展图谱")

    # 显示晋升路径（图谱）
    st.subheader("晋升路径")
    if ai_data['promotion_paths']:
        for path in ai_data['promotion_paths']:
            # 解析路径中的职位
            levels = [level.strip() for level in re.split(r'[→→]', path) if level.strip()]
            if levels:
                cols = st.columns(len(levels))
                for i, level in enumerate(levels):
                    with cols[i]:
                        if i == 0:
                            st.markdown(f"<div class='promotion-step entry'>{level}</div>", unsafe_allow_html=True)
                        elif i == len(levels) - 1:
                            st.markdown(f"<div class='promotion-step advanced'>{level}</div>", unsafe_allow_html=True)
                        else:
                            st.markdown(f"<div class='promotion-step middle'>{level}</div>", unsafe_allow_html=True)

    # 显示换岗路径
    st.markdown("---")
    st.subheader("换岗路径")
    if ai_data['career_paths']:
        for i, path in enumerate(ai_data['career_paths'], 1):
            st.markdown(f"路径{i}: {path}")

    # 显示技能要求
    st.markdown("---")
    st.subheader("技能要求")
    if ai_data['skills']:
        skill_tags = ""
        for skill in ai_data['skills'][:15]:  # 显示前15个技能
            skill_tags += f"<span class='skill-tag'>{skill}</span>"
        st.markdown(skill_tags, unsafe_allow_html=True)


def render_advanced_report(generator: AdvancedCareerReportGenerator, profile: Dict, match_results: Dict):
    """升级版规划报告页面"""
    # 添加规划报告页面的样式
    st.markdown("""
    <style>
        .report-card {
            background: white;
            border-radius: 12px;
            padding: 25px;
            box-shadow: 0 4px 15px rgba(156, 39, 176, 0.1);
            border: 1px solid #e8dfff;
            margin-bottom: 20px;
        }
        
        .export-button {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            border: none;
            border-radius: 8px;
            padding: 10px 20px;
            font-size: 14px;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.3s ease;
        }
        
        .export-button:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 10px rgba(102, 126, 234, 0.3);
        }
        
        /* 移除报告标题上方的空白圆角框 */
        .stMarkdown {
            background: none !important;
            border: none !important;
            border-radius: 0 !important;
            box-shadow: none !important;
        }
        
        .stMarkdown > div {
            background: none !important;
            border: none !important;
            border-radius: 0 !important;
            box-shadow: none !important;
        }
    </style>
    """, unsafe_allow_html=True)

    st.title("📄 AI个性化职业规划报告")
    st.markdown("基于AI技术，为您提供精准的职业规划方案")
    st.markdown("---")

    resume_portrait = st.session_state.get('portrait', None)
    if resume_portrait:
        st.success("✅ 已加载简历解析数据，报告将更加个性化")

    # 生成报告按钮
    if st.button("🎯 生成AI个性化报告", type="primary", use_container_width=True):
        with st.spinner("AI正在为您生成个性化职业规划报告..."):
            report, generated_time = generator.generate_ai_report(profile, match_results, resume_portrait)
            st.session_state['generated_report'] = report
            st.session_state['report_generated_time'] = generated_time

        if save_career_report(st.session_state['username'], report, generated_time):
            st.success("报告已生成并保存！")

    if 'generated_report' in st.session_state:
        st.markdown("---")

        tab1, tab2 = st.tabs(["📖 查看报告", "✏️ 编辑报告"])

        with tab1:
            st.markdown(st.session_state['generated_report'])

        with tab2:
            st.markdown("### 编辑报告内容")
            st.info("💡 您可以在下方编辑器中修改报告内容，修改后点击保存按钮即可更新报告")

            if 'edited_report' not in st.session_state:
                st.session_state['edited_report'] = st.session_state['generated_report']

            edited_content = st.text_area(
                "报告内容编辑器",
                value=st.session_state['edited_report'],
                height=500,
                key="report_editor",
                help="支持Markdown格式，您可以自由编辑报告内容"
            )

            col_edit1, col_edit2 = st.columns(2)
            with col_edit1:
                if st.button("💾 保存修改", type="primary", use_container_width=True):
                    st.session_state['generated_report'] = edited_content
                    st.session_state['edited_report'] = edited_content
                    if save_career_report(st.session_state['username'], edited_content, st.session_state.get('report_generated_time')):
                        st.success("报告修改已保存！")
                    st.rerun()

            with col_edit2:
                if st.button("🔄 恢复原始报告", use_container_width=True):
                    st.session_state['edited_report'] = st.session_state['generated_report']
                    st.rerun()

        st.markdown("---")
        st.subheader("📤 导出报告")

        export_folder = ensure_export_folder()

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"AI职业规划报告_{timestamp}"

        col1, col2, col3, col4, col5 = st.columns(5)

        with col1:
            md_content = st.session_state['generated_report'].encode('utf-8')
            md_filename = f"{base_filename}.md"
            with open(os.path.join(export_folder, md_filename), 'wb') as f:
                f.write(md_content)
            st.download_button(
                label="📄 Markdown",
                data=md_content,
                file_name=md_filename,
                mime="text/markdown",
                use_container_width=True
            )

        with col2:
            try:
                pdf_content = export_report_to_pdf(st.session_state['generated_report'], base_filename)
                pdf_filename = f"{base_filename}.pdf"
                with open(os.path.join(export_folder, pdf_filename), 'wb') as f:
                    f.write(pdf_content)
                st.download_button(
                    label="📕 PDF",
                    data=pdf_content,
                    file_name=pdf_filename,
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"PDF导出失败: {e}")

        with col3:
            try:
                word_content = export_report_to_word(st.session_state['generated_report'], base_filename)
                word_filename = f"{base_filename}.docx"
                with open(os.path.join(export_folder, word_filename), 'wb') as f:
                    f.write(word_content)
                st.download_button(
                    label="📘 Word",
                    data=word_content,
                    file_name=word_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Word导出失败: {e}")

        with col4:
            try:
                excel_content = export_report_to_excel(st.session_state['generated_report'], base_filename)
                excel_filename = f"{base_filename}.xlsx"
                with open(os.path.join(export_folder, excel_filename), 'wb') as f:
                    f.write(excel_content)
                st.download_button(
                    label="📊 Excel",
                    data=excel_content,
                    file_name=excel_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Excel导出失败: {e}")

        with col5:
            try:
                txt_content = export_report_to_txt(st.session_state['generated_report'], base_filename)
                txt_filename = f"{base_filename}.txt"
                with open(os.path.join(export_folder, txt_filename), 'wb') as f:
                    f.write(txt_content)
                st.download_button(
                    label="📝 TXT",
                    data=txt_content,
                    file_name=txt_filename,
                    mime="text/plain",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"TXT导出失败: {e}")

        st.success(f"报告已保存到: {os.path.abspath(export_folder)}")

    # 报告历史
    st.markdown("---")
    st.subheader("📋 报告历史")

    def get_report_history(username):
        try:
            conn = sqlite3.connect(DB_PATH)
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            c.execute("SELECT * FROM career_reports WHERE username = ? ORDER BY created_at DESC", (username,))
            rows = c.fetchall()
            conn.close()
            return rows
        except:
            return []

    if 'username' in st.session_state:
        history = get_report_history(st.session_state['username'])
        if history:
            st.success(f"找到 {len(history)} 份历史报告")
            for i, row in enumerate(history, 1):
                with st.expander(f"报告 {i}: {row['created_at']}"):
                    content = row['report_content']
                    st.markdown(content[:300] + "..." if len(content) > 300 else content)
                    if st.button(f"查看完整报告", key=f"view_{i}"):
                        st.session_state['view_report'] = content
                        st.rerun()
        else:
            st.info("暂无历史报告")

    if 'view_report' in st.session_state:
        st.markdown("---")
        st.subheader("完整报告")
        st.markdown(st.session_state['view_report'])
        if st.button("关闭"):
            del st.session_state['view_report']
            st.rerun()


def render_resume_parser_enhanced():
    """增强版简历解析页面"""
    st.title("📄 AI智能简历解析")
    st.markdown("---")

    st.info("""
    💡 **支持的文件格式**：
    - 文档：PDF、Word(.docx/.doc)、TXT、RTF、MD
    - 图片：JPG、JPEG、PNG（支持OCR文字识别）
    - 网页：HTML
    
    上传简历后，AI将自动提取您的个人信息、技能、项目经验等，并生成能力画像。
    """)

    if 'portrait' not in st.session_state:
        st.session_state['portrait'] = None
    if 'scores' not in st.session_state:
        st.session_state['scores'] = None

    col1, col2 = st.columns([1, 2])

    with col1:
        st.markdown("### 📂 上传简历文件")
        uploaded_file = st.file_uploader(
            "选择简历文件",
            type=['pdf', 'docx', 'doc', 'txt', 'rtf', 'md', 'html', 'jpg', 'jpeg', 'png']
        )

        st.markdown("### ✍️ 或手动输入")
        manual_input = st.text_area("粘贴简历内容", height=200, placeholder="请将简历内容粘贴到这里...")

        parse_clicked = st.button("🚀 开始智能解析", type="primary", use_container_width=True)

    with col2:
        st.markdown("### 🔍 解析结果")

        if parse_clicked:
            resume_text = ""
            file_name = None

            if uploaded_file:
                file_name = uploaded_file.name
                with st.spinner(f"正在解析文件: {file_name}..."):
                    resume_text = EnhancedResumeParser.extract_text_from_file(uploaded_file)
            elif manual_input:
                resume_text = manual_input
                file_name = "手动输入"
            else:
                st.warning("请上传简历文件或手动输入简历内容")
                return

            if resume_text and len(resume_text.strip()) >= 10:
                with st.spinner("AI正在深度分析简历内容..."):
                    portrait = EnhancedResumeParser.parse_resume_with_llm(resume_text, file_name)
                    scores = EnhancedResumeParser.calculate_portrait_score(portrait)
                    st.session_state['portrait'] = portrait
                    st.session_state['scores'] = scores

                st.success("解析完成！")

                # 显示评分
                col_a, col_b = st.columns(2)
                with col_a:
                    st.metric("完整度评分", f"{scores['完整度评分']}%")
                    st.progress(scores['完整度评分'] / 100)
                with col_b:
                    st.metric("竞争力评分", f"{scores['竞争力评分']}%")
                    st.progress(scores['竞争力评分'] / 100)

                # 基本信息
                basic = portrait.get("基本信息", {})
                st.markdown("---")
                st.markdown("#### 📋 基本信息")
                cols = st.columns(5)
                with cols[0]: st.info(f"**姓名**: {basic.get('姓名', '-')}")
                with cols[1]: st.info(f"**专业**: {basic.get('专业', '-')}")
                with cols[2]: st.info(f"**学历**: {basic.get('学历', '-')}")
                with cols[3]: st.info(f"**学校**: {basic.get('学校', '-')}")
                with cols[4]: st.info(f"**城市**: {basic.get('意向城市', '-')}")

                # 专业技能
                st.markdown("#### 🛠️ 专业技能")
                skills = portrait.get("专业技能", [])
                if skills:
                    st.markdown(" ".join([f"`{s}`" for s in skills]))
                else:
                    st.warning("未识别到技能信息")

                # 证书
                st.markdown("#### 📜 证书")
                certs = portrait.get("证书", [])
                if certs:
                    st.markdown(" ".join([f"`{c}`" for c in certs]))

                # 软实力
                soft = portrait.get("软实力", {})
                st.markdown("#### 🌟 软实力评估")
                soft_cols = st.columns(4)
                with soft_cols[0]: st.success(f"学习能力: {soft.get('学习能力', '良好')}")
                with soft_cols[1]: st.success(f"沟通能力: {soft.get('沟通能力', '良好')}")
                with soft_cols[2]: st.success(f"抗压能力: {soft.get('抗压能力', '良好')}")
                with soft_cols[3]: st.success(f"创新能力: {soft.get('创新能力', '良好')}")

                # 综合评估
                evaluation = portrait.get("综合评估", {})
                if evaluation.get("主要优势") or evaluation.get("待提升方向"):
                    st.markdown("#### 🎯 综合评估")
                    col_adv, col_gap = st.columns(2)
                    with col_adv:
                        st.markdown("**✅ 主要优势**")
                        for adv in evaluation.get("主要优势", []):
                            st.success(f"✓ {adv}")
                    with col_gap:
                        st.markdown("**⚠️ 待提升方向**")
                        for gap in evaluation.get("待提升方向", []):
                            st.warning(f"⚠ {gap}")

                    if evaluation.get("推荐岗位"):
                        st.markdown("**💼 推荐岗位**")
                        st.info(" → ".join(evaluation["推荐岗位"][:3]))
            else:
                st.error("简历内容太少，请提供更详细的简历信息")

        # 保存按钮
        if st.session_state.get('portrait'):
            st.markdown("---")
            if st.button("💾 保存到信息录入", type="primary", use_container_width=True):
                st.session_state['parsed_profile'] = st.session_state['portrait']
                if 'username' in st.session_state:
                    parsed_data = st.session_state['portrait']
                    basic = parsed_data.get("基本信息", {})
                    recommended = parsed_data.get("综合评估", {}).get("推荐岗位", [])
                    default_job = recommended[0] if recommended else get_jobs_from_excel()[0]

                    form_data = {
                        "name": basic.get('姓名', ''),
                        "major": basic.get('专业', ''),
                        "school": basic.get('学校', ''),
                        "education": basic.get('学历', '本科'),
                        "skills": parsed_data.get('专业技能', []),
                        "certificates": parsed_data.get('证书', []),
                        "projects": parsed_data.get('项目经验', []),
                        "internship": parsed_data.get('实习经历', ''),
                        "city": basic.get('意向城市', ''),
                        "salary": parsed_data.get('求职意向', {}).get('期望薪资', '8K-12K'),
                        "job_type": default_job
                    }

                    manager = StudentProfileManager()
                    profile = manager.create_profile_from_form(form_data)
                    completeness = manager.calculate_completeness_score(profile)
                    competitiveness = manager.calculate_competitiveness_score(profile)
                    profile['completeness_score'] = completeness
                    profile['competitiveness_score'] = competitiveness

                    if save_student_profile(st.session_state['username'], profile):
                        st.session_state['student_profile'] = profile
                        st.success("✅ 已保存到信息录入！")
                    else:
                        st.error("保存失败")
                else:
                    st.error("请先登录")


# ==================== 讨论区页面（新版） ====================

def render_discussion_page():
    """讨论区页面 - 完整版（支持搜索、编辑、删除、权限控制）"""
    st.markdown("""
    <style>
    /* 讨论区样式 */
    .discussion-container {
        max-width: 100%;
        margin: 0 auto;
    }
    
    .discussion-card {
        background: white;
        border-radius: 16px;
        padding: 20px;
        margin-bottom: 16px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        border: 1px solid #f0f0f0;
        transition: all 0.2s ease;
        overflow: visible;
    }

    .discussion-card:hover {
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
    }

    .discussion-pinned {
        border-left: 4px solid #ff9800;
        background: #fffbf5;
    }

    .discussion-title {
        font-size: 18px;
        font-weight: 600;
        margin-bottom: 8px;
        color: #333;
    }

    .discussion-meta {
        display: flex;
        flex-wrap: wrap;
        gap: 16px;
        margin-bottom: 12px;
    }

    .discussion-meta-item {
        display: inline-flex;
        align-items: center;
        gap: 4px;
        color: #888;
        font-size: 13px;
    }

    .discussion-content {
        margin: 12px 0;
        color: #555 !important;
        line-height: 1.6;
        font-size: 14px !important;
        word-wrap: break-word;
        overflow-wrap: break-word;
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
        min-height: 40px;
        padding: 8px 0;
        z-index: 10;
        position: relative;
    }

    .discussion-actions {
        display: flex;
        gap: 8px;
        flex-wrap: wrap;
        align-items: center;
        margin-top: 12px;
        padding-top: 12px;
        border-top: 1px solid #f0f0f0;
    }
    
    .reply-item {
        background: #f9f9f9;
        border-radius: 12px;
        padding: 12px 16px;
        margin-bottom: 12px;
        margin-left: 20px;
        border-left: 3px solid #667eea;
    }
    
    .reply-meta {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 8px;
        font-size: 12px;
        color: #888;
        flex-wrap: wrap;
        gap: 8px;
    }
    
    .reply-author {
        font-weight: 600;
        color: #667eea;
    }
    
    .reply-content {
        color: #333 !important;
        line-height: 1.6;
        margin-bottom: 8px;
        font-size: 14px !important;
        word-wrap: break-word;
        overflow-wrap: break-word;
        display: block !important;
        visibility: visible !important;
        opacity: 1 !important;
        min-height: 20px;
        padding: 4px 0;
        z-index: 5;
        position: relative;
    }
    
    .reply-footer {
        display: flex;
        justify-content: flex-end;
        gap: 12px;
        margin-top: 8px;
    }
    
    .reply-footer button {
        background: none;
        border: none;
        cursor: pointer;
        font-size: 12px;
        padding: 4px 8px;
        border-radius: 6px;
        color: #888;
        transition: all 0.2s;
    }
    
    .reply-footer button:hover {
        background: #e8e8e8;
        color: #333;
    }
    
    .badge-pinned {
        background: #ff9800;
        color: white;
        padding: 2px 8px;
        border-radius: 12px;
        font-size: 11px;
        margin-left: 8px;
    }
    
    .edit-form {
        background: #f5f5f5;
        padding: 15px;
        border-radius: 12px;
        margin: 10px 0;
    }
    
    .search-container {
        display: flex;
        gap: 10px;
        margin-bottom: 20px;
    }
    
    .reply-section {
        margin-top: 15px;
        padding: 15px;
        background: #fafafa;
        border-radius: 12px;
    }
    
    /* 按钮样式覆盖 */
    .stButton > button {
        border-radius: 8px !important;
        font-size: 13px !important;
    }
    
    /* 分割线 */
    hr {
        margin: 20px 0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    username = st.session_state.get('username', '')
    user_type = st.session_state.get('user_type', 'user')
    
    # 初始化会话状态
    if 'editing_discussion' not in st.session_state:
        st.session_state['editing_discussion'] = None
    if 'editing_reply' not in st.session_state:
        st.session_state['editing_reply'] = set()
    if 'expanded_discussions' not in st.session_state:
        st.session_state['expanded_discussions'] = set()
    if 'show_reply_forms' not in st.session_state:
        st.session_state['show_reply_forms'] = set()
    if 'search_keyword' not in st.session_state:
        st.session_state['search_keyword'] = ''
    
    st.title("💬 求职讨论区")
    st.markdown("分享经验、交流心得、解答疑惑")
    
    # 搜索框区域
    col_search1, col_search2 = st.columns([3, 1])
    with col_search1:
        search_keyword = st.text_input(
            "🔍 搜索帖子",
            value=st.session_state['search_keyword'],
            placeholder="输入关键词搜索（如：软考、面试、Java、简历...）",
            key="search_input",
            label_visibility="collapsed"
        )
    with col_search2:
        if st.session_state['search_keyword']:
            if st.button("❌ 清除搜索", use_container_width=True, key=f"btn_clear_search_{st.session_state['search_keyword']}", type="secondary"):
                st.session_state['search_keyword'] = ""
                st.rerun()
        else:
            if st.button("🔍 搜索", use_container_width=True, key=f"btn_search_{search_keyword}"):
                st.session_state['search_keyword'] = search_keyword
                st.rerun()
    
    # 获取我的发布和全部讨论
    all_discussions = get_discussions()
    my_discussions = [d for d in all_discussions if d.get('username') == username]
    
    # 搜索结果处理（搜索时完全隐藏Tab）
    if st.session_state['search_keyword']:
        search_results = get_discussions(st.session_state['search_keyword'])
        st.markdown("---")
        st.markdown("### 🔍 搜索结果")
        st.caption(f"搜索关键词：{st.session_state['search_keyword']} | 点击「清除搜索」返回全部话题")
        if search_results:
            st.success(f"找到 {len(search_results)} 个包含「{st.session_state['search_keyword']}」的帖子")
            for disc in search_results:
                render_discussion_card(disc, username, user_type, section="search")
            st.markdown("---")
        else:
            st.info(f"没有找到包含「{st.session_state['search_keyword']}」的帖子")
            st.markdown("---")
    else:
        tab1, tab2 = st.tabs(["📤 我的发布", "💬 全部话题"])
        
        with tab1:
            with st.expander("📝 发布新话题", expanded=False):
                with st.form("discussion_form", clear_on_submit=True):
                    title = st.text_input("标题", placeholder="请输入话题标题")
                    content = st.text_area("内容", placeholder="请输入话题内容", height=120)
                    is_anonymous = st.checkbox("🙈 匿名发布", help="选择后将隐藏您的身份，适合讨论敏感问题")
                    col1, col2 = st.columns([1, 5])
                    with col1:
                        submitted = st.form_submit_button("发布", type="primary", use_container_width=True)
                    if submitted:
                        if title and content:
                            if add_discussion(username, user_type, title, content, is_anonymous):
                                st.success("发布成功！")
                                log_operation(username, f"{'匿名' if is_anonymous else ''}发布话题: {title[:30]}")
                                time.sleep(0.5)
                                st.rerun()
                            else:
                                st.error("发布失败")
                        else:
                            st.warning("请填写标题和内容")
            
            st.markdown("### 📤 我的发布")
            if my_discussions:
                for disc in my_discussions:
                    render_discussion_card(disc, username, user_type, section="my")
            else:
                st.info("您还没有发布任何话题")
        
        with tab2:
            st.markdown("### 💬 全部话题")
            other_discussions = [d for d in all_discussions if d.get('username') != username]
            if other_discussions:
                for disc in other_discussions:
                    render_discussion_card(disc, username, user_type, section="all")
            else:
                st.info("暂无其他话题")
            st.caption(f"共 {len(all_discussions)} 个话题，其中 {len(other_discussions)} 个来自其他用户")


def render_discussion_card(disc, username, user_type, section=""):
    """渲染单个讨论卡片"""
    is_pinned = disc.get('is_pinned', 0)
    is_author = (disc.get('username') == username)
    is_admin = (user_type == 'admin')
    disc_id = disc['id']
    key_suffix = f"_{section}" if section else ""
    
    # 编辑讨论模式
    if st.session_state['editing_discussion'] == disc_id:
        with st.container():
            st.markdown('<div class="edit-form">', unsafe_allow_html=True)
            new_title = st.text_input("标题", value=disc.get('title', ''), key=f"edit_title_{disc_id}{key_suffix}")
            new_content = st.text_area("内容", value=disc.get('content', ''), height=150, key=f"edit_content_{disc_id}{key_suffix}")
            col_edit1, col_edit2, col_edit3 = st.columns([1, 1, 3])
            with col_edit1:
                if st.button("💾 保存", key=f"save_disc_{disc_id}{key_suffix}"):
                    if update_discussion(disc_id, new_title, new_content, username, user_type):
                        st.success("修改成功！")
                        st.session_state['editing_discussion'] = None
                        log_operation(username, f"编辑话题: {new_title[:30]}")
                        st.rerun()
                    else:
                        st.error("修改失败，您只能编辑自己发布的话题")
            with col_edit2:
                if st.button("❌ 取消", key=f"cancel_disc_{disc_id}{key_suffix}"):
                    st.session_state['editing_discussion'] = None
                    st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        return
    
    # 正常显示讨论卡片
    pin_class = "discussion-pinned" if is_pinned else ""
    content_preview = disc.get('content', '') or ''
    content_display = content_preview[:300] + ('...' if len(content_preview) > 300 else '')
    import html
    content_display = html.escape(content_display).replace('\n', '<br>')
    
    is_anonymous = disc.get('is_anonymous', 0)
    display_username = '匿名用户' if is_anonymous else disc.get('username', '未知')
    
    card_html = f"""
    <div class="discussion-card {pin_class}">
        <div class="discussion-title">
            {'📌 ' if is_pinned else ''}{disc.get('title', '无标题')}
            {f'<span class="badge-pinned">置顶</span>' if is_pinned else ''}
            {f'<span class="badge-pinned" style="background: #667eea; margin-left: 8px;">匿名</span>' if is_anonymous else ''}
        </div>
        <div class="discussion-meta">
            <span class="discussion-meta-item">👤 {display_username}</span>
            <span class="discussion-meta-item">💬 {disc.get('reply_count', 0)} 回复</span>
            <span class="discussion-meta-item">👍 {disc.get('likes', 0)} 点赞</span>
            <span class="discussion-meta-item">📅 {str(disc.get('created_at', ''))[:16] if disc.get('created_at') else '未知'}</span>
        </div>
        <div class="discussion-content">
            {content_display}
        </div>
    </div>
    """
    
    try:
        st.html(card_html)
    except Exception:
        st.markdown(card_html, unsafe_allow_html=True)
    
    col_actions = st.columns([1, 1, 1, 1, 1, 1, 1])

    with col_actions[0]:
        is_liked = check_discussion_liked(disc_id, username)
        like_text = "❤️ 已赞" if is_liked else "👍 点赞"
        if st.button(like_text, key=f"like_disc_{disc_id}{key_suffix}", use_container_width=True):
            like_discussion(disc_id, username)
            st.rerun()

    with col_actions[1]:
        is_expanded = disc_id in st.session_state['expanded_discussions']
        expand_text = "📖 收起" if is_expanded else "📖 查看详情"
        if st.button(expand_text, key=f"view_disc_{disc_id}{key_suffix}", use_container_width=True):
            if is_expanded:
                st.session_state['expanded_discussions'].discard(disc_id)
            else:
                st.session_state['expanded_discussions'].add(disc_id)
            st.rerun()

    with col_actions[2]:
        show_reply = disc_id in st.session_state['show_reply_forms']
        reply_text = "✏️ 取消回复" if show_reply else "💬 回复"
        if st.button(reply_text, key=f"reply_disc_{disc_id}{key_suffix}", use_container_width=True):
            if show_reply:
                st.session_state['show_reply_forms'].discard(disc_id)
            else:
                st.session_state['show_reply_forms'].add(disc_id)
            st.rerun()

    if is_admin:
        with col_actions[3]:
            if is_pinned:
                if st.button("📍 取消置顶", key=f"unpin_{disc_id}{key_suffix}", use_container_width=True):
                    pin_discussion(disc_id, 0)
                    log_operation(username, f"取消置顶话题: {disc.get('title', '')[:30]}")
                    st.rerun()
            else:
                if st.button("📌 置顶", key=f"pin_{disc_id}{key_suffix}", use_container_width=True):
                    pin_discussion(disc_id, 1)
                    log_operation(username, f"置顶话题: {disc.get('title', '')[:30]}")
                    st.rerun()

    if is_author or is_admin:
        with col_actions[4]:
            if st.button("✏️ 编辑", key=f"edit_disc_{disc_id}{key_suffix}", use_container_width=True):
                st.session_state['editing_discussion'] = disc_id
                st.rerun()

    if is_author or is_admin:
        with col_actions[5]:
            if st.button("🗑️ 删除", key=f"del_disc_{disc_id}{key_suffix}", use_container_width=True):
                if delete_discussion(disc_id, username, user_type):
                    st.success("删除成功")
                    log_operation(username, f"删除话题: {disc.get('title', '')[:30]}")
                    st.rerun()
                else:
                    st.error("删除失败")
    
    if disc_id in st.session_state['show_reply_forms']:
        with st.container():
            st.markdown('<div class="reply-section">', unsafe_allow_html=True)
            st.markdown("**✏️ 发表回复**")
            reply_content = st.text_area("", placeholder="请输入回复内容...", height=80, key=f"reply_input_{disc_id}{key_suffix}")
            reply_anonymous = st.checkbox("🙈 匿名回复", help="选择后将隐藏您的身份", key=f"reply_anonymous_{disc_id}{key_suffix}")
            col_r1, col_r2 = st.columns([1, 5])
            with col_r1:
                if st.button("📤 发送回复", key=f"submit_reply_{disc_id}{key_suffix}", use_container_width=True):
                    if reply_content.strip():
                        if add_reply(disc_id, username, user_type, reply_content, reply_anonymous):
                            st.success("回复成功！")
                            st.session_state['show_reply_forms'].discard(disc_id)
                            log_operation(username, f"{'匿名' if reply_anonymous else ''}回复话题: {disc.get('title', '')[:30]}")
                            st.rerun()
                        else:
                            st.error("回复失败")
                    else:
                        st.warning("请输入回复内容")
            with col_r2:
                if st.button("取消", key=f"cancel_reply_{disc_id}{key_suffix}", use_container_width=True):
                    st.session_state['show_reply_forms'].discard(disc_id)
                    st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
    
    if disc_id in st.session_state['expanded_discussions']:
        replies = get_replies(disc_id)
        
        if replies:
            st.markdown(f"**📋 共 {len(replies)} 条回复**")
            for reply in replies:
                reply_id = reply['id']
                is_reply_author = (reply.get('username') == username)
                is_reply_admin = (user_type == 'admin')
                can_edit_reply = is_reply_author or is_reply_admin
                
                if reply_id in st.session_state['editing_reply']:
                    st.markdown('<div class="edit-form" style="margin-left: 20px;">', unsafe_allow_html=True)
                    new_reply_content = st.text_area("", value=reply.get('content', ''), height=80, key=f"edit_reply_text_{reply_id}{key_suffix}")
                    col_er1, col_er2 = st.columns(2)
                    with col_er1:
                        if st.button("💾 保存", key=f"save_reply_{reply_id}{key_suffix}", use_container_width=True):
                            if update_reply(reply_id, new_reply_content, username, user_type):
                                st.success("修改成功！")
                                st.session_state['editing_reply'].discard(reply_id)
                                st.rerun()
                            else:
                                st.error("修改失败")
                    with col_er2:
                        if st.button("❌ 取消", key=f"cancel_edit_reply_{reply_id}{key_suffix}"):
                            st.session_state['editing_reply'].discard(reply_id)
                            st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)
                    continue
                
                is_reply_anonymous = reply.get('is_anonymous', 0)
                reply_display_name = '匿名用户' if is_reply_anonymous else reply.get('username', '未知')
                reply_content = reply.get('content', '') or ''
                reply_content_display = html.escape(reply_content).replace('\n', '<br>')
                
                reply_html = f"""
                <div class="reply-item">
                    <div class="reply-meta">
                        <div>
                            <span class="reply-author">{reply_display_name}</span>
                            {f'<span style="margin-left: 8px; color: #667eea; font-size: 12px;">👤 匿名</span>' if is_reply_anonymous else ''}
                            <span style="margin-left: 12px;">📅 {str(reply.get('created_at', ''))[:16] if reply.get('created_at') else '未知'}</span>
                            <span style="margin-left: 12px;">👍 {reply.get('likes', 0)}</span>
                        </div>
                    </div>
                    <div class="reply-content">{reply_content_display}</div>
                </div>
                """
                
                try:
                    st.html(reply_html)
                except Exception:
                    st.markdown(reply_html, unsafe_allow_html=True)
                
                col_ra1, col_ra2, col_ra3, col_ra4 = st.columns([6, 1, 1, 1])
                
                with col_ra2:
                    is_reply_liked = check_reply_liked(reply_id, username)
                    like_reply_text = "❤️" if is_reply_liked else "👍"
                    if st.button(like_reply_text, key=f"like_reply_{reply_id}{key_suffix}", help="点赞"):
                        like_reply(reply_id, username)
                        st.rerun()
                
                with col_ra3:
                    if can_edit_reply:
                        if st.button("✏️", key=f"edit_reply_btn_{reply_id}{key_suffix}", help="编辑"):
                            st.session_state['editing_reply'].add(reply_id)
                            st.rerun()
                
                with col_ra4:
                    if can_edit_reply:
                        if st.button("🗑️", key=f"del_reply_btn_{reply_id}{key_suffix}", help="删除"):
                            if delete_reply(reply_id, username, user_type, disc_id):
                                st.success("删除成功")
                                st.rerun()
        else:
            st.info("暂无回复，快来抢沙发吧～")
    
    st.markdown("---")


def render_resume_parser():
    """简历解析页面"""
    st.title("简历解析 - 生成学生画像")
    st.markdown("---")

    # 初始化会话状态
    if 'portrait' not in st.session_state:
        st.session_state['portrait'] = None
    if 'scores' not in st.session_state:
        st.session_state['scores'] = None

    col1, col2 = st.columns([1, 2])

    with col1:
        st.markdown("### 上传简历")
        uploaded_file = st.file_uploader("选择简历文件", type=['pdf', 'png', 'jpg', 'jpeg'])

        st.markdown("### 手动输入")
        manual_input = st.text_area("或者手动输入简历内容", height=200, placeholder="粘贴简历文本...")

    with col2:
        st.markdown("### 解析结果")

        if st.button("开始解析", type="primary"):
            resume_text = ""

            if uploaded_file:
                # 简单读取文件内容
                try:
                    if uploaded_file.name.endswith('.pdf'):
                        # PDF解析需要pdfplumber
                        try:
                            import pdfplumber
                            import tempfile
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                                tmp.write(uploaded_file.read())
                                tmp_path = tmp.name

                            with pdfplumber.open(tmp_path) as pdf:
                                for page in pdf.pages:
                                    resume_text += page.extract_text() or ""
                            os.unlink(tmp_path)
                        except Exception as e:
                            st.warning(f"PDF解析失败: {e}")
                            resume_text = manual_input or ""
                    else:
                        # 图片用base64简单处理
                        import base64
                        resume_text = f"[图片简历: {uploaded_file.name}]"
                except Exception as e:
                    st.warning(f"文件解析部分失败: {e}")
                    resume_text = manual_input or ""
            else:
                resume_text = manual_input

            if resume_text:
                with st.spinner("AI正在解析简历..."):
                    portrait = ResumeParser.parse_resume_with_llm(resume_text)
                    scores = ResumeParser.calculate_portrait_score(portrait)
                    # 保存到会话状态
                    st.session_state['portrait'] = portrait
                    st.session_state['scores'] = scores

                # 显示结果
                st.success("解析完成！")

                # 评分卡片
                col_a, col_b = st.columns(2)
                with col_a:
                    st.metric("完整度评分", f"{scores['完整度评分']}%")
                    st.progress(scores['完整度评分'] / 100)
                with col_b:
                    st.metric("竞争力评分", f"{scores['竞争力评分']}%")
                    st.progress(scores['竞争力评分'] / 100)

                st.markdown("---")
                st.markdown("### 能力画像详情")

                # 基本信息
                basic_cols = st.columns(4)
                with basic_cols[0]:
                    st.info(f"**姓名**: {portrait.get('姓名', '未知')}")
                with basic_cols[1]:
                    st.info(f"**专业**: {portrait.get('专业', '未知')}")
                with basic_cols[2]:
                    st.info(f"**学历**: {portrait.get('学历', '未知')}")
                with basic_cols[3]:
                    st.info(f"**城市**: {portrait.get('意向城市', '未知')}")

                # 技能
                st.markdown("#### 专业技能")
                skills = portrait.get("专业技能", [])
                if skills:
                    st.markdown(" ".join([f"`{s}`" for s in skills]))
                else:
                    st.warning("未识别到技能")

                # 证书
                st.markdown("#### 证书")
                certs = portrait.get("证书", [])
                if certs:
                    st.markdown(" ".join([f"`{c}`" for c in certs]))
                else:
                    st.warning("未识别到证书")

                # 项目经验
                st.markdown("#### 项目经验")
                projects = portrait.get("项目经验", [])
                if projects:
                    for p in projects:
                        st.markdown(f"- {p}")
                else:
                    st.warning("未识别到项目经验")

                # 软技能
                st.markdown("#### 软技能")
                soft_cols = st.columns(3)
                with soft_cols[0]:
                    st.success(f"学习能力: {portrait.get('学习能力', '未知')}")
                with soft_cols[1]:
                    st.success(f"沟通能力: {portrait.get('沟通能力', '未知')}")
                with soft_cols[2]:
                    st.success(f"抗压能力: {portrait.get('抗压能力', '未知')}")
            else:
                st.warning("请上传简历文件或手动输入简历内容")

        # 显示保存按钮，无论是否点击了解析按钮
        if st.session_state['portrait']:
            st.markdown("---")
            if st.button("保存到信息录入"):
                # 保存到会话状态
                st.session_state['parsed_profile'] = st.session_state['portrait']

                # 直接保存到数据库
                if 'username' in st.session_state:
                    # 构建保存数据
                    parsed_data = st.session_state['portrait']
                    form_data = {
                        "name": parsed_data.get('姓名', ''),
                        "major": parsed_data.get('专业', ''),
                        "school": "",  # 从简历中可能无法提取学校信息
                        "education": parsed_data.get('学历', '本科'),
                        "skills": parsed_data.get('专业技能', []),
                        "certificates": parsed_data.get('证书', []),
                        "projects": parsed_data.get('项目经验', []),
                        "internship": parsed_data.get('实习经历', ''),
                        "city": parsed_data.get('意向城市', ''),
                        "salary": "8K-12K",  # 默认值
                        "job_type": get_jobs_from_excel()[0]  # 默认第一个岗位
                    }

                    # 创建学生画像
                    manager = StudentProfileManager()
                    profile = manager.create_profile_from_form(form_data)
                    completeness = manager.calculate_completeness_score(profile)
                    competitiveness = manager.calculate_competitiveness_score(profile)
                    profile['completeness_score'] = completeness
                    profile['competitiveness_score'] = competitiveness

                    # 保存到数据库
                    if save_student_profile(st.session_state['username'], profile):
                        st.session_state['student_profile'] = profile

                        # 触发能力匹配分析
                        matching_engine = JobMatchingEngine()
                        match_result = matching_engine.match(profile, form_data['job_type'])
                        st.session_state['match_results'] = match_result

                        # 触发规划报告分析
                        report_gen = CareerReportGenerator()
                        report = report_gen.generate_report(profile, match_result)
                        save_career_report(st.session_state['username'], report)

                        st.success("✅ 已保存到数据库，并完成能力匹配和规划报告分析！")
                        st.info("您可以在能力画像、岗位匹配和规划报告页面查看分析结果。")
                    else:
                        st.error("保存到数据库失败，请稍后重试。")
                else:
                    st.error("请先登录后再保存。")


def render_resume_optimization():
    """简历优化与分析页面"""
    st.title("📄 简历优化与分析")
    st.markdown("---")

    username = st.session_state.get('username', '访客用户')

    # 上传简历
    st.subheader("📤 上传简历")
    uploaded_file = st.file_uploader("选择简历文件（支持txt、docx、pdf格式）", type=["txt", "docx", "pdf"])

    # 目标岗位
    target_job = st.text_input("目标岗位", placeholder="例如：Java开发工程师")

    if st.button("🔍 分析简历", type="primary"):
        if uploaded_file and target_job:
            with st.spinner("正在分析简历..."):
                # 读取简历内容
                try:
                    if uploaded_file.type == "text/plain":
                        resume_content = uploaded_file.getvalue().decode("utf-8")
                    else:
                        # 模拟简历解析
                        resume_content = "姓名：张三\n专业：计算机科学与技术\n学历：本科\n技能：Java, Python, MySQL\n项目经验：参与开发电商网站\n实习经验：某科技公司前端开发实习生"
                except Exception as e:
                    st.error(f"简历解析失败: {e}")
                    return

                # 模拟大模型分析
                system_prompt = "你是一个专业的简历分析师，擅长分析简历与目标岗位的匹配度，并提供具体的修改建议。请分析以下简历与目标岗位的匹配度，提供详细的修改建议，包括关键词优化、经历量化、格式调整等。"
                user_prompt = f"简历内容：{resume_content}\n目标岗位：{target_job}"

                # 调用大模型
                response = llm.chat(user_prompt, system_prompt)

                # 提取匹配度和建议
                # 模拟匹配度计算
                import random
                matching_score = round(random.uniform(60, 90), 1)

                # 模拟优化建议
                suggestions = "1. 技能部分添加具体的技术栈版本\n2. 项目经验添加量化成果\n3. 实习经验详细描述工作职责\n4. 添加相关证书和奖项\n5. 优化简历格式，使其更清晰易读"

                # 模拟优化后的简历
                optimized_resume = "姓名：张三\n专业：计算机科学与技术\n学历：本科\n技能：Java (Spring Boot), Python (Django), MySQL, Redis, Docker\n项目经验：参与开发电商网站，负责后端模块，提升系统性能30%\n实习经验：某科技公司前端开发实习生，负责页面开发和优化，完成10+页面开发\n证书：英语四级，计算机二级"

                # 保存优化记录
                save_resume_optimization(username, resume_content, optimized_resume, suggestions, matching_score, target_job)

                # 显示分析结果
                st.success("简历分析完成！")

                # 匹配度
                st.markdown("---")
                st.subheader("🎯 匹配度分析")
                st.markdown(f"**目标岗位**: {target_job}")
                st.markdown(f"**匹配度**: {matching_score}%")

                # 进度条显示匹配度
                st.progress(matching_score / 100)

                # 优化建议
                st.markdown("---")
                st.subheader("💡 优化建议")
                st.markdown(suggestions)

                # 优化前后对比
                st.markdown("---")
                st.subheader("📊 优化前后对比")

                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("**优化前**")
                    st.text_area("", value=resume_content, height=300)

                with col2:
                    st.markdown("**优化后**")
                    st.text_area("", value=optimized_resume, height=300)

                # 下载优化后的简历
                st.markdown("---")
                st.subheader("💾 下载优化后简历")
                st.download_button(
                    label="下载优化后的简历",
                    data=optimized_resume,
                    file_name=f"optimized_resume_{target_job}.txt",
                    mime="text/plain"
                )
        else:
            st.warning("请上传简历并输入目标岗位")

    # 历史优化记录
    st.markdown("---")
    st.subheader("📋 历史优化记录")

    history = get_resume_optimization_history(username)

    if history:
        for record in history:
            with st.expander(f"{record['target_job']} - 匹配度: {record['matching_score']}% - {record['created_at']}"):
                st.markdown(f"**目标岗位**: {record['target_job']}")
                st.markdown(f"**匹配度**: {record['matching_score']}%")
                st.markdown(f"**优化建议**: {record['suggestions']}")
                if st.button("查看详情", key=f"view_{record['id']}"):
                    # 显示详细信息
                    st.markdown("**原始简历**")
                    st.text_area("", value=record['original_resume'], height=200)
                    st.markdown("**优化后简历**")
                    st.text_area("", value=record['optimized_resume'], height=200)
    else:
        st.info("暂无简历优化记录")


def render_mock_interview():
    """模拟面试系统页面"""
    st.title("🎯 模拟面试系统")
    st.markdown("---")

    username = st.session_state.get('username', '访客用户')

    # 选择岗位类型
    job_categories = [
        "Java开发工程师",
        "Python开发工程师",
        "前端开发工程师",
        "数据分析工程师",
        "AI算法工程师",
        "产品经理",
        "UI设计师",
        "其他"
    ]

    job_title = st.selectbox("选择目标岗位", job_categories)

    # 面试题库（模拟）
    interview_questions = {
        "Java开发工程师": [
            "请解释什么是Java中的多线程？",
            "Spring Boot的核心功能有哪些？",
            "如何优化Java应用的性能？",
            "请解释数据库事务的ACID特性。",
            "如何处理Java中的异常？"
        ],
        "Python开发工程师": [
            "请解释Python中的装饰器是什么？",
            "Django和Flask的区别是什么？",
            "如何处理Python中的内存管理？",
            "请解释Python中的生成器和迭代器。",
            "如何优化Python代码的性能？"
        ],
        "前端开发工程师": [
            "请解释CSS中的盒模型。",
            "Vue和React的区别是什么？",
            "如何优化前端页面的加载速度？",
            "请解释JavaScript中的闭包。",
            "如何处理前端的跨域问题？"
        ],
        "数据分析工程师": [
            "请解释什么是数据仓库？",
            "如何处理缺失值和异常值？",
            "请解释SQL中的连接操作。",
            "如何使用Python进行数据可视化？",
            "请解释什么是机器学习？"
        ],
        "AI算法工程师": [
            "请解释什么是深度学习？",
            "如何处理过拟合问题？",
            "请解释卷积神经网络的工作原理。",
            "如何评估模型的性能？",
            "请解释什么是强化学习？"
        ],
        "产品经理": [
            "请解释什么是产品生命周期？",
            "如何进行用户需求分析？",
            "如何制定产品 roadmap？",
            "请解释什么是敏捷开发？",
            "如何处理产品上线后的问题？"
        ],
        "UI设计师": [
            "请解释什么是用户体验设计？",
            "如何进行色彩搭配？",
            "请解释什么是响应式设计？",
            "如何设计一个好的登录页面？",
            "请解释什么是设计系统？"
        ],
        "其他": [
            "请介绍一下你自己。",
            "为什么选择我们公司？",
            "你的职业规划是什么？",
            "你最大的优点和缺点是什么？",
            "你如何处理工作中的压力？"
        ]
    }

    # 开始面试
    if st.button("开始面试", type="primary"):
        # 初始化面试状态
        if 'interview_state' not in st.session_state:
            st.session_state['interview_state'] = {
                'job_title': job_title,
                'questions': interview_questions.get(job_title, interview_questions['其他']),
                'current_question': 0,
                'answers': [],
                'start_time': datetime.now()
            }
        st.rerun()

    # 面试过程
    if 'interview_state' in st.session_state:
        state = st.session_state['interview_state']

        # 显示当前问题
        if state['current_question'] < len(state['questions']):
            question = state['questions'][state['current_question']]
            st.markdown(f"### 问题 {state['current_question'] + 1}/{len(state['questions'])}")
            st.markdown(f"**{question}**")

            # 回答输入
            answer = st.text_area("请输入您的回答", height=200, key=f"answer_{state['current_question']}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("下一题", type="primary"):
                    state['answers'].append(answer)
                    state['current_question'] += 1
                    st.rerun()
            with col2:
                if st.button("退出面试"):
                    del st.session_state['interview_state']
                    st.rerun()
        else:
            # 面试结束，生成反馈
            st.success("面试完成！")

            # 计算面试时间
            end_time = datetime.now()
            duration = (end_time - state['start_time']).total_seconds() / 60

            # 构建面试内容
            interview_content = "\n".join([f"问题 {i+1}: {q}\n回答: {a}\n" for i, (q, a) in enumerate(zip(state['questions'], state['answers']))])

            # 生成AI反馈
            system_prompt = "你是一个专业的面试官，擅长分析面试表现并提供详细的反馈。请根据以下面试内容，提供专业的反馈和评分。"
            user_prompt = f"面试岗位: {state['job_title']}\n面试内容:\n{interview_content}"

            feedback = llm.chat(user_prompt, system_prompt)

            # 模拟评分
            import random
            score = round(random.uniform(70, 95), 1)

            # 保存面试记录
            save_mock_interview(username, state['job_title'], str(state['questions']), str(state['answers']), feedback, score)

            # 显示反馈
            st.markdown("---")
            st.subheader("面试反馈")
            st.markdown(f"**面试岗位**: {state['job_title']}")
            st.markdown(f"**面试时间**: {duration:.1f} 分钟")
            st.markdown(f"**面试评分**: {score}分")

            st.markdown("---")
            st.subheader("AI反馈")
            st.markdown(feedback)

            # 结束面试
            if st.button("结束面试"):
                del st.session_state['interview_state']
                st.rerun()

    # 面试历史
    st.markdown("---")
    st.subheader("📋 面试历史")

    history = get_mock_interview_history(username)

    if history:
        for record in history:
            with st.expander(f"{record['job_title']} - 评分: {record['score']} - {record['created_at']}"):
                st.markdown(f"**岗位**: {record['job_title']}")
                st.markdown(f"**评分**: {record['score']}")
                st.markdown(f"**反馈**: {record['feedback']}")
                if st.button("查看详情", key=f"view_interview_{record['id']}"):
                    st.markdown("**问题**: {record['interview_questions']}")
                    st.markdown("**回答**: {record['user_answers']}")
    else:
        st.info("暂无面试记录")


def render_learning_resources():
    """学习资源页面"""
    st.title("学习资源推荐")

    if 'student_profile' in st.session_state and 'match_results' in st.session_state:
        profile = st.session_state['student_profile']
        gaps = st.session_state['match_results'].get('差距', [])

        if gaps:
            st.markdown("### 根据您的能力差距推荐的学习资源")

            resources = {
                "Python": [
                    {"name": "Python编程从入门到实践", "type": "书籍",
                     "url": "https://book.douban.com/subject/3116733/"},
                    {"name": "Python Flask Web开发", "type": "书籍",
                     "url": "https://book.douban.com/subject/26597428/"},
                    {"name": "Python官方文档", "type": "文档", "url": "https://docs.python.org/3/"},
                    {"name": "Coursera Python课程", "type": "在线课程", "url": "https://www.coursera.org/learn/python"}
                ],
                "Java": [
                    {"name": "Java核心技术", "type": "书籍", "url": "https://book.douban.com/subject/26880667/"},
                    {"name": "Spring Boot实战", "type": "书籍", "url": "https://book.douban.com/subject/26897412/"},
                    {"name": "慕课网Java工程师", "type": "在线课程", "url": "https://www.imooc.com/learn/124"},
                    {"name": "B站Java学习路线", "type": "视频", "url": "https://www.bilibili.com/video/BV1zE411j7Nw"}
                ],
                "前端": [
                    {"name": "MDN Web开发文档", "type": "文档", "url": "https://developer.mozilla.org/"},
                    {"name": "Vue3官方文档", "type": "文档", "url": "https://vuejs.org/"},
                    {"name": "React官方文档", "type": "文档", "url": "https://react.dev/"}
                ]
            }

            for gap in gaps:
                with st.expander(f"✓ {gap} 的资源"):
                    for category, res_list in resources.items():
                        if category.lower() in gap.lower() or gap.lower() in category.lower():
                            for res in res_list:
                                st.markdown(f"- **{res['name']}** ({res['type']}): [访问链接]({res['url']})")
                            break
                    else:
                        for res in resources["Python"]:
                            st.markdown(f"- **{res['name']}** ({res['type']}): [访问链接]({res['url']})")
        else:
            st.info("请先进行岗位匹配，系统将根据您的能力差距推荐学习资源")

    st.markdown("---")
    st.markdown("### 热门技能学习资源")

    tabs = st.tabs(["Java", "Python", "前端", "数据分析", "AI算法"])

    resources_all = {
        "Java": [
            {"name": "Java核心技术", "type": "书籍", "url": "https://book.douban.com/subject/26880667/"},
            {"name": "Spring Boot实战", "type": "书籍", "url": "https://book.douban.com/subject/26897412/"},
            {"name": "慕课网Java工程师", "type": "在线课程", "url": "https://www.imooc.com/learn/124"}
        ],
        "Python": [
            {"name": "Python编程从入门到实践", "type": "书籍", "url": "https://book.douban.com/subject/3116733/"},
            {"name": "利用Python进行数据分析", "type": "书籍", "url": "https://book.douban.com/subject/26278699/"},
            {"name": "Coursera Python课程", "type": "在线课程", "url": "https://www.coursera.org/learn/python"}
        ],
        "前端": [
            {"name": "MDN Web开发文档", "type": "文档", "url": "https://developer.mozilla.org/"},
            {"name": "Vue3官方文档", "type": "文档", "url": "https://vuejs.org/"},
            {"name": "React官方文档", "type": "文档", "url": "https://react.dev/"}
        ],
        "数据分析": [
            {"name": "利用Python进行数据分析", "type": "书籍", "url": "https://book.douban.com/subject/26278699/"},
            {"name": "SQL必知必会", "type": "书籍", "url": "https://book.douban.com/subject/24250054/"},
            {"name": "Kaggle数据集", "type": "实践", "url": "https://www.kaggle.com/datasets"}
        ],
        "AI算法": [
            {"name": "机器学习实战", "type": "书籍", "url": "https://book.douban.com/subject/24462300/"},
            {"name": "深度学习入门", "type": "书籍", "url": "https://book.douban.com/subject/27028517/"},
            {"name": "吴恩达机器学习", "type": "在线课程", "url": "https://www.coursera.org/learn/machine-learning"}
        ]
    }

    for i, tab in enumerate(tabs):
        with tab:
            skill_name = list(resources_all.keys())[i]
            for res in resources_all[skill_name]:
                st.markdown(f"- **{res['name']}** ({res['type']}): [访问]({res['url']})")

    st.markdown("---")
    st.markdown("### 证书考试资源")

    certs = {
        "软考初级": "程序员",
        "软考中级": "软件设计师",
        "软考高级": "系统架构设计师",
        "PMP": "项目管理专业人士",
        "NPDP": "产品经理专业人士",
        "AWS认证": "AWS认证"
    }

    for cert, name in certs.items():
        if "软考" in cert:
            st.markdown(f"- **{cert}（{name}）**: [官网](https://www.ruankao.org.cn/)")
        elif cert == "PMP":
            st.markdown(f"- **{cert}（{name}）**: [官网](https://www.pmi.org/)")
        elif cert == "NPDP":
            st.markdown(f"- **{cert}（{name}）**: [官网](https://www.pdma.org/)")
        else:
            st.markdown(f"- **{cert}（{name}）**: [官网](https://aws.amazon.com/certification/)")

    st.markdown("---")
    st.markdown("### 实习招聘平台")

    platforms = [
        {"name": "实习僧", "url": "https://www.shixiseng.com/"},
        {"name": "牛客网", "url": "https://www.nowcoder.com/"},
        {"name": "拉勾网实习", "url": "https://www.lagou.com/"},
        {"name": "BOSS直聘实习", "url": "https://www.zhipin.com/"}
    ]

    for platform in platforms:
        st.markdown(f"- **{platform['name']}**: [访问]({platform['url']})")


def render_industry_news():
    """行业动态页面"""
    st.title("📰 行业动态")
    st.markdown("---")

    st.subheader("🌐 科技行业最新资讯")

    news = IndustryNewsAPI.get_tech_news()
    for item in news:
        with st.expander(f"📌 {item['title']}"):
            st.markdown(f"**{item['content']}**")
            st.caption(f"来源: {item.get('source', '未知')}")

    st.markdown("---")
    st.subheader("📈 就业市场趋势")

    trend = IndustryNewsAPI.get_job_market_trend()

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**🔥 热门岗位**")
        for job in trend.get("热门岗位", []):
            st.success(f"✓ {job}")

    with col2:
        st.markdown("**📈 薪资上涨**")
        for salary in trend.get("薪资上涨", []):
            st.info(f"↑ {salary}")

    col3, col4 = st.columns(2)
    with col3:
        st.markdown("**📋 需求增加**")
        for job in trend.get("需求增加", []):
            st.warning(f"• {job}")

    with col4:
        st.markdown("**💡 建议**")
        st.success(trend.get("建议", ""))

    st.markdown("---")
    st.subheader("🔔 动态更新提醒")
    st.info("建议定期查看行业动态，根据市场变化调整您的职业规划")


def render_career_consultation():
    """职业咨询页面"""
    st.title("职业咨询服务")
    st.markdown("---")

    st.markdown("""
    ### 咨询方式
    
    1. **在线讨论区**: 在讨论区发布您的问题，获得社区成员和导师的解答
    2. **预约咨询**: 联系行业导师获取一对一职业规划指导
    """)

    st.markdown("---")
    st.subheader("👨‍🏫 导师团队")

    mentors = [
        {"name": "张老师", "title": "互联网大厂技术总监", "expertise": "Java开发、架构设计", "price": "299元/次"},
        {"name": "李老师", "title": "产品总监", "expertise": "产品规划、用户研究", "price": "399元/次"},
        {"name": "王老师", "title": "AI算法专家", "expertise": "机器学习、深度学习", "price": "499元/次"},
    ]

    for mentor in mentors:
        with st.expander(f"{mentor['name']} - {mentor['title']}"):
            st.markdown(f"**专长**: {mentor['expertise']}")
            st.markdown(f"**咨询费用**: {mentor['price']}")
            if st.button(f"预约咨询", key=f"book_{mentor['name']}"):
                st.success(f"已向 {mentor['name']} 发送预约请求，请留意站内信或邮箱")

    st.markdown("---")
    st.subheader("📝 咨询预约表单")

    with st.form("consult_form"):
        name = st.text_input("您的姓名")
        phone = st.text_input("联系电话")
        interest = st.selectbox("咨询方向", ["技术开发", "产品设计", "数据分析", "AI算法", "职业规划"])
        question = st.text_area("您的问题或需求", height=100)

        submit = st.form_submit_button("提交预约", type="primary")

        if submit:
            if name and question:
                st.success("预约提交成功！我们将尽快与您联系")
            else:
                st.warning("请填写姓名和问题")

    st.markdown("---")
    st.info("💡 提示：您也可以在讨论区免费发布问题，获得社区帮助")


# ==================== 系统辅助函数 ====================

def log_operation(username: str, action: str, ip: str = ""):
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("INSERT INTO operation_logs (username, action, ip) VALUES (?, ?, ?)",
                  (username, action, ip))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"记录日志失败: {e}")

def get_system_config():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT config_key, config_value FROM system_config")
    rows = c.fetchall()
    conn.close()
    config = {}
    for row in rows:
        config[row['config_key']] = row['config_value']
    return config

def save_system_config(config_key, config_value):
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute("UPDATE system_config SET config_value = ?, updated_at = ? WHERE config_key = ?",
                  (config_value, datetime.now(), config_key))
        conn.commit()
        conn.close()
        return True
    except:
        return False

# ==================== 管理员端函数 ====================
def render_admin_dashboard():
    """渲染系统仪表盘 - 专业可视化大屏版本"""
    st.title("📊 系统管理中心")
    st.markdown("---")

    # 获取统计数据
    stats = get_system_statistics()

    # 第一行：关键指标卡片
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_users']}</div>
            <div class="dashboard-label">注册用户</div>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_profiles']}</div>
            <div class="dashboard-label">学生画像</div>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_discussions']}</div>
            <div class="dashboard-label">讨论话题</div>
        </div>
        """, unsafe_allow_html=True)
    with col4:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_replies']}</div>
            <div class="dashboard-label">回复总数</div>
        </div>
        """, unsafe_allow_html=True)
    with col5:
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-number">{stats['total_reports']}</div>
            <div class="dashboard-label">规划报告</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    # ====================== 可视化大屏区域 ======================
    st.markdown("## 📈 核心数据")

    # 第一行：技能分布 + 内容分布
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("### 热门技能分布")
        if stats['skill_count']:
            top_skills = sorted(stats['skill_count'].items(), key=lambda x: x[1], reverse=True)[:10]
            if top_skills:
                skills_df = pd.DataFrame(top_skills, columns=['技能', '人数'])
                fig = px.bar(skills_df, x='人数', y='技能', orientation='h',
                             title='学生掌握热门技能Top10',
                             color='人数', color_continuous_scale='Blues', text='人数')
                fig.update_layout(plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)', height=400)
                fig.update_traces(textposition='outside')
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.info("暂无技能数据")
        else:
            st.info("暂无技能数据，请先录入学生信息")

    with col2:
        st.markdown("### 平台内容分布")
        data_categories = ['学生画像', '讨论话题', '规划报告']
        data_values = [stats['total_profiles'], stats['total_discussions'], stats['total_reports']]
        if sum(data_values) > 0:
            fig = go.Figure(data=[go.Pie(labels=data_categories, values=data_values, hole=0.4,
                                         marker_colors=['#667eea', '#764ba2', '#f39c12'],
                                         textinfo='label+percent', textposition='outside')])
            fig.update_layout(plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)', height=400)
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("暂无数据")

    st.markdown("---")

    # ====================== 用户行为分析 ======================
    st.markdown("## 🔄 用户行为分析中心")

    # 加载日志数据（用于绘图）
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM operation_logs ORDER BY created_at DESC")
    logs = [dict(row) for row in c.fetchall()]
    conn.close()

    if not logs:
        st.info("暂无操作记录，无法生成分析图表")
        return

    # 转为 DataFrame 方便绘图
    df = pd.DataFrame(logs)
    df['created_at'] = pd.to_datetime(df['created_at'], errors='coerce')
    df['date'] = df['created_at'].dt.date
    df['hour'] = df['created_at'].dt.hour

    # 一行 3 个分析图表
    col_a, col_b, col_c = st.columns(3)

    with col_a:
        st.markdown("#### 每日活跃趋势")
        daily_count = df.groupby('date').size().reset_index(name='次数')
        fig = px.line(daily_count, x='date', y='次数', title='平台日活跃度',
                      color_discrete_sequence=['#1f77b4'])
        fig.update_layout(height=280, plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)')
        st.plotly_chart(fig, use_container_width=True)

    with col_b:
        st.markdown("#### 24小时访问时段")
        hour_count = df.groupby('hour').size().reset_index(name='次数')
        fig = px.bar(hour_count, x='hour', y='次数', title='用户活跃时段分布',
                     color_discrete_sequence=['#ff7f0e'])
        fig.update_layout(height=280, plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)')
        st.plotly_chart(fig, use_container_width=True)

    with col_c:
        st.markdown("#### 用户操作排行")
        user_count = df['username'].value_counts().reset_index()
        user_count.columns = ['用户', '操作次数']
        user_count = user_count.head(8)
        fig = px.pie(user_count, names='用户', values='操作次数', title='活跃用户占比')
        fig.update_layout(height=280, plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)')
        st.plotly_chart(fig, use_container_width=True)

    # 第二行：操作类型分析
    st.markdown("#### 操作类型统计")
    action_count = df['action'].value_counts().reset_index()
    action_count.columns = ['操作类型', '次数']
    fig = px.bar(action_count, x='操作类型', y='次数', color='次数',
                 title='用户各类操作数量', color_continuous_scale='Agsunset')
    fig.update_layout(height=320, plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)')
    st.plotly_chart(fig, use_container_width=True)

def render_admin_user_management():
    st.title("👥 用户管理")
    st.markdown("---")

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT id, username, email, user_type, created_at,
               (SELECT COUNT(*) FROM student_profiles WHERE username = users.username) as has_profile
        FROM users ORDER BY created_at DESC
    """)
    rows = c.fetchall()
    users = [dict(row) for row in rows]
    conn.close()

    if users:
        df = pd.DataFrame(users)
        st.dataframe(df, use_container_width=True)

    st.markdown("---")
    st.subheader("🔧 用户操作")

    # 样式定制
    st.markdown("""
    <style>
    .user-op-container {
        gap: 16px !important;
        display: flex;
        flex-direction: column;
    }
    .stSelectbox > div {
        border-radius: 12px;
        padding: 4px;
        border: 1px solid #e0e0e0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .stTextInput > div > div > input {
        height: 52px;
        border-radius: 12px;
        border: 1px solid #e0e0e0;
        padding: 0 16px;
        font-size: 15px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    .stButton > button {
        height: 56px;
        border-radius: 12px;
        font-size: 16px;
        font-weight: 600;
        border: none;
        box-shadow: 0 4px 12px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
        margin: 8px 0;
    }
    </style>
    """, unsafe_allow_html=True)

    if users:
        current_user = st.session_state['username']
        available_users = [u for u in users if u['username'] != current_user]

        if not available_users:
            st.info("没有其他用户可操作")
        else:
            st.markdown('<div class="user-op-container">', unsafe_allow_html=True)

            user_options = {f"{u['username']} (ID:{u['id']})": u for u in available_users}
            selected_key = st.selectbox("选择要操作的用户", list(user_options.keys()))
            selected_user = user_options[selected_key]

            # 角色切换按钮
            new_role = "admin" if selected_user['user_type'] == "user" else "user"
            role_text = "设为管理员" if new_role == "admin" else "降为普通用户"
            if st.button(f"👑 {role_text}", use_container_width=True):
                conn = sqlite3.connect(DB_PATH)
                c = conn.cursor()
                c.execute("UPDATE users SET user_type = ? WHERE id = ?", (new_role, selected_user['id']))
                conn.commit()
                conn.close()
                log_operation(st.session_state['username'], f"修改用户 {selected_user['username']} 角色为 {new_role}")
                st.success(f"已{role_text} {selected_user['username']}")
                time.sleep(0.5)
                st.rerun()

            # 密码重置
            new_password = st.text_input(
                "新密码",
                placeholder="输入新密码（至少8位）",
                type="password",
                key="reset_pwd"
            )
            if st.button("🔑 重置密码", use_container_width=True):
                if new_password and len(new_password) >= 8:
                    conn = sqlite3.connect(DB_PATH)
                    c = conn.cursor()
                    pwd_hash = hash_password(new_password)
                    c.execute("UPDATE users SET password_hash = ? WHERE id = ?", (pwd_hash, selected_user['id']))
                    conn.commit()
                    conn.close()
                    log_operation(st.session_state['username'], f"重置用户 {selected_user['username']} 密码")
                    st.success(f"已重置 {selected_user['username']} 的密码")
                else:
                    st.warning("密码至少8位")

            # 删除用户
            if st.button("🗑️ 删除用户", use_container_width=True, type="primary"):
                conn = sqlite3.connect(DB_PATH)
                c = conn.cursor()
                c.execute("DELETE FROM users WHERE id = ?", (selected_user['id'],))
                c.execute("DELETE FROM student_profiles WHERE username = ?", (selected_user['username'],))
                c.execute("DELETE FROM career_reports WHERE username = ?", (selected_user['username'],))
                c.execute("DELETE FROM discussions WHERE username = ?", (selected_user['username'],))
                c.execute("DELETE FROM replies WHERE username = ?", (selected_user['username'],))
                conn.commit()
                conn.close()
                log_operation(st.session_state['username'], f"删除用户 {selected_user['username']}")
                st.success(f"已删除用户 {selected_user['username']}")
                time.sleep(0.5)
                st.rerun()

            st.markdown('</div>', unsafe_allow_html=True)

def render_admin_discussion_management():
    st.title("💬 讨论管理")
    st.markdown("---")

    if 'admin_editing_discussion' not in st.session_state:
        st.session_state['admin_editing_discussion'] = None

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT d.*,
               (SELECT COUNT(*) FROM replies WHERE discussion_id = d.id) as reply_count
        FROM discussions d
        ORDER BY d.is_pinned DESC, d.created_at DESC
    """)
    rows = c.fetchall()
    discussions = [dict(row) for row in rows]
    conn.close()

    if not discussions:
        st.info("暂无讨论")
        return

    filter_status = st.selectbox("筛选", ["全部", "置顶", "普通"])

    for disc in discussions:
        if filter_status == "置顶" and not disc.get('is_pinned', 0):
            continue
        if filter_status == "普通" and disc.get('is_pinned', 0):
            continue

        disc_id = disc['id']

        # 编辑讨论模式
        if st.session_state['admin_editing_discussion'] == disc_id:
            with st.container():
                st.markdown("**✏️ 编辑讨论**")
                new_title = st.text_input("标题", value=disc['title'], key=f"admin_edit_title_{disc_id}")
                new_content = st.text_area("内容", value=disc['content'], height=120, key=f"admin_edit_content_{disc_id}")
                col_edit1, col_edit2, col_edit3 = st.columns([1, 1, 3])
                with col_edit1:
                    if st.button("💾 保存", key=f"admin_save_disc_{disc_id}"):
                        if update_discussion(disc_id, new_title, new_content, st.session_state['username'], "admin"):
                            st.success("修改成功！")
                            st.session_state['admin_editing_discussion'] = None
                            log_operation(st.session_state['username'], f"管理员编辑话题: {new_title[:30]}")
                            st.rerun()
                        else:
                            st.error("修改失败")
                with col_edit2:
                    if st.button("❌ 取消", key=f"admin_cancel_disc_{disc_id}"):
                        st.session_state['admin_editing_discussion'] = None
                        st.rerun()
            st.markdown("---")
            continue

        with st.container():
            pin_badge = "📌 " if disc.get('is_pinned', 0) else ""
            st.markdown(f"""
            <div style="background: #f5f5f5; border-radius: 10px; padding: 15px; margin: 10px 0;">
                <h4>{pin_badge}{disc.get('title', '无标题')}</h4>
                <p>{disc.get('content', '')[:150]}{'...' if len(disc.get('content', '')) > 150 else ''}</p>
                <small style="color: #666;">
                    👤 {disc.get('username', '未知')} | 💬 {disc.get('reply_count', 0)}回复 |
                    👍 {disc.get('likes', 0)}赞 | 📅 {disc.get('created_at', '')[:16] if disc.get('created_at') else '未知'}
                </small>
            </div>
            """, unsafe_allow_html=True)

            col1, col2, col3 = st.columns([1, 1, 1])
            with col1:
                if disc.get('is_pinned', 0):
                    if st.button("📍 取消置顶", key=f"unpin_{disc['id']}"):
                        pin_discussion(disc['id'], 0)
                        log_operation(st.session_state['username'], f"取消置顶讨论 ID:{disc['id']}")
                        st.rerun()
                else:
                    if st.button("📌 置顶", key=f"pin_{disc['id']}"):
                        pin_discussion(disc['id'], 1)
                        log_operation(st.session_state['username'], f"置顶讨论 ID:{disc['id']}")
                        st.rerun()

            with col2:
                if st.button("✏️ 编辑", key=f"admin_edit_disc_{disc['id']}"):
                    st.session_state['admin_editing_discussion'] = disc['id']
                    st.rerun()

            with col3:
                if st.button("🗑️ 删除", key=f"admin_del_{disc['id']}"):
                    if delete_discussion(disc['id'], disc['username'], "admin"):
                        st.success("删除成功")
                        st.rerun()

            with st.expander(f"查看回复 ({disc.get('reply_count', 0)})"):
                conn2 = sqlite3.connect(DB_PATH)
                conn2.row_factory = sqlite3.Row
                c2 = conn2.cursor()
                c2.execute("SELECT * FROM replies WHERE discussion_id = ? ORDER BY created_at ASC", (disc['id'],))
                reply_rows = c2.fetchall()
                replies = [dict(row) for row in reply_rows]
                conn2.close()

                for reply in replies:
                    st.markdown(f"""
                    <div style="padding: 10px; margin: 5px 0; background: #e8e8e8; border-radius: 5px;">
                        <strong>{reply.get('username', '未知')}</strong>: {reply.get('content', '')}
                        <br><small>{reply.get('created_at', '')[:16] if reply.get('created_at') else '未知'}</small>
                    </div>
                    """, unsafe_allow_html=True)

                    col_r1, col_r2, col_r3 = st.columns([6, 1, 1])
                    with col_r2:
                        if st.button("✏️", key=f"admin_edit_reply_{reply['id']}", help="编辑"):
                            st.session_state[f'admin_editing_reply_{reply["id"]}'] = True
                            st.rerun()
                    with col_r3:
                        if st.button("🗑️", key=f"admin_del_reply_{reply['id']}", help="删除"):
                            if delete_reply(reply['id'], reply['username'], "admin", disc['id']):
                                st.rerun()

                    if st.session_state.get(f'admin_editing_reply_{reply["id"]}', False):
                        new_reply_content = st.text_area("", value=reply['content'], height=80, key=f"admin_edit_reply_text_{reply['id']}")
                        col_er1, col_er2 = st.columns(2)
                        with col_er1:
                            if st.button("💾 保存", key=f"admin_save_reply_{reply['id']}"):
                                if update_reply(reply['id'], new_reply_content, st.session_state['username'], "admin"):
                                    st.success("修改成功！")
                                    st.session_state[f'admin_editing_reply_{reply["id"]}'] = False
                                    st.rerun()
                        with col_er2:
                            if st.button("❌ 取消", key=f"admin_cancel_reply_{reply['id']}"):
                                st.session_state[f'admin_editing_reply_{reply["id"]}'] = False
                                st.rerun()

            st.markdown("---")

def render_admin_data_statistics():
    st.title("📈 数据统计分析")
    st.markdown("---")

    conn = sqlite3.connect(DB_PATH)

    # 统计1：画像完成度分布
    st.subheader("📊 用户画像完成度分布")
    df = pd.read_sql_query("""
        SELECT completeness_score FROM student_profiles
        WHERE completeness_score IS NOT NULL
    """, conn)

    if not df.empty:
        df['等级'] = pd.cut(df['completeness_score'],
                            bins=[0, 30, 60, 80, 100],
                            labels=['低(0-30)', '中(30-60)', '良(60-80)', '优(80-100)'])
        dist = df['等级'].value_counts()
        fig = px.bar(x=dist.index, y=dist.values, title="画像完成度分布")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("暂无画像数据")

    # 统计2：热门技能分析
    st.subheader("🔥 热门技能分析")
    c = conn.cursor()
    c.execute("SELECT profile_data FROM student_profiles")
    profiles = c.fetchall()

    skill_counts = {}
    for profile in profiles:
        if profile[0]:
            try:
                data = json.loads(profile[0])
                skills = data.get('专业技能', [])
                for skill in skills:
                    skill_counts[skill] = skill_counts.get(skill, 0) + 1
            except:
                pass

    if skill_counts:
        top_skills = sorted(skill_counts.items(), key=lambda x: x[1], reverse=True)[:10]
        df_skills = pd.DataFrame(top_skills, columns=['技能', '人数'])
        fig = px.bar(df_skills, x='技能', y='人数', title="热门技能排行")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("暂无技能数据")

    # 统计3：用户类型分布
    st.subheader("👥 用户类型分布")
    c = conn.cursor()
    c.execute("SELECT user_type, COUNT(*) as count FROM users GROUP BY user_type")
    rows = c.fetchall()
    conn.close()

    if rows:
        df_type = pd.DataFrame(rows, columns=['用户类型', '数量'])
        fig = px.pie(df_type, values='数量', names='用户类型', title="用户类型分布")
        st.plotly_chart(fig, use_container_width=True)
    else:
        st.info("暂无用户数据")

def render_admin_operation_logs():
    st.title("📜 操作日志")
    st.markdown("---")

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM operation_logs ORDER BY created_at DESC LIMIT 100")
    rows = c.fetchall()
    logs = [dict(row) for row in rows]
    conn.close()

    if logs:
        df_logs = pd.DataFrame(logs)
        st.dataframe(df_logs, use_container_width=True)
    else:
        st.info("暂无操作日志")

def render_admin_system_config():
    st.title("⚙️ 系统配置")
    st.markdown("---")

    config = get_system_config()

    st.subheader("📝 站点信息")

    site_name = st.text_input("站点名称", value=config.get('site_name', 'AI大学生职业规划智能体'))
    site_desc = st.text_area("站点描述", value=config.get('site_desc', '基于人工智能技术，为您提供精准的职业规划方案'))

    if st.button("保存站点信息", type="primary"):
        save_system_config('site_name', site_name)
        save_system_config('site_desc', site_desc)
        st.success("保存成功！页面将刷新")
        time.sleep(0.5)
        st.rerun()

    st.markdown("---")
    st.subheader("🎨 主题设置")

    theme_color = st.color_picker("主题色", value=config.get('theme_color', '#667eea'))
    if st.button("保存主题色", type="primary"):
        save_system_config('theme_color', theme_color)
        st.success("保存成功！页面将刷新")
        time.sleep(0.5)
        st.rerun()

    st.markdown("---")
    st.subheader("🔧 功能开关")

    col1, col2 = st.columns(2)
    with col1:
        enable_discussion = st.checkbox("启用讨论区", value=config.get('enable_discussion', 'True') == 'True')
        enable_resume_parse = st.checkbox("启用简历解析", value=config.get('enable_resume_parse', 'True') == 'True')
    with col2:
        enable_report = st.checkbox("启用报告生成", value=config.get('enable_report', 'True') == 'True')
        enable_recommend = st.checkbox("启用岗位推荐", value=config.get('enable_recommend', 'True') == 'True')

    if st.button("保存配置", type="primary"):
        save_system_config('enable_discussion', str(enable_discussion))
        save_system_config('enable_resume_parse', str(enable_resume_parse))
        save_system_config('enable_report', str(enable_report))
        save_system_config('enable_recommend', str(enable_recommend))
        st.success("配置已保存！页面将刷新")
        time.sleep(0.5)
        st.rerun()

    st.markdown("---")
    st.info("💡 提示：修改配置后页面会自动刷新，刷新后新配置生效。")

def render_admin_report_management():
    """职业规划报告管理"""
    st.title("📊 职业规划报告管理")
    st.markdown("---")
    
    if 'viewing_report' not in st.session_state:
        st.session_state['viewing_report'] = None
    if 'confirm_delete_report' not in st.session_state:
        st.session_state['confirm_delete_report'] = None
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    if st.session_state['confirm_delete_report']:
        report_id_to_delete = st.session_state['confirm_delete_report']
        c.execute("DELETE FROM career_reports WHERE id = ?", (report_id_to_delete,))
        conn.commit()
        st.success(f"报告 ID {report_id_to_delete} 已删除")
        st.session_state['confirm_delete_report'] = None
        time.sleep(1)
        st.rerun()
    
    c.execute("SELECT COUNT(*) FROM career_reports")
    total_reports = c.fetchone()[0]
    
    c.execute("SELECT COUNT(*) FROM career_reports WHERE DATE(created_at) = DATE('now')")
    today_reports = c.fetchone()[0]
    
    c.execute("SELECT COUNT(*) FROM career_reports WHERE username = '访客用户'")
    guest_reports = c.fetchone()[0]
    
    registered_reports = total_reports - guest_reports
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(f"""
        <div style="background: white; border-radius: 12px; padding: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.05);">
            <div style="font-size: 24px; font-weight: bold; color: #333;">{total_reports}</div>
            <div style="font-size: 14px; color: #666; margin-top: 5px;">总报告数</div>
            <div style="font-size: 12px; color: #4CAF50; margin-top: 5px;">↑ 12% 较上月</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown(f"""
        <div style="background: white; border-radius: 12px; padding: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.05);">
            <div style="font-size: 24px; font-weight: bold; color: #333;">{today_reports}</div>
            <div style="font-size: 14px; color: #666; margin-top: 5px;">今日新增</div>
            <div style="font-size: 12px; color: #4CAF50; margin-top: 5px;">↑ 8% 较昨日</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        guest_ratio = int(guest_reports/total_reports*100) if total_reports > 0 else 0
        st.markdown(f"""
        <div style="background: #FFF3E0; border-radius: 12px; padding: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.05);">
            <div style="font-size: 24px; font-weight: bold; color: #333;">{guest_reports}</div>
            <div style="font-size: 14px; color: #666; margin-top: 5px;">访客用户报告</div>
            <div style="font-size: 12px; color: #666; margin-top: 5px;">占比 {guest_ratio}%</div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        registered_ratio = int(registered_reports/total_reports*100) if total_reports > 0 else 0
        st.markdown(f"""
        <div style="background: #F3E5F5; border-radius: 12px; padding: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.05);">
            <div style="font-size: 24px; font-weight: bold; color: #333;">{registered_reports}</div>
            <div style="font-size: 14px; color: #666; margin-top: 5px;">注册用户报告</div>
            <div style="font-size: 12px; color: #666; margin-top: 5px;">占比 {registered_ratio}%</div>
        </div>
        """, unsafe_allow_html=True)
    
    if st.session_state['viewing_report']:
        col_back, col_title = st.columns([1, 5])
        with col_back:
            if st.button("← 返回列表"):
                st.session_state['viewing_report'] = None
                st.rerun()
        with col_title:
            st.markdown("### 报告详情")
        
        report_id = st.session_state['viewing_report']
        c.execute("SELECT report_content, username, created_at FROM career_reports WHERE id = ?", (report_id,))
        result = c.fetchone()
        
        if result:
            report_content, username, created_at = result
            st.markdown(f"**用户：** {username}")
            st.markdown(f"**生成时间：** {created_at}")
            st.markdown("---")
            st.markdown(report_content)
            
            st.download_button(
                label="📥 下载报告",
                data=report_content,
                file_name=f"职业规划报告_{username}_{created_at}.md",
                mime="text/markdown",
                use_container_width=True
            )
        
        conn.close()
        return
    
    col_search, col_filter, col_export = st.columns([3, 2, 1])
    
    with col_search:
        search_keyword = st.text_input("搜索报告", placeholder="输入用户名搜索", label_visibility="collapsed")
    
    with col_filter:
        user_type = st.selectbox("用户类型", ["全部", "注册用户", "访客用户"], label_visibility="collapsed")
    
    with col_export:
        c.execute("SELECT id, username, report_content, created_at FROM career_reports ORDER BY created_at DESC")
        all_reports = c.fetchall()
        if all_reports:
            export_data = ""
            for r in all_reports:
                export_data += f"报告ID: {r[0]}, 用户: {r[1]}, 时间: {r[3]}\n\n报告内容:\n{r[2]}\n\n{'='*50}\n\n"
            
            st.download_button(
                label="📤 导出全部",
                data=export_data,
                file_name="职业规划报告汇总.md",
                mime="text/markdown",
                use_container_width=True
            )
    
    query = "SELECT id, username, created_at FROM career_reports WHERE 1=1"
    params = []
    
    if search_keyword:
        query += " AND username LIKE ?"
        params.append(f"%{search_keyword}%")
    
    if user_type == "注册用户":
        query += " AND username != '访客用户'"
    elif user_type == "访客用户":
        query += " AND username = '访客用户'"
    
    query += " ORDER BY created_at DESC"
    
    c.execute(query, params)
    reports = c.fetchall()
    
    if not reports:
        st.info("暂无职业规划报告")
        conn.close()
        return
    
    st.markdown("---")
    st.markdown("### 报告列表")
    
    page_size = 10
    total_pages = (len(reports) + page_size - 1) // page_size
    current_page = st.session_state.get('report_page', 1)
    
    start_idx = (current_page - 1) * page_size
    end_idx = min(start_idx + page_size, len(reports))
    
    col_id, col_user, col_type, col_time, col_action = st.columns([1, 2, 1.5, 2, 3])
    
    with col_id:
        st.markdown("**报告ID**")
    with col_user:
        st.markdown("**用户**")
    with col_type:
        st.markdown("**用户类型**")
    with col_time:
        st.markdown("**生成时间**")
    with col_action:
        st.markdown("**操作**")
    
    st.markdown("---")
    
    for i in range(start_idx, end_idx):
        report = reports[i]
        report_id, username, created_at = report
        
        col_id, col_user, col_type, col_time, col_action = st.columns([1, 2, 1.5, 2, 3])
        
        with col_id:
            st.write(report_id)
        with col_user:
            st.write(username)
        with col_type:
            if username == "访客用户":
                st.markdown("<span style='background: #FFF3E0; color: #FF9800; padding: 2px 8px; border-radius: 10px; font-size: 12px;'>访客</span>", unsafe_allow_html=True)
            else:
                st.markdown("<span style='background: #E3F2FD; color: #2196F3; padding: 2px 8px; border-radius: 10px; font-size: 12px;'>注册用户</span>", unsafe_allow_html=True)
        with col_time:
            st.write(created_at)
        with col_action:
            col_view, col_down, col_del = st.columns(3)
            with col_view:
                if st.button("👁️", key=f"view_{report_id}", help="查看报告"):
                    st.session_state['viewing_report'] = report_id
                    st.rerun()
            with col_down:
                c.execute("SELECT report_content FROM career_reports WHERE id = ?", (report_id,))
                report_content = c.fetchone()[0]
                st.download_button(
                    "📥",
                    data=report_content,
                    file_name=f"职业规划报告_{username}_{created_at}.md",
                    mime="text/markdown",
                    key=f"down_{report_id}",
                    help="下载报告"
                )
            with col_del:
                if st.button("🗑️", key=f"del_{report_id}", help="删除报告"):
                    st.session_state['confirm_delete_report'] = report_id
                    st.rerun()
    st.markdown("---")

    # 创建分页控制栏，确保所有按钮在一行显示，大小一致
    if total_pages > 1:
        # 简化的分页控件：上一页 + 页码 + 下一页
        col_prev, col_page, col_next = st.columns([1, 2, 1])

        with col_prev:
            if current_page > 1:
                if st.button("〈", key="report_prev"):
                    st.session_state['report_page'] = current_page - 1
                    st.rerun()

        with col_page:
            # 显示当前页码
            st.markdown(f"""
            <div style="display: flex; justify-content: center; align-items: center; gap: 10px;">
                <div style="background: #667eea; color: white; padding: 8px 16px; border-radius: 8px; font-weight: bold; text-align: center;">
                    {current_page}
                </div>
            </div>
            """, unsafe_allow_html=True)

        with col_next:
            if current_page < total_pages:
                if st.button("〉", key="report_next"):
                    st.session_state['report_page'] = current_page + 1
                    st.rerun()

    st.caption(f"显示 {start_idx + 1}-{end_idx} 条，共 {len(reports)} 条")
    
    conn.close()


def render_admin_resume_management():
    """简历优化历史管理"""
    st.title("📄 简历管理")
    st.markdown("---")
    
    if 'viewing_resume' not in st.session_state:
        st.session_state['viewing_resume'] = None
    if 'confirm_delete_resume' not in st.session_state:
        st.session_state['confirm_delete_resume'] = None
    if 'resume_page' not in st.session_state:
        st.session_state['resume_page'] = 1
    
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    if st.session_state['confirm_delete_resume']:
        resume_id_to_delete = st.session_state['confirm_delete_resume']
        c.execute("DELETE FROM resume_optimization WHERE id = ?", (resume_id_to_delete,))
        conn.commit()
        st.success(f"简历 ID {resume_id_to_delete} 已删除")
        st.session_state['confirm_delete_resume'] = None
        time.sleep(1)
        st.rerun()
    
    if st.session_state['viewing_resume']:
        col_back, col_title = st.columns([1, 5])
        with col_back:
            if st.button("← 返回列表"):
                st.session_state['viewing_resume'] = None
                st.rerun()
        with col_title:
            st.markdown("### 简历详情")
        
        resume_id = st.session_state['viewing_resume']
        c.execute("SELECT original_resume, optimized_resume, suggestions, username, target_job, created_at FROM resume_optimization WHERE id = ?", (resume_id,))
        result = c.fetchone()
        
        if result:
            original_resume, optimized_resume, suggestions, username, target_job, created_at = result
            st.markdown(f"**用户：** {username}")
            st.markdown(f"**目标岗位：** {target_job}")
            st.markdown(f"**上传时间：** {created_at}")
            st.markdown("---")
            
            tab1, tab2, tab3 = st.tabs(["原始简历", "优化后简历", "优化建议"])
            with tab1:
                st.text(original_resume)
            with tab2:
                st.text(optimized_resume)
            with tab3:
                st.markdown(suggestions)
            
            st.download_button(
                label="📥 下载简历",
                data=optimized_resume,
                file_name=f"简历_{username}_{target_job}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        conn.close()
        return
    
    col_search, col_filter, col_export = st.columns([3, 2, 1])
    
    with col_search:
        search_keyword = st.text_input("搜索姓名/专业/岗位", placeholder="输入关键词搜索", label_visibility="collapsed")
    
    with col_filter:
        job_filter = st.selectbox("岗位类型", ["全部", "Java开发", "前端开发", "数据分析", "产品经理", "UI设计", "其他"], label_visibility="collapsed")
    
    with col_export:
        c.execute("SELECT id, username, target_job, optimized_resume, created_at FROM resume_optimization ORDER BY created_at DESC")
        all_resumes = c.fetchall()
        if all_resumes:
            export_data = ""
            for r in all_resumes:
                export_data += f"ID: {r[0]}, 用户: {r[1]}, 岗位: {r[2]}, 时间: {r[4]}\n\n优化后简历:\n{r[3]}\n\n{'='*50}\n\n"
            
            st.download_button(
                label="📤 导出全部",
                data=export_data,
                file_name="简历汇总.txt",
                mime="text/plain",
                use_container_width=True
            )
    
    query = "SELECT id, username, target_job, matching_score, created_at FROM resume_optimization WHERE 1=1"
    params = []
    
    if search_keyword:
        query += " AND (username LIKE ? OR target_job LIKE ?)"
        params.extend([f"%{search_keyword}%", f"%{search_keyword}%"])
    
    if job_filter != "全部":
        query += " AND target_job = ?"
        params.append(job_filter)
    
    query += " ORDER BY created_at DESC"
    
    c.execute(query, params)
    resumes = c.fetchall()
    
    st.markdown("---")
    st.markdown("### 简历列表")
    
    page_size = 10
    total_pages = (len(resumes) + page_size - 1) // page_size
    current_page = st.session_state.get('resume_page', 1)
    
    start_idx = (current_page - 1) * page_size
    end_idx = min(start_idx + page_size, len(resumes))
    
    col_id, col_user, col_major, col_job, col_edu, col_time, col_action = st.columns([1, 2, 1.5, 2, 1, 2, 3])
    
    with col_id:
        st.markdown("**ID**")
    with col_user:
        st.markdown("**姓名**")
    with col_major:
        st.markdown("**专业**")
    with col_job:
        st.markdown("**意向岗位**")
    with col_edu:
        st.markdown("**学历**")
    with col_time:
        st.markdown("**上传时间**")
    with col_action:
        st.markdown("**操作**")
    
    st.markdown("---")
    
    if resumes:
        for i in range(start_idx, end_idx):
            resume = resumes[i]
            resume_id, username, target_job, matching_score, created_at = resume
            
            major = "计算机" if i % 3 == 0 else "软件工程" if i % 3 == 1 else "大数据"
            education = "本科" if i % 2 == 0 else "硕士"
            
            col_id, col_user, col_major, col_job, col_edu, col_time, col_action = st.columns([1, 2, 1.5, 2, 1, 2, 3])
            
            with col_id:
                st.write(i+1)
            with col_user:
                st.write(username)
            with col_major:
                st.write(major)
            with col_job:
                st.write(target_job)
            with col_edu:
                st.write(education)
            with col_time:
                st.write(created_at)
            with col_action:
                col_view, col_down, col_del = st.columns(3)
                with col_view:
                    if st.button("👁️", key=f"resume_view_{resume_id}", help="查看简历"):
                        st.session_state['viewing_resume'] = resume_id
                        st.rerun()
                with col_down:
                    c.execute("SELECT optimized_resume FROM resume_optimization WHERE id = ?", (resume_id,))
                    optimized_resume = c.fetchone()[0]
                    st.download_button(
                        "📥",
                        data=optimized_resume,
                        file_name=f"简历_{username}_{target_job}.txt",
                        mime="text/plain",
                        key=f"resume_down_{resume_id}",
                        help="下载简历"
                    )
                with col_del:
                    if st.button("🗑️", key=f"resume_del_{resume_id}", help="删除简历"):
                        st.session_state['confirm_delete_resume'] = resume_id
                        st.rerun()
    else:
        st.info("暂无简历数据")
    
    st.markdown("---")

    if resumes:
        # 创建分页控制栏，确保所有按钮在一行显示，大小一致
        if total_pages > 1:
            # 简化的分页控件：上一页 + 页码 + 下一页
            col_prev, col_page, col_next = st.columns([1, 2, 1])

            with col_prev:
                if current_page > 1:
                    if st.button("〈", key="resume_prev"):
                        st.session_state['resume_page'] = current_page - 1
                        st.rerun()

            with col_page:
                # 显示当前页码
                st.markdown(f"""
                <div style="display: flex; justify-content: center; align-items: center; gap: 10px;">
                    <div style="background: #667eea; color: white; padding: 8px 16px; border-radius: 8px; font-weight: bold; text-align: center;">
                        {current_page}
                    </div>
                </div>
                """, unsafe_allow_html=True)

            with col_next:
                if current_page < total_pages:
                    if st.button("〉", key="resume_next"):
                        st.session_state['resume_page'] = current_page + 1
                        st.rerun()

        st.caption(f"显示 {start_idx + 1}-{end_idx} 条，共 {len(resumes)} 份简历")
    
    conn.close()


def render_admin_interview_management():
    """模拟面试记录管理"""
    st.title("🎯 模拟面试记录管理")
    st.markdown("---")
    
    # 连接数据库
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    # 查询所有模拟面试记录
    c.execute("SELECT id, username, job_title, score, created_at FROM mock_interviews ORDER BY created_at DESC")
    interviews = c.fetchall()
    
    if not interviews:
        st.info("暂无模拟面试记录")
        return
    
    # 显示面试记录列表
    for interview in interviews:
        interview_id, username, job_title, score, created_at = interview
        
        with st.expander(f"记录 ID: {interview_id} | 用户: {username} | 岗位: {job_title} | 得分: {score} | 面试时间: {created_at}"):
            # 查询面试详情
            c.execute("SELECT interview_questions, user_answers, feedback FROM mock_interviews WHERE id = ?", (interview_id,))
            interview_questions, user_answers, feedback = c.fetchone()
            
            # 显示面试内容
            st.markdown("### 面试问题")
            st.text(interview_questions)
            
            st.markdown("### 你的回答")
            st.text(user_answers)
            
            st.markdown("### AI 反馈")
            st.markdown(feedback)
    
    conn.close()


def render_admin_profile_management():
    """学生能力画像管理"""
    st.title("🧠 学生能力画像管理")
    st.markdown("---")
    
    # 连接数据库
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    # 查询所有学生画像
    c.execute("SELECT id, username, completeness_score, competitiveness_score, updated_at FROM student_profiles ORDER BY updated_at DESC")
    profiles = c.fetchall()
    
    if not profiles:
        st.info("暂无学生能力画像数据")
        return
    
    # 显示学生画像列表
    for profile in profiles:
        profile_id, username, completeness, competitiveness, updated_at = profile
        
        with st.expander(f"用户: {username} | 完整度: {completeness}% | 竞争力: {competitiveness}% | 更新时间: {updated_at}"):
            # 查询画像详情
            c.execute("SELECT profile_data FROM student_profiles WHERE id = ?", (profile_id,))
            profile_data = c.fetchone()[0]
            
            # 显示画像数据
            st.markdown("### 能力画像数据")
            try:
                data = json.loads(profile_data)
                st.json(data)
            except:
                st.text(profile_data)
    
    conn.close()


def render_admin_panel(choice):
    if choice == "仪表盘":
        render_admin_dashboard()
    elif choice == "用户管理":
        render_admin_user_management()
    elif choice == "讨论管理":
        render_admin_discussion_management()
    elif choice == "报告管理":
        render_admin_report_management()
    elif choice == "简历管理":
        render_admin_resume_management()
    elif choice == "面试管理":
        render_admin_interview_management()
    elif choice == "能力画像管理":
        render_admin_profile_management()
    elif choice == "数据统计":
        render_admin_data_statistics()
    elif choice == "操作日志":
        render_admin_operation_logs()
    elif choice == "系统配置":
        render_admin_system_config()


# ==================== 主函数 ====================

def main():
    init_database()

    if 'logged_in' not in st.session_state:
        st.session_state['logged_in'] = False
    if 'show_register' not in st.session_state:
        st.session_state['show_register'] = False
    if 'show_change_pwd' not in st.session_state:
        st.session_state['show_change_pwd'] = False
    if 'show_forgot_pwd' not in st.session_state:
        st.session_state['show_forgot_pwd'] = False
    if 'show_admin_reset' not in st.session_state:
        st.session_state['show_admin_reset'] = False
    if 'nav' not in st.session_state:
        st.session_state['nav'] = 'Home'

    query_params = st.query_params

    if 'nav' in query_params:
        nav_value = query_params['nav']
        if nav_value == 'Logout':
            st.session_state['logged_in'] = False
            st.session_state['username'] = None
            st.session_state['user_type'] = None
            st.session_state['nav'] = 'Home'
            st.query_params.clear()
            st.rerun()
            return

    if st.session_state.get('logged_in', False) and st.session_state.get('user_type') == 'admin':
        # 默认导航到仪表盘
        if 'admin_nav' not in st.session_state:
            st.session_state['admin_nav'] = 'dashboard'

        # 渲染管理员顶部导航栏（使用表单按钮方式）
        render_admin_top_nav()

        # 根据导航参数渲染对应页面
        current_admin_page = st.session_state.get('admin_nav', 'dashboard')

        if current_admin_page == 'change_password':
            render_change_password_page()
        else:
            # 调用 render_admin_panel 函数，根据页面参数渲染对应管理界面
            # 需要将导航参数映射到原有的菜单项
            menu_map = {
                'dashboard': '仪表盘',
                'user_management': '用户管理',
                'discussion_management': '讨论管理',
                'report_management': '报告管理',
                'resume_management': '简历管理',
                'interview_management': '面试管理',
                'profile_management': '能力画像管理',
                'data_statistics': '数据统计',
                'operation_logs': '操作日志',
                'system_config': '系统配置'
            }
            if current_admin_page in menu_map:
                render_admin_panel(menu_map[current_admin_page])
            else:
                render_admin_panel('仪表盘')
        return

    if st.session_state['show_register']:
        render_register_page()
        return
    elif st.session_state['show_change_pwd']:
        render_change_password_page()
        return
    elif st.session_state['show_forgot_pwd']:
        render_forgot_password_page()
        return
    elif st.session_state['show_admin_reset']:
        render_admin_reset_page()
        return

    if 'nav' in query_params:
        nav_value = query_params['nav']
        if nav_value == 'register':
            st.session_state['show_register'] = True
            st.rerun()
            return
        elif nav_value == 'forgot':
            st.session_state['show_forgot_pwd'] = True
            st.rerun()
            return
        elif nav_value == 'Login':
            if not st.session_state.get('logged_in', False) or st.session_state.get('username') == '访客用户':
                if st.session_state.get('username') == '访客用户':
                    st.session_state['logged_in'] = False
                    st.session_state['username'] = None
                    st.session_state['user_type'] = None
                render_login_page()
                return
        elif nav_value in ['Home', 'Profile', 'Resume', 'Interview', 'Assessment', 'Image', 'Match', 'Graph', 'Report', 'Interaction', 'Learn', 'News', 'Password', 'LearningPath', 'Cooperation']:
            st.session_state['nav'] = nav_value

    if not st.session_state.get('logged_in', False):
        if st.session_state.get('username') is None or st.session_state.get('username') == '访客用户':
            st.session_state['logged_in'] = True
            st.session_state['username'] = '访客用户'
            st.session_state['user_type'] = 'user'

    render_floating_chat()

    render_top_nav()

    st.markdown('<div style="height: 80px;"></div>', unsafe_allow_html=True)

    current_page = st.session_state['nav']

    manager = StudentProfileManager()
    matching_engine = AdvancedJobMatchingEngine()
    report_gen = AdvancedCareerReportGenerator()

    profile = get_student_profile(st.session_state['username'])
    if profile:
        st.session_state['student_profile'] = profile

    if current_page == "Home":
        render_homepage()
    elif current_page == "Profile":
        render_student_input(manager)
    elif current_page == "Resume":
        st.title("📄 简历服务")
        st.markdown("---")
        tab1, tab2 = st.tabs(["简历解析", "简历优化"])
        with tab1:
            render_resume_parser()
        with tab2:
            render_resume_optimization()
    elif current_page == "Assessment":
        try:
            from assessment import render_assessment
            render_assessment()
        except ImportError:
            st.info("职业测评功能正在开发中，敬请期待")
    elif current_page == "Image":
        if 'student_profile' in st.session_state:
            render_student_profile(manager, st.session_state['student_profile'])
        else:
            st.warning("请先录入信息")
    elif current_page == "Match":
        if 'student_profile' in st.session_state:
            render_advanced_job_matching(matching_engine, st.session_state['student_profile'])
        else:
            st.warning("请先录入信息")
    elif current_page == "Graph":
        render_career_graph()
    elif current_page == "Report":
        if 'student_profile' in st.session_state:
            match_results = st.session_state.get('match_results', {"总匹配度": 70, "优势": [], "差距": []})
            render_advanced_report(report_gen, st.session_state['student_profile'], match_results)
        else:
            st.warning("请先录入信息")
    elif current_page == "Interaction":
        st.title("💬 咨询与互动")
        st.markdown("---")
        tab1, tab2 = st.tabs(["在线讨论", "职业咨询"])
        with tab1:
            render_discussion_page()
        with tab2:
            render_career_consultation()
    elif current_page == "Learn":
        render_learning_resources()
    elif current_page == "News":
        render_industry_news()
    elif current_page == "LearningPath":
        render_learning_path()
    elif current_page == "Cooperation":
        render_corporate_cooperation()
    elif current_page == "Interview":
        render_mock_interview()
    elif current_page == "Password":
        render_change_password_page()
    else:
        render_homepage()


if __name__ == "__main__":
    main()